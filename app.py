import os
import shutil
import uuid
import json
import zipfile
import re
from datetime import datetime
from urllib.parse import urlparse

from modules.env_loader import load_dotenv_if_present

load_dotenv_if_present(os.path.dirname(__file__))
from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    send_file,
    send_from_directory,
    abort,
    jsonify,
    session,
    g,
    flash,
)
from werkzeug.utils import secure_filename
from functools import wraps

from modules.rbac_store import (
    RBACConfigError,
    ROLE_ADMIN,
    ROLE_EDITOR,
    ROLE_LABELS_ZH,
    PERM_USER_MANAGE,
    authenticate,
    build_mssql_engine_from_env,
    create_user,
    ensure_schema,
    get_user_by_id,
    get_user_roles,
    list_users,
    seed_defaults,
    set_user_active,
    set_user_password,
    set_user_role,
    user_has_permission,
)

def _optional_dependency_stub(feature: str):
    def _stub(*_args, **_kwargs):
        raise RuntimeError(
            f"{feature} requires optional document-processing dependencies "
            "(e.g. spire.doc / python-docx / PyMuPDF)."
        )

    return _stub


try:
    from modules.workflow import SUPPORTED_STEPS, run_workflow
except Exception:  # optional dependency (spire.doc) may be missing
    SUPPORTED_STEPS = {}
    run_workflow = _optional_dependency_stub("Workflow execution")

try:
    from modules.Extract_AllFile_to_FinalWord import (
        center_table_figure_paragraphs,
        apply_basic_style,
        remove_hidden_runs,
        hide_paragraphs_with_text,
        remove_paragraphs_with_text,
    )
except Exception:  # optional dependencies may be missing
    center_table_figure_paragraphs = _optional_dependency_stub("center_table_figure_paragraphs")
    apply_basic_style = _optional_dependency_stub("apply_basic_style")
    remove_hidden_runs = _optional_dependency_stub("remove_hidden_runs")
    hide_paragraphs_with_text = _optional_dependency_stub("hide_paragraphs_with_text")
    remove_paragraphs_with_text = _optional_dependency_stub("remove_paragraphs_with_text")

try:
    from modules.Edit_Word import renumber_figures_tables_file
except Exception:
    renumber_figures_tables_file = _optional_dependency_stub("renumber_figures_tables_file")

try:
    from modules.translate_with_bedrock import translate_file
except Exception:
    translate_file = _optional_dependency_stub("translate_file")

from modules.file_copier import copy_files

app = Flask(__name__, instance_relative_config=True)
app.config["SECRET_KEY"] = "dev-secret"
BASE_DIR = os.path.dirname(__file__)
app.config["OUTPUT_FOLDER"] = os.path.join(BASE_DIR, "output")
app.config["TASK_FOLDER"] = os.path.join(BASE_DIR, "tasks")
app.config["ALLOWED_SOURCE_ROOTS"] = []


def parse_bool(value, default=False):
    if value is None:
        return default
    return str(value).strip().lower() in {"1", "true", "yes", "y", "on"}


app.config["AUTH_ENABLED"] = parse_bool(os.environ.get("AUTH_ENABLED"), True)


def get_rbac_engine():
    engine = getattr(app, "_rbac_engine", None)
    if engine is not None:
        return engine
    engine = build_mssql_engine_from_env()
    app._rbac_engine = engine
    return engine


def permission_required(permission_name: str):
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            if not app.config.get("AUTH_ENABLED", True):
                return func(*args, **kwargs)
            if not getattr(g, "user", None):
                abort(401)
            engine = get_rbac_engine()
            if not user_has_permission(engine, g.user.id, permission_name):
                abort(403)
            return func(*args, **kwargs)

        return wrapper

    return decorator


@app.context_processor
def inject_auth_context():
    def _has_perm(perm: str) -> bool:
        if not app.config.get("AUTH_ENABLED", True):
            return True
        if not getattr(g, "user", None):
            return False
        try:
            return user_has_permission(get_rbac_engine(), g.user.id, perm)
        except RBACConfigError:
            return False

    def _role_labels(role_names: list[str]) -> str:
        return "、".join(ROLE_LABELS_ZH.get(name, name) for name in role_names)

    return {
        "auth_enabled": app.config.get("AUTH_ENABLED", True),
        "current_user": getattr(g, "user", None),
        "current_user_roles": getattr(g, "user_roles", []),
        "has_permission": _has_perm,
        "role_labels": _role_labels,
        "ROLE_ADMIN": ROLE_ADMIN,
        "ROLE_EDITOR": ROLE_EDITOR,
        "PERM_USER_MANAGE": PERM_USER_MANAGE,
    }


@app.before_request
def enforce_login():
    g.user = None
    g.user_roles = []
    if not app.config.get("AUTH_ENABLED", True):
        return

    public_endpoints = {"login", "logout", "static"}
    if request.endpoint in public_endpoints:
        return

    user_id = session.get("user_id")
    if user_id:
        try:
            engine = get_rbac_engine()
            user = get_user_by_id(engine, int(user_id))
            if user and user.is_active:
                g.user = user
                g.user_roles = get_user_roles(engine, user.id)
                return
        except RBACConfigError:
            g.user = None
        except Exception:
            g.user = None

    next_url = request.full_path
    if next_url.endswith("?"):
        next_url = next_url[:-1]
    return redirect(url_for("login", next=next_url))


@app.route("/login", methods=["GET", "POST"])
def login():
    if not app.config.get("AUTH_ENABLED", True):
        return redirect(url_for("tasks"))

    if session.get("user_id"):
        return redirect(url_for("tasks"))

    error = ""
    if request.method == "POST":
        username = request.form.get("username", "")
        password = request.form.get("password", "")
        try:
            user = authenticate(get_rbac_engine(), username=username, password=password)
        except RBACConfigError as exc:
            error = str(exc)
            user = None
        except Exception:
            app.logger.exception("登入驗證失敗")
            error = "登入時發生錯誤，請稍後再試或聯絡系統管理員"
            user = None

        if user:
            session["user_id"] = user.id
            flash("登入成功", "success")
            next_url = request.args.get("next")
            if next_url:
                parsed = urlparse(next_url)
                if parsed.scheme or parsed.netloc or not next_url.startswith("/"):
                    next_url = None
            return redirect(next_url or url_for("tasks"))
        if not error:
            error = "帳號或密碼錯誤，或帳號已停用"

    return render_template("login.html", error=error)


@app.get("/logout")
def logout():
    session.clear()
    flash("已登出", "info")
    return redirect(url_for("login"))


@app.route("/admin/users", methods=["GET", "POST"])
@permission_required(PERM_USER_MANAGE)
def admin_users():
    engine = get_rbac_engine()

    message = ""
    if request.method == "POST":
        action = request.form.get("action", "")
        try:
            if action == "create":
                username = request.form.get("username", "")
                password = request.form.get("password", "")
                role = request.form.get("role", ROLE_EDITOR)
                ensure_schema(engine)
                seed_defaults(engine)
                create_user(engine, username=username, password=password, role=role)
                message = f"已新增使用者 {username}"
            elif action == "set_role":
                user_id = int(request.form.get("user_id", "0"))
                role = request.form.get("role", ROLE_EDITOR)
                set_user_role(engine, user_id=user_id, role=role)
                message = "已更新角色"
            elif action == "set_active":
                user_id = int(request.form.get("user_id", "0"))
                is_active = parse_bool(request.form.get("is_active"), True)
                set_user_active(engine, user_id=user_id, is_active=is_active)
                message = "已更新狀態"
            elif action == "reset_password":
                user_id = int(request.form.get("user_id", "0"))
                new_password = request.form.get("new_password", "")
                if not new_password:
                    raise ValueError("請輸入新密碼")
                set_user_password(engine, user_id=user_id, new_password=new_password)
                message = "已重設密碼"
        except Exception as exc:
            message = str(exc)

    ensure_schema(engine)
    seed_defaults(engine)
    users_data = list_users(engine)
    return render_template(
        "admin_users.html",
        users=users_data,
        message=message,
        role_labels_map=ROLE_LABELS_ZH,
    )


@app.errorhandler(403)
def forbidden(_exc):
    return render_template("403.html"), 403


@app.cli.command("init-rbac")
def init_rbac_command():
    """初始化 RBAC 相關資料表並建立預設角色/權限。"""
    engine = get_rbac_engine()
    ensure_schema(engine)
    seed_defaults(engine)
    print("RBAC schema ready. (roles/permissions seeded)")


@app.cli.command("create-user")
def create_user_command():
    """互動式建立使用者（需先 init-rbac）。"""
    import getpass

    engine = get_rbac_engine()
    ensure_schema(engine)
    seed_defaults(engine)

    username = input("Username: ").strip()
    role = input("Role (admin/editor) [editor]: ").strip() or ROLE_EDITOR
    password = getpass.getpass("Password: ")
    user_id = create_user(engine, username=username, password=password, role=role)
    print(f"Created user id={user_id}")

nas_roots_env = os.environ.get("ALLOWED_NAS_ROOTS")
nas_allowed_roots = [
    os.path.abspath(p)
    for p in nas_roots_env.split(os.pathsep)
    if p.strip()
]
app.config["ALLOWED_NAS_ROOTS"] = nas_allowed_roots
app.config["NAS_ALLOWED_ROOTS"] = nas_allowed_roots
app.config["NAS_ALLOW_RECURSIVE"] = parse_bool(
    os.environ.get("NAS_ALLOW_RECURSIVE"), True
)
max_copy_size_mb = os.environ.get("NAS_MAX_COPY_FILE_SIZE_MB")
try:
    app.config["NAS_MAX_COPY_FILE_SIZE"] = (
        int(max_copy_size_mb) * 1024 * 1024 if max_copy_size_mb else 500 * 1024 * 1024
    )
except ValueError:
    app.config["NAS_MAX_COPY_FILE_SIZE"] = 500 * 1024 * 1024

os.makedirs(app.config["OUTPUT_FOLDER"], exist_ok=True)
os.makedirs(app.config["TASK_FOLDER"], exist_ok=True)

ALLOWED_DOCX = {".docx"}
ALLOWED_PDF = {".pdf"}
ALLOWED_ZIP = {".zip"}

DOCUMENT_FORMAT_PRESETS = {
    "none": {
        "label": "無（保留原文件格式）",
        "western_font": "",
        "east_asian_font": "",
        "font_size": 0,
    },
    "default": {
        "label": "Times New Roman / 新細明體（12 pt）",
        "western_font": "Times New Roman",
        "east_asian_font": "新細明體",
        "font_size": 12,
        "space_before": 6,
        "space_after": 6,
    },
    "modern": {
        "label": "Calibri / 微軟正黑體（12 pt）",
        "western_font": "Calibri",
        "east_asian_font": "微軟正黑體",
        "font_size": 12,
        "space_before": 6,
        "space_after": 6,
    },
}

DEFAULT_DOCUMENT_FORMAT_KEY = "default"
DEFAULT_LINE_SPACING = 1.5
LINE_SPACING_CHOICES = [
    ("none", "無（保留原行距）"),
    ("1", "單行（1.0）"),
    ("1.15", "1.15 倍行距"),
    ("1.5", "1.5 倍行距"),
    ("2", "雙行（2.0）"),
]
DEFAULT_APPLY_FORMATTING = False


def normalize_document_format(key: str) -> str:
    if not key or key not in DOCUMENT_FORMAT_PRESETS:
        return DEFAULT_DOCUMENT_FORMAT_KEY
    return key


def coerce_line_spacing(value) -> float:
    if isinstance(value, str) and value.strip().lower() == "none":
        return DEFAULT_LINE_SPACING
    try:
        spacing = float(value)
        if spacing <= 0:
            raise ValueError
        return spacing
    except (TypeError, ValueError):
        return DEFAULT_LINE_SPACING


def allowed_file(filename, kinds=("docx", "pdf", "zip")):
    ext = os.path.splitext(filename)[1].lower()
    if "docx" in kinds and ext in ALLOWED_DOCX:
        return True
    if "pdf" in kinds and ext in ALLOWED_PDF:
        return True
    if "zip" in kinds and ext in ALLOWED_ZIP:
        return True
    return False


def load_allowed_roots_from_env():
    roots = []
    raw = os.environ.get("TASK_ALLOWED_ROOTS", "")
    for entry in raw.split(os.path.pathsep):
        candidate = entry.strip()
        if not candidate:
            continue
        abs_path = os.path.abspath(candidate)
        if os.path.isdir(abs_path):
            roots.append(abs_path)
    return roots


def ensure_allowed_roots_loaded():
    if not app.config.get("ALLOWED_SOURCE_ROOTS"):
        loaded_roots = load_allowed_roots_from_env()
        if loaded_roots:
            app.config["ALLOWED_SOURCE_ROOTS"] = loaded_roots
        elif app.config.get("NAS_ALLOWED_ROOTS"):
            app.config["ALLOWED_SOURCE_ROOTS"] = list(app.config["NAS_ALLOWED_ROOTS"])


def normalize_relative_path(raw_path: str, allow_recursive: bool) -> str:
    if not raw_path or not raw_path.strip():
        raise ValueError("請提供要匯入的檔案或資料夾路徑")
    cleaned = raw_path.strip().replace("\\", "/")
    # Make POSIX-style absolute paths invalid on Windows too (e.g. "/abs/path").
    if cleaned.startswith("/"):
        raise ValueError("路徑不可為絕對路徑，請填寫相對於允許根目錄的路徑")
    if os.path.isabs(cleaned):
        raise ValueError("路徑不可為絕對路徑，請填寫相對於允許根目錄的路徑")
    norm_rel = os.path.normpath(cleaned)
    if norm_rel.startswith(".."):
        raise ValueError("路徑不可包含 .. 或跳脫允許的根目錄")
    if not allow_recursive and os.sep in norm_rel:
        raise ValueError("目前僅允許存取根層級的項目")
    return norm_rel


def validate_nas_path(raw_path: str, allowed_roots=None, allow_recursive=None):
    allow_recursive = (
        app.config.get("NAS_ALLOW_RECURSIVE", True)
        if allow_recursive is None
        else allow_recursive
    )
    norm_rel = normalize_relative_path(raw_path, allow_recursive)
    ensure_allowed_roots_loaded()
    allowed_roots = (
        allowed_roots
        or app.config.get("NAS_ALLOWED_ROOTS")
        or app.config.get("ALLOWED_SOURCE_ROOTS", [])
    )
    if not allowed_roots:
        raise ValueError("尚未設定允許的根目錄，請聯絡系統管理員")
    for root in allowed_roots:
        root_abs = os.path.abspath(root)
        candidate = os.path.abspath(os.path.join(root_abs, norm_rel))
        try:
            if os.path.commonpath([root_abs, candidate]) != root_abs:
                continue
        except ValueError:
            continue
        if os.path.exists(candidate):
            return candidate
    raise FileNotFoundError("找不到指定的路徑，或不在允許的根目錄內")


def get_configured_nas_roots() -> list[str]:
    ensure_allowed_roots_loaded()
    roots = app.config.get("NAS_ALLOWED_ROOTS") or app.config.get("ALLOWED_SOURCE_ROOTS", [])
    return list(roots) if roots else []


def resolve_nas_path_in_root(raw_path: str, root_index: int, allow_recursive=None) -> str:
    allow_recursive = (
        app.config.get("NAS_ALLOW_RECURSIVE", True)
        if allow_recursive is None
        else allow_recursive
    )
    roots = get_configured_nas_roots()
    if not roots:
        raise ValueError("NAS roots are not configured")
    if root_index < 0 or root_index >= len(roots):
        raise ValueError("Invalid NAS root index")

    root_abs = os.path.abspath(roots[root_index])
    norm_rel = normalize_relative_path(raw_path, allow_recursive)
    candidate = os.path.abspath(os.path.join(root_abs, norm_rel))
    try:
        if os.path.commonpath([root_abs, candidate]) != root_abs:
            raise ValueError("Path escapes the allowed NAS root")
    except ValueError:
        raise ValueError("Invalid path")

    if os.path.exists(candidate):
        return candidate
    raise FileNotFoundError("Path does not exist in the selected NAS root")


def resolve_nas_path(raw_path: str, allowed_roots=None, allow_recursive=None, root_index=None) -> str:
    if root_index is None or str(root_index).strip() == "":
        return validate_nas_path(raw_path, allowed_roots=allowed_roots, allow_recursive=allow_recursive)
    try:
        root_index_int = int(root_index)
    except (TypeError, ValueError):
        raise ValueError("Invalid NAS root index")
    return resolve_nas_path_in_root(raw_path, root_index_int, allow_recursive=allow_recursive)


def deduplicate_name(base_dir: str, name: str) -> str:
    candidate = name
    stem, ext = os.path.splitext(name)
    counter = 1
    while os.path.exists(os.path.join(base_dir, candidate)):
        candidate = f"{stem} ({counter}){ext}"
        counter += 1
    return candidate


def ensure_windows_long_path(path: str) -> str:
    """Add the Windows long-path prefix to avoid MAX_PATH issues."""
    if os.name != "nt" or not path:
        return path
    normalized = os.path.abspath(path)
    if normalized.startswith("\\\\?\\"):
        return normalized
    if normalized.startswith("\\\\"):
        return "\\\\?\\UNC\\" + normalized[2:]
    return "\\\\?\\" + normalized


def enforce_max_copy_size(path: str):
    max_bytes = app.config.get("NAS_MAX_COPY_FILE_SIZE")
    if not max_bytes:
        return
    checked_path = ensure_windows_long_path(path)

    def _check(target: str):
        try:
            return os.path.getsize(target)
        except OSError:
            return 0

    if os.path.isfile(checked_path):
        if _check(checked_path) > max_bytes:
            raise ValueError("檔案超過允許的大小限制，請分批處理或聯絡系統管理員")
        return

    for root, _, files in os.walk(checked_path):
        for fn in files:
            fpath = os.path.join(root, fn)
            if _check(fpath) > max_bytes:
                app.logger.warning("檔案大小超過限制：%s", fpath)
                raise ValueError("檔案超過允許的大小限制，請分批處理或聯絡系統管理員")


@app.get("/api/nas/dirs")
def api_nas_list_dirs():
    """List sub-directories under a configured NAS root.

    Query params:
      - root_index: int (required)
      - path: str (optional, relative path; empty means the root itself)
    """
    root_index = request.args.get("root_index", type=int)
    if root_index is None:
        return jsonify({"error": "root_index is required"}), 400

    rel_path_raw = (request.args.get("path") or "").strip()
    allow_recursive = app.config.get("NAS_ALLOW_RECURSIVE", True)

    try:
        roots = get_configured_nas_roots()
        if not roots:
            return jsonify({"error": "NAS roots are not configured"}), 400
        if root_index < 0 or root_index >= len(roots):
            return jsonify({"error": "Invalid NAS root index"}), 400

        root_abs = os.path.abspath(roots[root_index])
        if rel_path_raw in {"", ".", "/"}:
            abs_dir = root_abs
            rel_path = ""
        else:
            rel_path = normalize_relative_path(rel_path_raw, allow_recursive=allow_recursive).replace("\\", "/")
            abs_dir = resolve_nas_path_in_root(rel_path, root_index, allow_recursive=allow_recursive)

        if not os.path.isdir(abs_dir):
            return jsonify({"error": "Path is not a directory"}), 400

        dirs = []
        for name in sorted(os.listdir(abs_dir), key=str.lower):
            full = os.path.join(abs_dir, name)
            if os.path.isdir(full):
                child_rel = f"{rel_path}/{name}" if rel_path else name
                dirs.append({"name": name, "path": child_rel.replace("\\", "/")})

        parent = None
        if rel_path:
            parent_parts = rel_path.split("/")
            parent = "/".join(parent_parts[:-1]) if len(parent_parts) > 1 else ""

        return jsonify(
            {
                "root_index": root_index,
                "path": rel_path,
                "parent": parent,
                "dirs": dirs,
                "allow_recursive": bool(allow_recursive),
            }
        )
    except (ValueError, FileNotFoundError) as exc:
        return jsonify({"error": str(exc)}), 400
    except PermissionError:
        return jsonify({"error": "Permission denied"}), 403


def list_files(base_dir):
    files = []
    for root, _, fns in os.walk(base_dir):
        for fn in fns:
            rel = os.path.relpath(os.path.join(root, fn), base_dir)
            files.append(rel)
    return sorted(files)


def build_file_tree(base_dir):
    tree = {"dirs": {}, "files": []}
    for root, dirs, files in os.walk(base_dir):
        rel = os.path.relpath(root, base_dir)
        node = tree
        if rel != ".":
            for part in rel.split(os.sep):
                node = node["dirs"].setdefault(part, {"dirs": {}, "files": []})
        node["files"].extend(sorted(files))
    return tree


def list_dirs(base_dir):
    dirs = []
    for root, dirnames, _ in os.walk(base_dir):
        rel_root = os.path.relpath(root, base_dir)
        for d in dirnames:
            path = os.path.normpath(os.path.join(rel_root, d))
            dirs.append(path)
    return sorted(dirs)


def collect_titles_to_hide(entries):
    titles = []
    seen = set()
    if not isinstance(entries, list):
        return titles
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        captured = entry.get("captured_titles")
        if not captured:
            result_meta = entry.get("result")
            if isinstance(result_meta, dict):
                captured = result_meta.get("captured_titles")
        if not captured:
            continue
        for title in captured:
            if not isinstance(title, str):
                continue
            trimmed = title.strip()
            normalized = " ".join(trimmed.split())
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            titles.append(trimmed)
    return titles


def load_titles_to_hide_from_log(job_dir):
    log_path = os.path.join(job_dir, "log.json")
    if not os.path.exists(log_path):
        return []
    try:
        with open(log_path, "r", encoding="utf-8") as f:
            entries = json.load(f)
        return collect_titles_to_hide(entries)
    except Exception:
        return []


def clean_compare_html_content(html_content):
    html_content = re.sub(
        r'<(\w+)[^>]*style="[^"]*display\s*:\s*none[^"]*"[^>]*>.*?</\1>',
        "",
        html_content,
        flags=re.IGNORECASE | re.DOTALL,
    )
    html_content = re.sub(
        r"<p[^>]*>(?:\s|&nbsp;|&#160;)*</p>",
        "",
        html_content,
        flags=re.IGNORECASE,
    )
    return html_content


def save_compare_output(
    job_dir,
    html_content,
    titles_to_hide,
    base_name="result",
    subdir=None,
):
    target_dir = job_dir if not subdir else os.path.join(job_dir, subdir)
    os.makedirs(target_dir, exist_ok=True)
    html_path = os.path.join(target_dir, f"{base_name}.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    from spire.doc import Document, FileFormat

    doc = Document()
    doc.LoadFromFile(html_path, FileFormat.Html)
    doc.SaveToFile(os.path.join(target_dir, f"{base_name}.docx"), FileFormat.Docx)
    doc.Close()
    result_docx = os.path.join(target_dir, f"{base_name}.docx")
    remove_hidden_runs(result_docx, preserve_texts=titles_to_hide)
    apply_basic_style(result_docx)
    hide_paragraphs_with_text(result_docx, titles_to_hide)
    return html_path, result_docx


def load_version_metadata(versions_dir):
    metadata = {"versions": []}
    if not os.path.isdir(versions_dir):
        return metadata
    meta_path = os.path.join(versions_dir, "metadata.json")
    if not os.path.exists(meta_path):
        return metadata
    try:
        with open(meta_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, dict) and isinstance(data.get("versions"), list):
            metadata = data
    except Exception:
        metadata = {"versions": []}
    return metadata


def save_version_metadata(versions_dir, metadata):
    os.makedirs(versions_dir, exist_ok=True)
    meta_path = os.path.join(versions_dir, "metadata.json")
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)


def sanitize_version_slug(name):
    if not name:
        return "version"
    slug = re.sub(r"[^\w\-]+", "_", name.strip(), flags=re.UNICODE)
    slug = slug.strip("_")
    if not slug:
        slug = "version"
    return slug[:60]


def build_version_context(task_id, job_id, job_dir):
    versions_dir = os.path.join(job_dir, "versions")
    metadata = load_version_metadata(versions_dir)
    context = []
    versions = metadata.get("versions", [])
    for item in sorted(versions, key=lambda v: v.get("created_at", ""), reverse=True):
        version_id = item.get("id")
        base_name = item.get("base_name")
        if not version_id or not base_name:
            continue
        html_rel = f"versions/{base_name}.html"
        docx_rel = os.path.join(versions_dir, f"{base_name}.docx")
        html_abs = os.path.join(versions_dir, f"{base_name}.html")
        if not os.path.exists(docx_rel) or not os.path.exists(html_abs):
            continue
        created_display = item.get("created_at", "")
        created_at = item.get("created_at")
        if created_at:
            try:
                created_display = datetime.fromisoformat(created_at).strftime(
                    "%Y-%m-%d %H:%M:%S"
                )
            except ValueError:
                created_display = created_at
        context.append(
            {
                "id": version_id,
                "name": item.get("name") or version_id,
                "created_at_display": created_display,
                "html_url": url_for(
                    "task_view_file",
                    task_id=task_id,
                    job_id=job_id,
                    filename=html_rel,
                ),
                "docx_url": url_for(
                    "task_download_version",
                    task_id=task_id,
                    job_id=job_id,
                    version_id=version_id,
                ),
                "restore_url": url_for(
                    "task_compare_restore_version",
                    task_id=task_id,
                    job_id=job_id,
                    version_id=version_id,
                ),
                "delete_url": url_for(
                    "task_compare_delete_version",
                    task_id=task_id,
                    job_id=job_id,
                    version_id=version_id,
                ),
            }
        )
    return context


def task_name_exists(name, exclude_id=None):
    for tid in os.listdir(app.config["TASK_FOLDER"]):
        if exclude_id and tid == exclude_id:
            continue
        tdir = os.path.join(app.config["TASK_FOLDER"], tid)
        if not os.path.isdir(tdir):
            continue
        meta_path = os.path.join(tdir, "meta.json")
        tname = tid
        if os.path.exists(meta_path):
            with open(meta_path, "r", encoding="utf-8") as f:
                tname = json.load(f).get("name", tid)
        if tname == name:
            return True
    return False

@app.route("/tasks/<task_id>/copy-files", methods=["GET", "POST"], endpoint="task_copy_files")
def task_copy_files(task_id):
    base = os.path.join(app.config["TASK_FOLDER"], task_id, "files")
    if not os.path.isdir(base):
        abort(404)

    def _safe_path(rel: str) -> str:
        norm = os.path.normpath(rel)
        if not rel or os.path.isabs(norm) or norm.startswith(".."):
            raise ValueError("資料夾名稱不合法")
        return os.path.join(base, norm)

    message = ""
    if request.method == "POST":
        action = request.form.get("action")
        if action == "create_dir":
            new_rel = request.form.get("new_dir", "").strip()
            try:
                os.makedirs(_safe_path(new_rel), exist_ok=True)
                message = f"已建立資料夾 {os.path.normpath(new_rel)}"
            except ValueError:
                message = "資料夾名稱不合法"
        else:
            source_rel = request.form.get("source_dir", "").strip()
            dest_rel = request.form.get("dest_dir", "").strip()
            keywords_raw = request.form.get("keywords", "")
            keywords = [k.strip() for k in keywords_raw.split(",") if k.strip()]
            if not source_rel or not dest_rel or not keywords:
                message = "請完整輸入資料"
            else:
                try:
                    src = _safe_path(source_rel)
                    dest = _safe_path(dest_rel)
                    copied = copy_files(src, dest, keywords)
                    message = f"已複製 {len(copied)} 個檔案"
                except ValueError:
                    message = "資料夾名稱不合法"
                except Exception as e:
                    message = str(e)
    dirs = list_dirs(base)
    dirs.insert(0, ".")
    return render_template("copy_files.html", dirs=dirs, message=message, task_id=task_id)


@app.route("/tasks/<task_id>/mapping", methods=["GET", "POST"], endpoint="task_mapping")
def task_mapping(task_id):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    if not os.path.isdir(tdir):
        abort(404)
    files_dir = os.path.join(tdir, "files")
    out_dir = os.path.join(app.config["OUTPUT_FOLDER"], task_id)
    messages = []
    outputs = []
    if request.method == "POST":
        f = request.files.get("mapping_file")
        if not f or not f.filename:
            messages.append("請選擇檔案")
        else:
            path = os.path.join(tdir, secure_filename(f.filename))
            f.save(path)
            try:
                from modules.mapping_processor import process_mapping_excel
                result = process_mapping_excel(path, files_dir, out_dir)
                messages = result["logs"]
                outputs = result["outputs"]
            except Exception as e:
                messages = [str(e)]
    rel_outputs = [os.path.basename(p) for p in outputs]
    return render_template("mapping.html", task_id=task_id, messages=messages, outputs=rel_outputs)


@app.get("/tasks/<task_id>/output/<filename>")
def task_download_output(task_id, filename):
    out_dir = os.path.join(app.config["OUTPUT_FOLDER"], task_id)
    file_path = os.path.join(out_dir, filename)
    if not os.path.isfile(file_path):
        abort(404)
    return send_from_directory(out_dir, filename, as_attachment=True)


@app.get("/")
def tasks():
    task_list = []
    for tid in os.listdir(app.config["TASK_FOLDER"]):
        tdir = os.path.join(app.config["TASK_FOLDER"], tid)
        if os.path.isdir(tdir):
            meta_path = os.path.join(tdir, "meta.json")
            name = tid
            description = ""
            created = None
            if os.path.exists(meta_path):
                with open(meta_path, "r", encoding="utf-8") as f:
                    meta = json.load(f)
                    name = meta.get("name", tid)
                    description = meta.get("description", "")
                    created = meta.get("created")
            if not created:
                created = datetime.fromtimestamp(os.path.getmtime(tdir)).strftime("%Y-%m-%d %H:%M")
            task_list.append({"id": tid, "name": name, "description": description, "created": created})
    task_list.sort(key=lambda x: x["created"], reverse=True)
    return render_template(
        "tasks.html",
        tasks=task_list,
        allowed_nas_roots=app.config.get("ALLOWED_NAS_ROOTS", []),
    )


@app.post("/tasks")
def create_task():
    def _fail(message: str):
        flash(message, "danger")
        return redirect(url_for("tasks"))

    nas_path = request.form.get("nas_path", "")
    try:
        nas_root_index = request.form.get("nas_root_index", "").strip()
        resolved_path = resolve_nas_path(
            nas_path,
            allowed_roots=app.config.get("NAS_ALLOWED_ROOTS"),
            allow_recursive=app.config.get("NAS_ALLOW_RECURSIVE", True),
            root_index=nas_root_index or None,
        )
        if not os.path.isdir(resolved_path):
            return _fail("指定的 NAS 路徑不是資料夾")
        enforce_max_copy_size(resolved_path)
    except ValueError as exc:
        return _fail(str(exc))
    except FileNotFoundError as exc:
        return _fail(str(exc))
    task_name = request.form.get("task_name", "").strip() or "未命名任務"
    task_desc = request.form.get("task_desc", "").strip()
    if task_name_exists(task_name):
        return _fail("任務名稱已存在")
    tid = str(uuid.uuid4())[:8]
    tdir = os.path.join(app.config["TASK_FOLDER"], tid)
    files_dir = os.path.join(tdir, "files")
    os.makedirs(files_dir, exist_ok=True)
    src_dir = ensure_windows_long_path(resolved_path)
    dest_dir = ensure_windows_long_path(files_dir)
    try:
        shutil.copytree(src_dir, dest_dir, dirs_exist_ok=True)
    except PermissionError:
        shutil.rmtree(tdir, ignore_errors=True)
        return _fail("沒有足夠的權限讀取或複製指定路徑")
    except shutil.Error as exc:
        app.logger.exception("複製 NAS 目錄失敗")
        shutil.rmtree(tdir, ignore_errors=True)
        detail = ""
        if exc.args and isinstance(exc.args[0], list) and exc.args[0]:
            first_error = exc.args[0][0]
            if len(first_error) >= 3:
                detail = f"：{first_error[2]}"
        return _fail(f"複製 NAS 目錄時發生錯誤{detail or ''}，請稍後再試")
    except Exception:
        app.logger.exception("複製 NAS 目錄失敗")
        shutil.rmtree(tdir, ignore_errors=True)
        return _fail("複製 NAS 目錄時發生錯誤，請稍後再試")
    with open(os.path.join(tdir, "meta.json"), "w", encoding="utf-8") as meta:
        json.dump(
            {
                "name": task_name,
                "description": task_desc,
                "created" : datetime.now().strftime("%Y-%m-%d %H:%M"),
            },
            meta,
            ensure_ascii=False,
            indent=2,
        )
    return redirect(url_for("task_detail", task_id=tid))


@app.post("/tasks/<task_id>/delete")
def delete_task(task_id):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    if os.path.isdir(tdir):
        import shutil
        shutil.rmtree(tdir)
    return redirect(url_for("tasks"))


@app.post("/tasks/<task_id>/rename")
def rename_task(task_id):
    new_name = request.form.get("name", "").strip()
    if not new_name:
        return "缺少名稱", 400
    if task_name_exists(new_name, exclude_id=task_id):
        return "任務名稱已存在", 400
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    if not os.path.isdir(tdir):
        abort(404)
    meta_path = os.path.join(tdir, "meta.json")
    meta = {}
    if os.path.exists(meta_path):
        with open(meta_path, "r", encoding="utf-8") as f:
            meta = json.load(f)
    meta["name"] = new_name
    if "created" not in meta:
        meta["created"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)
    return redirect(url_for("tasks"))


@app.get("/tasks/<task_id>")
def task_detail(task_id):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    files_dir = os.path.join(tdir, "files")
    if not os.path.isdir(files_dir):
        abort(404)
    meta_path = os.path.join(tdir, "meta.json")
    name = task_id
    description = ""
    if os.path.exists(meta_path):
        with open(meta_path, "r", encoding="utf-8") as f:
            meta = json.load(f)
            name = meta.get("name", task_id)
            description = meta.get("description", "")
    tree = build_file_tree(files_dir)
    return render_template(
        "task_detail.html",
        task={"id": task_id, "name": name, "description": description},
        files_tree=tree,
    )


@app.post("/tasks/<task_id>/files")
def upload_task_file(task_id):
    """Upload additional files to an existing task."""
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    files_dir = os.path.join(tdir, "files")
    if not os.path.isdir(files_dir):
        abort(404)

    nas_input = request.form.get("nas_file_path", "").strip()
    try:
        source_path = validate_nas_path(
            nas_input,
            allowed_roots=app.config.get("ALLOWED_SOURCE_ROOTS", []),
        )
        enforce_max_copy_size(source_path)
    except ValueError as e:
        return str(e), 400
    except FileNotFoundError as e:
        return str(e), 404

    try:
        source_path = ensure_windows_long_path(source_path)
        if os.path.isdir(source_path):
            dest_name = deduplicate_name(files_dir, os.path.basename(source_path))
            dest_path = ensure_windows_long_path(os.path.join(files_dir, dest_name))
            shutil.copytree(source_path, dest_path)
        else:
            if not allowed_file(source_path):
                return "僅支援 DOCX、PDF 或 ZIP 檔案，或複製整個資料夾", 400
            dest_name = deduplicate_name(files_dir, os.path.basename(source_path))
            dest_path = ensure_windows_long_path(os.path.join(files_dir, dest_name))
            if dest_name.lower().endswith(".zip"):
                shutil.copy2(source_path, dest_path)
                with zipfile.ZipFile(dest_path, "r") as zf:
                    zf.extractall(files_dir)
                os.remove(dest_path)
            else:
                shutil.copy2(source_path, dest_path)
    except PermissionError:
        return "沒有足夠的權限讀取或複製指定路徑", 400
    except FileNotFoundError:
        return "找不到指定的檔案或資料夾", 404
    except shutil.Error:
        app.logger.exception("複製檔案時發生錯誤")
        return "複製檔案時發生錯誤，請稍後再試", 400
    except Exception:
        app.logger.exception("處理 NAS 檔案時發生未預期錯誤")
        return "處理檔案時發生錯誤，請稍後再試", 400

    return redirect(url_for("task_detail", task_id=task_id))

def gather_available_files(files_dir):
    mapping = {"docx": [], "pdf": [], "zip": [], "dir": []}
    for rel in list_files(files_dir):
        ext = os.path.splitext(rel)[1].lower()
        if ext == ".docx":
            mapping["docx"].append(rel)
        elif ext == ".pdf":
            mapping["pdf"].append(rel)
        elif ext == ".zip":
            mapping["zip"].append(rel)
    dirs = list_dirs(files_dir)
    dirs.insert(0, ".")
    mapping["dir"] = dirs
    return mapping


@app.get("/tasks/<task_id>/flows")
def flow_builder(task_id):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    files_dir = os.path.join(tdir, "files")
    if not os.path.isdir(files_dir):
        abort(404)
    flow_dir = os.path.join(tdir, "flows")
    os.makedirs(flow_dir, exist_ok=True)
    flows = []
    for fn in os.listdir(flow_dir):
        if fn.endswith(".json"):
            path = os.path.join(flow_dir, fn)
            created = datetime.fromtimestamp(os.path.getmtime(path)).strftime("%Y-%m-%d %H:%M")
            has_copy = False
            steps_data = []
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, dict):
                    steps_data = data.get("steps", [])
                    created = data.get("created", created)
                elif isinstance(data, list):
                    steps_data = data
                has_copy = any(
                    isinstance(s, dict) and s.get("type") == "copy_files"
                    for s in steps_data
                )
            except Exception:
                pass
            flows.append(
                {
                    "name": os.path.splitext(fn)[0],
                    "created": created,
                    "has_copy": has_copy,
                }
            )
    preset = None
    center_titles = True
    loaded_name = request.args.get("flow")
    if loaded_name:
        p = os.path.join(flow_dir, f"{loaded_name}.json")
        if os.path.exists(p):
            with open(p, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                steps_data = data.get("steps", [])
                center_titles = data.get("center_titles", True) or any(
                    isinstance(s, dict) and s.get("type") == "center_table_figure_paragraphs" for s in steps_data
                )
            else:
                steps_data = data
                center_titles = True
            preset = [
                s for s in steps_data
                if isinstance(s, dict) and s.get("type") in SUPPORTED_STEPS
            ]
    avail = gather_available_files(files_dir)
    tree = build_file_tree(files_dir)
    return render_template(
        "flow.html",
        task={"id": task_id},
        steps=SUPPORTED_STEPS,
        files=avail,
        flows=flows,
        preset=preset,
        loaded_name=loaded_name,
        center_titles=center_titles,
        files_tree=tree,
    )


@app.post("/tasks/<task_id>/flows/run")
def run_flow(task_id):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    files_dir = os.path.join(tdir, "files")
    if not os.path.isdir(files_dir):
        abort(404)
    action = request.form.get("action", "run")
    flow_name = request.form.get("flow_name", "").strip()
    ordered_ids = request.form.get("ordered_ids", "").split(",")
    center_titles = request.form.get("center_titles") == "on"
    apply_formatting = False
    workflow = []
    for sid in ordered_ids:
        sid = sid.strip()
        if not sid:
            continue
        stype = request.form.get(f"step_{sid}_type", "")
        if not stype or stype not in SUPPORTED_STEPS:
            continue
        schema = SUPPORTED_STEPS.get(stype, {})
        params = {}
        for k in schema.get("inputs", []):
            field = f"step_{sid}_{k}"
            val = request.form.get(field, "")
            params[k] = val
        workflow.append({"type": stype, "params": params})
    flow_dir = os.path.join(tdir, "flows")
    os.makedirs(flow_dir, exist_ok=True)
    if action == "save":
        if not flow_name:
            return "缺少流程名稱", 400
        path = os.path.join(flow_dir, f"{flow_name}.json")
        created = datetime.now().strftime("%Y-%m-%d %H:%M")
        if os.path.exists(path):
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, dict) and "created" in data:
                    created = data["created"]
            except Exception:
                pass
        data = {
            "created": created,
            "steps": workflow,
            "center_titles": center_titles,
        }
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return redirect(url_for("flow_builder", task_id=task_id))

    runtime_steps = []
    for step in workflow:
        stype = step["type"]
        if stype not in SUPPORTED_STEPS:
            continue
        schema = SUPPORTED_STEPS.get(stype, {})
        params = {}
        for k, v in step["params"].items():
            accept = schema.get("accepts", {}).get(k, "text")
            if accept.startswith("file") and v:
                params[k] = os.path.join(files_dir, v)
            else:
                params[k] = v
        runtime_steps.append({"type": stype, "params": params})

    job_id = str(uuid.uuid4())[:8]
    job_dir = os.path.join(tdir, "jobs", job_id)
    os.makedirs(job_dir, exist_ok=True)
    workflow_result = run_workflow(runtime_steps, workdir=job_dir)
    result_path = workflow_result.get("result_docx") or os.path.join(job_dir, "result.docx")
    titles_to_hide = collect_titles_to_hide(workflow_result.get("log_json", []))
    renumber_figures_tables_file(result_path)
    if center_titles:
        center_table_figure_paragraphs(result_path)
    remove_hidden_runs(result_path, preserve_texts=titles_to_hide)
    hide_paragraphs_with_text(result_path, titles_to_hide)
    return redirect(url_for("task_result", task_id=task_id, job_id=job_id))


@app.post("/tasks/<task_id>/flows/execute/<flow_name>")
def execute_flow(task_id, flow_name):
    """Execute a previously saved flow."""
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    files_dir = os.path.join(tdir, "files")
    if not os.path.isdir(files_dir):
        abort(404)
    flow_path = os.path.join(tdir, "flows", f"{flow_name}.json")
    if not os.path.exists(flow_path):
        abort(404)
    with open(flow_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    document_format = DEFAULT_DOCUMENT_FORMAT_KEY
    line_spacing = DEFAULT_LINE_SPACING
    apply_formatting = False
    if isinstance(data, dict):
        workflow = data.get("steps", [])
        center_titles = data.get("center_titles", True) or any(
            isinstance(s, dict) and s.get("type") == "center_table_figure_paragraphs" for s in workflow
        )
        document_format = normalize_document_format(data.get("document_format"))
        line_spacing = coerce_line_spacing(data.get("line_spacing", DEFAULT_LINE_SPACING))
    else:
        workflow = data
        center_titles = True
    runtime_steps = []
    for step in workflow:
        stype = step.get("type")
        if stype not in SUPPORTED_STEPS:
            continue
        schema = SUPPORTED_STEPS.get(stype, {})
        params = {}
        for k, v in step.get("params", {}).items():
            accept = schema.get("accepts", {}).get(k, "text")
            if accept.startswith("file") and v:
                params[k] = os.path.join(files_dir, v)
            else:
                params[k] = v
        runtime_steps.append({"type": stype, "params": params})
    job_id = str(uuid.uuid4())[:8]
    job_dir = os.path.join(tdir, "jobs", job_id)
    os.makedirs(job_dir, exist_ok=True)
    workflow_result = run_workflow(runtime_steps, workdir=job_dir)
    result_path = workflow_result.get("result_docx") or os.path.join(job_dir, "result.docx")
    titles_to_hide = collect_titles_to_hide(workflow_result.get("log_json", []))
    renumber_figures_tables_file(result_path)
    if center_titles:
        center_table_figure_paragraphs(result_path)
    remove_hidden_runs(result_path, preserve_texts=titles_to_hide)
    hide_paragraphs_with_text(result_path, titles_to_hide)
    return redirect(url_for("task_result", task_id=task_id, job_id=job_id))


@app.post("/tasks/<task_id>/flows/update-format/<flow_name>")
def update_flow_format(task_id, flow_name):
    """Update the document formatting metadata for a saved flow."""
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    flow_dir = os.path.join(tdir, "flows")
    flow_path = os.path.join(flow_dir, f"{flow_name}.json")
    if not os.path.exists(flow_path):
        abort(404)

    document_format = normalize_document_format(request.form.get("document_format"))
    line_spacing_raw = request.form.get("line_spacing")
    line_spacing_value = (line_spacing_raw or f"{DEFAULT_LINE_SPACING:g}").strip()
    line_spacing_none = line_spacing_value.lower() == "none"
    line_spacing = DEFAULT_LINE_SPACING if line_spacing_none else coerce_line_spacing(line_spacing_value)
    apply_formatting_param = request.form.get("apply_formatting")
    apply_formatting = parse_bool(apply_formatting_param, DEFAULT_APPLY_FORMATTING)
    if document_format == "none" or line_spacing_none:
        apply_formatting = False

    try:
        with open(flow_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except json.JSONDecodeError:
        return "流程檔案格式錯誤", 400
    except Exception:
        data = {}

    if isinstance(data, dict):
        payload = data
    elif isinstance(data, list):
        payload = {"steps": data}
    else:
        payload = {"steps": []}

    current_apply = parse_bool(payload.get("apply_formatting"), DEFAULT_APPLY_FORMATTING)
    new_apply = apply_formatting if apply_formatting_param is not None else current_apply

    payload["document_format"] = document_format
    payload["line_spacing"] = line_spacing_value
    payload["apply_formatting"] = new_apply

    if "created" not in payload:
        created = datetime.fromtimestamp(os.path.getmtime(flow_path)).strftime("%Y-%m-%d %H:%M")
        payload["created"] = created
    payload.setdefault("center_titles", True)

    with open(flow_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)

    return redirect(url_for("flow_builder", task_id=task_id))


@app.post("/tasks/<task_id>/flows/delete/<flow_name>")
def delete_flow(task_id, flow_name):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    flow_dir = os.path.join(tdir, "flows")
    path = os.path.join(flow_dir, f"{flow_name}.json")
    if os.path.exists(path):
        os.remove(path)
    return redirect(url_for("flow_builder", task_id=task_id))


@app.post("/tasks/<task_id>/flows/rename/<flow_name>")
def rename_flow(task_id, flow_name):
    new_name = request.form.get("name", "").strip()
    if not new_name:
        return "缺少流程名稱", 400
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    flow_dir = os.path.join(tdir, "flows")
    old_path = os.path.join(flow_dir, f"{flow_name}.json")
    new_path = os.path.join(flow_dir, f"{new_name}.json")
    if not os.path.exists(old_path):
        abort(404)
    if os.path.exists(new_path):
        return "流程名稱已存在", 400
    os.rename(old_path, new_path)
    return redirect(url_for("flow_builder", task_id=task_id))

@app.get("/tasks/<task_id>/flows/export/<flow_name>")
def export_flow(task_id, flow_name):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    path = os.path.join(tdir, "flows", f"{flow_name}.json")
    if not os.path.exists(path):
        abort(404)
    return send_file(path, as_attachment=True, download_name=f"{flow_name}.json")


@app.post("/tasks/<task_id>/flows/import")
def import_flow(task_id):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    flow_dir = os.path.join(tdir, "flows")
    os.makedirs(flow_dir, exist_ok=True)
    f = request.files.get("flow_file")
    if not f or not f.filename.endswith(".json"):
        return "請上傳 JSON 檔", 400
    name = os.path.splitext(secure_filename(f.filename))[0]
    path = os.path.join(flow_dir, f"{name}.json")
    f.save(path)
    return redirect(url_for("flow_builder", task_id=task_id))


@app.get("/tasks/<task_id>/result/<job_id>")
def task_result(task_id, job_id):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    docx_path = os.path.join(job_dir, "result.docx")
    if not os.path.exists(docx_path):
        return "Job not found or failed.", 404
    log_json_path = os.path.join(job_dir, "log.json")
    log_entries = []
    overall_status = "ok"
    if os.path.exists(log_json_path):
        with open(log_json_path, "r", encoding="utf-8") as f:
            log_entries = json.load(f)
        if any(e.get("status") == "error" for e in log_entries):
            overall_status = "error"
    return render_template(
        "run.html",
        job_id=job_id,
        docx_path=url_for("task_download", task_id=task_id, job_id=job_id, kind="docx"),
        log_path=url_for("task_download", task_id=task_id, job_id=job_id, kind="log"),
        translate_path=url_for("task_translate", task_id=task_id, job_id=job_id),
        compare_path=url_for("task_compare", task_id=task_id, job_id=job_id),
        back_link=url_for("flow_builder", task_id=task_id),
        log_entries=log_entries,
        overall_status=overall_status,
    )


@app.get("/tasks/<task_id>/translate/<job_id>")
def task_translate(task_id, job_id):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    src = os.path.join(job_dir, "result.docx")
    if not os.path.exists(src):
        abort(404)
    out_docx = os.path.join(job_dir, "translated.docx")
    if not os.path.exists(out_docx):
        tmp_md = os.path.join(job_dir, "translated.md")
        translate_file(src, tmp_md)
        import docx
        doc = docx.Document()
        with open(tmp_md, "r", encoding="utf-8") as f:
            for line in f.read().splitlines():
                doc.add_paragraph(line)
        doc.save(out_docx)
    return send_file(
        out_docx,
        as_attachment=True,
        download_name=f"translated_{job_id}.docx",
    )


@app.get("/tasks/<task_id>/compare/<job_id>")
def task_compare(task_id, job_id):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    docx_path = os.path.join(job_dir, "result.docx")
    log_path = os.path.join(job_dir, "log.json")
    if not os.path.exists(docx_path) or not os.path.exists(log_path):
        abort(404)

    from spire.doc import Document, FileFormat

    with open(log_path, "r", encoding="utf-8") as f:
        entries = json.load(f)
    titles_to_hide = collect_titles_to_hide(entries)

    html_name = "result.html"
    html_path = os.path.join(job_dir, html_name)
    if not os.path.exists(html_path):
        doc = Document()
        doc.LoadFromFile(docx_path)
        doc.HtmlExportOptions.ImageEmbedded = True
        doc.SaveToFile(html_path, FileFormat.Html)
        doc.Close()
        remove_hidden_runs(docx_path, preserve_texts=titles_to_hide)

    chapter_sources = {}
    source_urls = {}
    converted_docx = {}
    current = None
    for entry in entries:
        stype = entry.get("type")
        params = entry.get("params", {})
        if stype == "insert_roman_heading":
            current = params.get("text", "")
            chapter_sources.setdefault(current, [])
        elif stype == "extract_pdf_chapter_to_table":
            pdf_dir = os.path.join(job_dir, "pdfs_extracted")
            pdfs = []
            if os.path.isdir(pdf_dir):
                for fn in sorted(os.listdir(pdf_dir)):
                    if fn.lower().endswith(".pdf"):
                        pdfs.append(fn)
                        rel = os.path.join("pdfs_extracted", fn)
                        source_urls[fn] = url_for(
                            "task_view_file", task_id=task_id, job_id=job_id, filename=rel
                        )
            chapter_sources.setdefault(current or "未分類", []).extend(pdfs)
        elif stype == "extract_word_chapter":
            infile = params.get("input_file", "")
            base = os.path.basename(infile)
            sec = params.get("target_chapter_section", "")
            use_title = str(params.get("target_title", "")).lower() in ["1", "true", "yes", "on"]
            title = params.get("target_title_section", "") if use_title else ""
            info = base
            if sec:
                info += f" 章節 {sec}"
            if title:
                info += f" 標題 {title}"
            chapter_sources.setdefault(current or "未分類", []).append(info)
            if base not in converted_docx and infile and os.path.exists(infile):
                preview_dir = os.path.join(job_dir, "source_html")
                os.makedirs(preview_dir, exist_ok=True)
                html_name_src = f"{os.path.splitext(base)[0]}.html"
                html_rel = os.path.join("source_html", html_name_src)
                html_path_src = os.path.join(job_dir, html_rel)
                doc = Document()
                doc.LoadFromFile(infile)
                doc.HtmlExportOptions.ImageEmbedded = True
                doc.SaveToFile(html_path_src, FileFormat.Html)
                doc.Close()
                converted_docx[base] = html_rel
            if base in converted_docx:
                source_urls[info] = url_for(
                    "task_view_file", task_id=task_id, job_id=job_id, filename=converted_docx[base]
                )
        elif stype == "extract_word_all_content":
            infile = params.get("input_file", "")
            base = os.path.basename(infile)
            chapter_sources.setdefault(current or "未分類", []).append(base)
            if base not in converted_docx and infile and os.path.exists(infile):
                preview_dir = os.path.join(job_dir, "source_html")
                os.makedirs(preview_dir, exist_ok=True)
                html_name_src = f"{os.path.splitext(base)[0]}.html"
                html_rel = os.path.join("source_html", html_name_src)
                html_path_src = os.path.join(job_dir, html_rel)
                doc = Document()
                doc.LoadFromFile(infile)
                doc.HtmlExportOptions.ImageEmbedded = True
                doc.SaveToFile(html_path_src, FileFormat.Html)
                doc.Close()
                converted_docx[base] = html_rel
            if base in converted_docx:
                source_urls[base] = url_for(
                    "task_view_file", task_id=task_id, job_id=job_id, filename=converted_docx[base]
                )

    chapters = list(chapter_sources.keys())
    html_url = url_for("task_view_file", task_id=task_id, job_id=job_id, filename=html_name)
    versions = build_version_context(task_id, job_id, job_dir)
    return render_template(
        "compare.html",
        html_url=html_url,
        chapters=chapters,
        chapter_sources=chapter_sources,
        source_urls=source_urls,
        titles_to_hide=titles_to_hide,
        back_link=url_for("task_result", task_id=task_id, job_id=job_id),
        save_url=url_for("task_compare_save", task_id=task_id, job_id=job_id),
        save_as_url=url_for("task_compare_save_as", task_id=task_id, job_id=job_id),
        download_url=url_for("task_download", task_id=task_id, job_id=job_id, kind="docx"),
        versions=versions,
    )


@app.post("/tasks/<task_id>/compare/<job_id>/save")
def task_compare_save(task_id, job_id):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    titles_to_hide = load_titles_to_hide_from_log(job_dir)
    html_content = request.form.get("html")
    if not html_content:
        data = request.get_json(silent=True) or {}
        html_content = data.get("html", "")
    if not html_content:
        return "缺少內容", 400
    html_content = clean_compare_html_content(html_content)
    save_compare_output(job_dir, html_content, titles_to_hide)
    return "OK"


@app.post("/tasks/<task_id>/compare/<job_id>/save-as")
def task_compare_save_as(task_id, job_id):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    titles_to_hide = load_titles_to_hide_from_log(job_dir)
    payload = request.get_json(silent=True) or {}
    html_content = payload.get("html")
    name = payload.get("name") or ""
    if not html_content:
        html_content = request.form.get("html")
        name = request.form.get("name") or name
    if not html_content:
        return jsonify({"error": "缺少內容"}), 400
    version_name = (name or "").strip()
    if not version_name:
        return jsonify({"error": "缺少版本名稱"}), 400
    html_content = clean_compare_html_content(html_content)
    versions_dir = os.path.join(job_dir, "versions")
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    unique_suffix = uuid.uuid4().hex[:6]
    version_id = f"{timestamp}_{unique_suffix}"
    slug = sanitize_version_slug(version_name)
    base_name = f"{version_id}_{slug}" if slug else version_id
    save_compare_output(
        job_dir,
        html_content,
        titles_to_hide,
        base_name=base_name,
        subdir="versions",
    )
    metadata = load_version_metadata(versions_dir)
    versions = metadata.get("versions", [])
    versions = [v for v in versions if v.get("id") != version_id]
    created_ts = datetime.now()
    versions.append(
        {
            "id": version_id,
            "name": version_name,
            "slug": slug,
            "base_name": base_name,
            "created_at": created_ts.isoformat(timespec="seconds"),
        }
    )
    versions.sort(key=lambda v: v.get("created_at", ""), reverse=True)
    metadata["versions"] = versions
    save_version_metadata(versions_dir, metadata)
    version_payload = {
        "id": version_id,
        "name": version_name,
        "created_at_display": created_ts.strftime("%Y-%m-%d %H:%M:%S"),
        "html_url": url_for(
            "task_view_file",
            task_id=task_id,
            job_id=job_id,
            filename=f"versions/{base_name}.html",
        ),
        "docx_url": url_for(
            "task_download_version",
            task_id=task_id,
            job_id=job_id,
            version_id=version_id,
        ),
        "restore_url": url_for(
            "task_compare_restore_version",
            task_id=task_id,
            job_id=job_id,
            version_id=version_id,
        ),
        "delete_url": url_for(
            "task_compare_delete_version",
            task_id=task_id,
            job_id=job_id,
            version_id=version_id,
        ),
    }
    return jsonify({"status": "ok", "version": version_payload})


@app.get("/tasks/<task_id>/view/<job_id>/<path:filename>")
def task_view_file(task_id, job_id, filename):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    safe_filename = filename.replace("\\", "/")
    file_path = os.path.join(job_dir, safe_filename)
    if not os.path.isfile(file_path):
        abort(404)
    return send_from_directory(job_dir, safe_filename)


@app.post("/tasks/<task_id>/compare/<job_id>/restore/<version_id>")
def task_compare_restore_version(task_id, job_id, version_id):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    versions_dir = os.path.join(job_dir, "versions")
    metadata = load_version_metadata(versions_dir)
    versions = metadata.get("versions", [])
    version = next((v for v in versions if v.get("id") == version_id), None)
    if not version:
        return jsonify({"error": "找不到指定版本"}), 404
    base_name = version.get("base_name")
    if not base_name:
        return jsonify({"error": "版本資料不完整"}), 404
    html_src = os.path.join(versions_dir, f"{base_name}.html")
    docx_src = os.path.join(versions_dir, f"{base_name}.docx")
    if not os.path.exists(html_src) or not os.path.exists(docx_src):
        return jsonify({"error": "版本檔案不存在"}), 404
    shutil.copyfile(html_src, os.path.join(job_dir, "result.html"))
    shutil.copyfile(docx_src, os.path.join(job_dir, "result.docx"))
    return jsonify({"status": "ok"})


@app.post("/tasks/<task_id>/compare/<job_id>/delete/<version_id>")
def task_compare_delete_version(task_id, job_id, version_id):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    versions_dir = os.path.join(job_dir, "versions")
    metadata = load_version_metadata(versions_dir)
    versions = metadata.get("versions", [])
    version = next((v for v in versions if v.get("id") == version_id), None)
    if not version:
        return jsonify({"error": "找不到指定版本"}), 404
    metadata["versions"] = [v for v in versions if v.get("id") != version_id]
    save_version_metadata(versions_dir, metadata)
    base_name = version.get("base_name")
    if base_name:
        for ext in ("html", "docx"):
            path = os.path.join(versions_dir, f"{base_name}.{ext}")
            try:
                if os.path.exists(path):
                    os.remove(path)
            except OSError:
                pass
    return jsonify({"status": "ok"})


@app.get("/tasks/<task_id>/download/<job_id>/version/<version_id>")
def task_download_version(task_id, job_id, version_id):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    versions_dir = os.path.join(job_dir, "versions")
    metadata = load_version_metadata(versions_dir)
    versions = metadata.get("versions", [])
    version = next((v for v in versions if v.get("id") == version_id), None)
    if not version:
        abort(404)
    base_name = version.get("base_name")
    if not base_name:
        abort(404)
    docx_src = os.path.join(versions_dir, f"{base_name}.docx")
    if not os.path.exists(docx_src):
        abort(404)
    slug = version.get("slug") or version_id
    download_name = f"{slug}_{version_id}.docx"
    return send_file(docx_src, as_attachment=True, download_name=download_name)


@app.get("/tasks/<task_id>/download/<job_id>/<kind>")
def task_download(task_id, job_id, kind):
    tdir = os.path.join(app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    if kind == "docx":
        result_path = os.path.join(job_dir, "result.docx")
        if not os.path.exists(result_path):
            abort(404)
        titles_to_remove = []
        log_path = os.path.join(job_dir, "log.json")
        if os.path.exists(log_path):
            try:
                with open(log_path, "r", encoding="utf-8") as f:
                    entries = json.load(f)
                titles_to_remove = collect_titles_to_hide(entries)
            except Exception:
                titles_to_remove = []

        download_path = os.path.join(job_dir, "result_download.docx")
        shutil.copyfile(result_path, download_path)
        if titles_to_remove:
            remove_paragraphs_with_text(download_path, titles_to_remove)
        remove_hidden_runs(download_path)
        return send_file(
            download_path,
            as_attachment=True,
            download_name=f"result_{job_id}.docx",
        )
    elif kind == "log":
        return send_file(
            os.path.join(job_dir, "log.json"),
            as_attachment=True,
            download_name=f"log_{job_id}.json",
        )
    abort(404)


if __name__ == "__main__":
    app.run(debug=True)
