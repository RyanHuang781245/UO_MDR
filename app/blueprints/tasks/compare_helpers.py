from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import tempfile
import uuid
from pathlib import Path

from flask import current_app, has_app_context
from werkzeug.utils import secure_filename

from modules.docx_provenance import (
    PROVENANCE_PREVIEW_LABEL_PREFIX,
    build_provenance_cache_payload,
    create_provenance_preview_docx,
    extract_provenance_block_trace,
    load_cached_provenance_payload,
)

_LIBREOFFICE_CANDIDATES = (
    "soffice",
    "libreoffice",
    "/usr/bin/soffice",
    "/usr/bin/libreoffice",
    "/usr/local/bin/soffice",
    "/usr/local/bin/libreoffice",
    "/snap/bin/soffice",
    "/snap/bin/libreoffice",
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
)
_PROVENANCE_PREVIEW_DOCX_CACHE_VERSION = 5
_PAGE_SOURCE_MAP_CACHE_VERSION = 14
_HTML_PREVIEW_CACHE_VERSION = 2
_LIBREOFFICE_REQUIRED_PATHS = (
    "/usr/local/sbin",
    "/usr/local/bin",
    "/usr/sbin",
    "/usr/bin",
    "/sbin",
    "/bin",
    "/snap/bin",
)


def _configured_libreoffice_binary() -> str | None:
    if has_app_context():
        configured = (current_app.config.get("LIBREOFFICE_BIN") or "").strip()
        if configured:
            return configured

    configured = (os.environ.get("LIBREOFFICE_BIN") or "").strip()
    return configured or None


def _find_libreoffice_binary() -> str | None:
    configured = _configured_libreoffice_binary()
    candidates = ((configured,) + _LIBREOFFICE_CANDIDATES) if configured else _LIBREOFFICE_CANDIDATES
    for candidate in candidates:
        resolved = shutil.which(candidate) if not os.path.isabs(candidate) else candidate
        if resolved and os.path.isfile(resolved):
            return resolved
    return None


def _build_libreoffice_env() -> dict[str, str]:
    env = os.environ.copy()
    if os.name == "nt":
        return env

    current_path = env.get("PATH", "")
    path_parts = [part for part in current_path.split(os.pathsep) if part]
    seen = set(path_parts)
    for required in _LIBREOFFICE_REQUIRED_PATHS:
        if required not in seen:
            path_parts.append(required)
            seen.add(required)
    env["PATH"] = os.pathsep.join(path_parts)
    return env


def _build_preview_pdf_name(source_path: str) -> str:
    stem = secure_filename(Path(source_path).stem) or "preview"
    digest = uuid.uuid5(uuid.NAMESPACE_URL, os.path.abspath(source_path)).hex[:10]
    return f"{stem}_{digest}.pdf"


def _ensure_pdf_preview(source_path: str, job_dir: str, subdir: str) -> tuple[str | None, str | None]:
    if not source_path or not os.path.isfile(source_path):
        return None, "找不到要預覽的文件"

    output_dir = os.path.join(job_dir, subdir)
    os.makedirs(output_dir, exist_ok=True)
    pdf_name = _build_preview_pdf_name(source_path)
    pdf_rel = os.path.join(subdir, pdf_name).replace("\\", "/")
    pdf_path = os.path.join(job_dir, pdf_rel)

    try:
        if os.path.exists(pdf_path) and os.path.getmtime(pdf_path) >= os.path.getmtime(source_path):
            return pdf_rel, None
    except OSError:
        pass

    libreoffice_bin = _find_libreoffice_binary()
    if not libreoffice_bin:
        return None, "找不到 LibreOffice，無法建立 PDF 預覽"

    try:
        with tempfile.TemporaryDirectory(prefix="lo_pdf_") as temp_dir:
            result = subprocess.run(
                [
                    libreoffice_bin,
                    "--headless",
                    "--convert-to",
                    "pdf:writer_pdf_Export",
                    "--outdir",
                    temp_dir,
                    source_path,
                ],
                capture_output=True,
                env=_build_libreoffice_env(),
                text=True,
                timeout=180,
                check=False,
            )
            converted_pdf = os.path.join(temp_dir, f"{Path(source_path).stem}.pdf")
            if result.returncode != 0 or not os.path.exists(converted_pdf):
                stdout = (result.stdout or "").strip()
                stderr = (result.stderr or "").strip()
                current_app.logger.warning(
                    "LibreOffice PDF conversion failed for %s (code=%s, stdout=%r, stderr=%r)",
                    source_path,
                    result.returncode,
                    stdout,
                    stderr,
                )
                return None, "LibreOffice 轉 PDF 失敗"
            shutil.copyfile(converted_pdf, pdf_path)
            return pdf_rel, None
    except subprocess.TimeoutExpired:
        current_app.logger.warning("LibreOffice PDF conversion timed out for %s", source_path)
        return None, "LibreOffice 轉 PDF 逾時"
    except Exception:
        current_app.logger.exception("Unexpected error while converting %s to PDF preview", source_path)
        return None, "建立 PDF 預覽時發生錯誤"


def _ensure_html_preview(source_path: str, job_dir: str, subdir: str, base_name: str) -> tuple[str | None, str | None]:
    if not source_path or not os.path.isfile(source_path):
        return None, "找不到要預覽的文件"

    output_dir = os.path.join(job_dir, subdir)
    os.makedirs(output_dir, exist_ok=True)
    html_name = f"{base_name}.html"
    html_rel = os.path.join(subdir, html_name).replace("\\", "/")
    html_path = os.path.join(job_dir, html_rel)
    meta_path = os.path.join(output_dir, "_meta.json")

    try:
        if os.path.exists(html_path) and os.path.exists(meta_path):
            with open(meta_path, "r", encoding="utf-8") as meta_file:
                meta = json.load(meta_file)
            if (
                meta.get("version") == _HTML_PREVIEW_CACHE_VERSION
                and os.path.getmtime(html_path) >= os.path.getmtime(source_path)
            ):
                return html_rel, None
    except OSError:
        pass
    except (ValueError, TypeError, json.JSONDecodeError):
        pass

    libreoffice_bin = _find_libreoffice_binary()
    if not libreoffice_bin:
        return None, "找不到 LibreOffice，無法建立 HTML 預覽"

    try:
        with tempfile.TemporaryDirectory(prefix="lo_html_") as temp_dir:
            result = subprocess.run(
                [
                    libreoffice_bin,
                    "--headless",
                    "--convert-to",
                    "html",
                    "--outdir",
                    temp_dir,
                    source_path,
                ],
                capture_output=True,
                env=_build_libreoffice_env(),
                text=True,
                timeout=180,
            )
            if result.returncode != 0:
                current_app.logger.warning(
                    "LibreOffice HTML conversion failed for %s (code=%s, stdout=%s, stderr=%s)",
                    source_path,
                    result.returncode,
                    result.stdout,
                    result.stderr,
                )
                return None, "LibreOffice 轉 HTML 失敗"

            converted_html = os.path.join(temp_dir, f"{Path(source_path).stem}.html")
            if not os.path.isfile(converted_html):
                current_app.logger.warning(
                    "LibreOffice HTML conversion did not produce %s for %s",
                    converted_html,
                    source_path,
                )
                return None, "找不到 LibreOffice 產生的 HTML 預覽"

            for entry in os.listdir(output_dir):
                entry_path = os.path.join(output_dir, entry)
                if os.path.isdir(entry_path):
                    shutil.rmtree(entry_path, ignore_errors=True)
                else:
                    try:
                        os.remove(entry_path)
                    except FileNotFoundError:
                        pass

            for entry in os.listdir(temp_dir):
                src_entry = os.path.join(temp_dir, entry)
                dst_entry = os.path.join(output_dir, entry)
                if os.path.isdir(src_entry):
                    shutil.copytree(src_entry, dst_entry, dirs_exist_ok=True)
                else:
                    shutil.copy2(src_entry, dst_entry)

            generated_html = os.path.join(output_dir, f"{Path(source_path).stem}.html")
            if generated_html != html_path and os.path.exists(generated_html):
                shutil.move(generated_html, html_path)

            _normalize_html_preview_alignment(html_path)

            with open(meta_path, "w", encoding="utf-8") as meta_file:
                json.dump({"version": _HTML_PREVIEW_CACHE_VERSION}, meta_file, ensure_ascii=False, indent=2)
        return html_rel, None
    except subprocess.TimeoutExpired:
        current_app.logger.warning("LibreOffice HTML conversion timed out for %s", source_path)
        return None, "LibreOffice 轉 HTML 逾時"
    except Exception:
        current_app.logger.exception("Unexpected error while converting %s to HTML preview", source_path)
        return None, "建立 HTML 預覽時發生錯誤"


def _normalize_html_preview_alignment(html_path: str) -> None:
    if not html_path or not os.path.isfile(html_path):
        return

    try:
        with open(html_path, "r", encoding="utf-8") as html_file:
            html = html_file.read()
    except OSError:
        return

    align_pattern = re.compile(
        r"<(?P<tag>[a-zA-Z0-9]+)(?P<before>[^>]*?)\salign=\"(?P<align>[^\"]+)\"(?P<after>[^>]*?)>",
        re.IGNORECASE,
    )

    def _replace_align(match: re.Match[str]) -> str:
        tag = match.group("tag")
        before = match.group("before") or ""
        align = (match.group("align") or "").strip().lower()
        after = match.group("after") or ""
        if not align:
            return match.group(0)

        attrs = f"{before}{after}"
        style_match = re.search(r'style=\"([^\"]*)\"', attrs, re.IGNORECASE)
        if style_match:
            style_value = style_match.group(1)
            if "text-align" in style_value.lower():
                return match.group(0)
            updated_style = f'{style_value.rstrip("; ")}; text-align: {align}'
            updated_attrs = attrs.replace(style_match.group(0), f'style="{updated_style}"', 1)
            return f"<{tag}{updated_attrs}>"

        return f'<{tag}{before} align="{align}" style="text-align: {align}"{after}>'

    normalized_html = align_pattern.sub(_replace_align, html)
    if normalized_html == html:
        return

    with open(html_path, "w", encoding="utf-8") as html_file:
        html_file.write(normalized_html)


def _ensure_provenance_preview_docx(
    result_docx: str,
    log_path: str,
    job_dir: str,
    source_lookup: dict[str, dict[str, object]],
) -> tuple[str | None, str | None]:
    if not result_docx or not os.path.isfile(result_docx):
        return None, "找不到結果文件"
    if not source_lookup:
        return None, None

    preview_rel = os.path.join("preview_trace", "provenance_preview.docx").replace("\\", "/")
    preview_path = os.path.join(job_dir, preview_rel)
    preview_meta_path = os.path.join(job_dir, "preview_trace", "provenance_preview.meta.json")

    try:
        preview_meta = {}
        if os.path.isfile(preview_meta_path):
            with open(preview_meta_path, "r", encoding="utf-8") as f:
                preview_meta = json.load(f) or {}
        if (
            os.path.isfile(preview_path)
            and os.path.getmtime(preview_path) >= os.path.getmtime(result_docx)
            and (not os.path.isfile(log_path) or os.path.getmtime(preview_path) >= os.path.getmtime(log_path))
            and int(preview_meta.get("version") or 0) == _PROVENANCE_PREVIEW_DOCX_CACHE_VERSION
        ):
            return preview_rel, None
    except OSError:
        pass
    except Exception:
        current_app.logger.warning("Failed to load provenance preview cache metadata for %s", job_dir, exc_info=True)

    try:
        if create_provenance_preview_docx(result_docx, preview_path, source_lookup):
            try:
                Path(preview_meta_path).write_text(
                    json.dumps({"version": _PROVENANCE_PREVIEW_DOCX_CACHE_VERSION}, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                )
            except Exception:
                current_app.logger.warning("Failed to save provenance preview cache metadata for %s", job_dir, exc_info=True)
            return preview_rel, None
        return None, "建立來源標記預覽失敗"
    except Exception:
        current_app.logger.exception("Unexpected error while building provenance preview for %s", result_docx)
        return None, "建立來源標記預覽時發生錯誤"


def _normalize_trace_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\xa0", " ")).strip()


def _format_source_file_label(path_raw: str) -> str:
    path = os.path.abspath(str(path_raw))
    file_path = Path(path)
    name = file_path.name or str(path_raw)
    parent_parts = [part for part in file_path.parts[-3:-1] if part not in {"", os.path.sep}]
    context = "/".join(parent_parts)
    if context:
        return f"{name} ({context})"
    return name


def _trace_source_label(entry: dict) -> str:
    params = entry.get("params") or {}
    input_file = params.get("input_file")
    if input_file:
        return _format_source_file_label(str(input_file))

    stype = str(entry.get("type") or "").strip()
    if stype == "insert_text":
        return "流程文字"
    if stype in {"insert_roman_heading", "insert_numbered_heading", "insert_bulleted_heading"}:
        return "流程標題"
    if stype == "template_merge":
        template_file = entry.get("template_file")
        if template_file:
            return _format_source_file_label(str(template_file))
    return stype or "未知來源"


def _extract_docx_trace_paragraphs(
    docx_path: str,
    *,
    hide_set: set[str] | None = None,
) -> list[dict[str, object]]:
    from docx import Document as DocxDocument

    items: list[dict[str, object]] = []
    hide_values = hide_set or set()
    doc = DocxDocument(docx_path)
    for idx, para in enumerate(doc.paragraphs):
        raw_text = (para.text or "").strip()
        normalized = _normalize_trace_text(raw_text)
        if not normalized or normalized in hide_values:
            continue
        items.append(
            {
                "paragraph_index": idx,
                "text": raw_text,
                "normalized_text": normalized,
            }
        )
    return items


def _extract_docx_table_texts(
    docx_path: str,
    *,
    hide_set: set[str] | None = None,
) -> list[str]:
    from docx import Document as DocxDocument

    hide_values = hide_set or set()
    doc = DocxDocument(docx_path)
    texts: list[str] = []
    seen: set[str] = set()

    for table in doc.tables:
        for row in table.rows:
            row_parts: list[str] = []
            for cell in row.cells:
                cell_text = _normalize_trace_text(cell.text or "")
                if not cell_text or cell_text in hide_values:
                    continue
                row_parts.append(cell_text)
                if len(cell_text) >= 20 and cell_text not in seen:
                    seen.add(cell_text)
                    texts.append(cell_text)
            row_text = _normalize_trace_text(" ".join(row_parts))
            if len(row_text) >= 24 and row_text not in seen:
                seen.add(row_text)
                texts.append(row_text)

    return texts


def _build_object_trace_candidates(
    entries: list[dict],
    titles_to_hide: list[str],
) -> list[dict[str, object]]:
    hide_set = {_normalize_trace_text(title) for title in titles_to_hide if _normalize_trace_text(title)}
    candidates: list[dict[str, object]] = []

    for entry in entries:
        if not isinstance(entry, dict):
            continue
        if entry.get("status") == "error":
            continue

        stype = str(entry.get("type") or "")
        if stype not in {"extract_specific_figure_from_word", "extract_specific_table_from_word"}:
            continue

        params = entry.get("params") or {}
        source_file = _trace_source_label(entry)
        output_docx = entry.get("output_docx")
        primary_probe_texts: list[str] = []
        fallback_probe_texts: list[str] = []

        if output_docx and os.path.isfile(str(output_docx)):
            try:
                paragraphs = _extract_docx_trace_paragraphs(str(output_docx), hide_set=hide_set)
                primary_probe_texts.extend(
                    [str(item["normalized_text"]) for item in paragraphs if item.get("normalized_text")]
                )
                if stype == "extract_specific_table_from_word":
                    fallback_probe_texts.extend(_extract_docx_table_texts(str(output_docx), hide_set=hide_set))
            except Exception:
                current_app.logger.warning("Failed to extract object trace paragraphs from %s", output_docx, exc_info=True)

        for part in (
            params.get("target_caption_label"),
            params.get("target_figure_title"),
            params.get("target_table_title"),
            params.get("target_chapter_title"),
            params.get("target_subtitle"),
        ):
            synthetic_text = _normalize_trace_text(str(part).strip())
            if synthetic_text:
                primary_probe_texts.append(synthetic_text)

        def _dedupe_probes(raw_probes: list[str]) -> list[str]:
            deduped: list[str] = []
            seen: set[str] = set()
            for probe in raw_probes:
                cleaned = _normalize_trace_text(probe)
                if not cleaned or cleaned in seen:
                    continue
                seen.add(cleaned)
                deduped.append(cleaned)
            return deduped

        candidates.append(
            {
                "source_file": source_file,
                "source_step": stype,
                "primary_probe_texts": _dedupe_probes(primary_probe_texts),
                "fallback_probe_texts": _dedupe_probes(fallback_probe_texts),
                "allow_multi_page": stype == "extract_specific_table_from_word",
            }
        )

    return candidates


def _build_compare_source_label(entry: dict) -> str:
    params = entry.get("params") or {}
    base = _trace_source_label(entry)
    stype = str(entry.get("type") or "")

    if stype == "extract_specific_figure_from_word":
        label = params.get("target_caption_label") or params.get("target_figure_title") or "圖片"
        return f"{base} 圖片 {label}".strip()
    if stype == "extract_specific_table_from_word":
        label = params.get("target_caption_label") or params.get("target_table_title") or "表格"
        return f"{base} 表格 {label}".strip()
    return base


def _collect_provenance_probe_texts(entry: dict) -> list[str]:
    params = entry.get("params") or {}
    texts: list[str] = []
    for raw_value in (
        params.get("target_caption_label"),
        params.get("target_figure_title"),
        params.get("target_table_title"),
        params.get("target_chapter_title"),
        params.get("target_title_section"),
        params.get("target_subtitle"),
        params.get("subheading_text"),
    ):
        normalized = _normalize_trace_text(str(raw_value or ""))
        if normalized:
            texts.append(normalized)
    deduped: list[str] = []
    seen: set[str] = set()
    for text in texts:
        if text in seen:
            continue
        seen.add(text)
        deduped.append(text)
    return deduped


def _build_provenance_source_lookup(entries: list[dict]) -> dict[str, dict[str, object]]:
    lookup: dict[str, dict[str, object]] = {}
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        provenance = entry.get("provenance")
        if not isinstance(provenance, dict):
            continue
        source_id = str(provenance.get("source_id") or "").strip()
        if not source_id:
            continue
        lookup[source_id] = {
            "source_id": source_id,
            "source_file": _trace_source_label(entry),
            "source_step": str(entry.get("type") or ""),
            "content_type": str(provenance.get("content_type") or ""),
            "bookmark_start": str(provenance.get("bookmark_start") or ""),
            "bookmark_end": str(provenance.get("bookmark_end") or ""),
            "bookmark_id": provenance.get("bookmark_id"),
            "fragment_path": str(provenance.get("fragment_path") or entry.get("output_docx") or ""),
            "fragment_order": provenance.get("fragment_order"),
            "template_index": provenance.get("template_index", entry.get("template_index")),
            "template_mode": provenance.get("template_mode", entry.get("template_mode")),
            "primary_probe_texts": _collect_provenance_probe_texts(entry),
        }
    return lookup


def _build_trace_from_provenance_blocks(
    block_trace: list[dict[str, object]],
) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    paragraph_trace: list[dict[str, object]] = []
    object_candidates: list[dict[str, object]] = []

    for block in block_trace:
        block_type = str(block.get("block_type") or "")
        source_file = str(block.get("source_file") or "未知來源")
        source_step = str(block.get("source_step") or "")
        content_type = str(block.get("content_type") or "")
        probe_texts = [str(item) for item in (block.get("probe_texts") or []) if str(item).strip()]
        if block_type == "paragraph":
            text = str(block.get("text") or "").strip()
            if not text:
                continue
            source_id = str(block.get("source_id") or "").strip()
            paragraph_trace.append(
                {
                    "merged_paragraph_index": len(paragraph_trace),
                    "source_id": source_id,
                    "source_file": source_file,
                    "source_paragraph_index": block.get("block_index"),
                    "result_block_index": block.get("block_index"),
                    "source_step": source_step,
                    "content_type": content_type or "paragraph",
                    "match_status": "provenance" if source_id else "context",
                    "count_as_source": bool(source_id),
                    "text": text,
                    "probe_texts": probe_texts[:],
                }
            )
            continue

        if not probe_texts:
            continue
        object_candidates.append(
            {
                "source_file": source_file,
                "source_step": source_step,
                "source_id": str(block.get("source_id") or "").strip(),
                "content_type": content_type,
                "count_as_source": bool(block.get("source_id")),
                "result_block_index": block.get("block_index"),
                "primary_probe_texts": probe_texts[:6],
                "fallback_probe_texts": probe_texts[6:],
                "allow_multi_page": content_type == "table",
            }
        )

    return paragraph_trace, object_candidates


def _build_provenance_trace(
    job_dir: str,
    result_docx: str,
    log_path: str,
    entries: list[dict],
    titles_to_hide: list[str],
) -> tuple[list[dict[str, object]], list[dict[str, object]]] | None:
    trace_dir = os.path.join(job_dir, "preview_trace")
    os.makedirs(trace_dir, exist_ok=True)
    cache_path = os.path.join(trace_dir, "provenance_map.json")
    source_lookup = _build_provenance_source_lookup(entries)
    if not source_lookup:
        return None

    cached_payload = None
    try:
        if (
            os.path.isfile(cache_path)
            and os.path.getmtime(cache_path) >= os.path.getmtime(result_docx)
            and os.path.getmtime(cache_path) >= os.path.getmtime(log_path)
        ):
            cached_payload = load_cached_provenance_payload(cache_path)
    except Exception:
        current_app.logger.warning("Failed to load cached provenance trace for %s", job_dir, exc_info=True)

    if cached_payload:
        block_trace = cached_payload.get("block_trace") or []
        if isinstance(block_trace, list):
            return _build_trace_from_provenance_blocks(block_trace)

    hide_set = {_normalize_trace_text(title) for title in titles_to_hide if _normalize_trace_text(title)}
    block_trace = extract_provenance_block_trace(result_docx, source_lookup, hide_set=hide_set)
    if not block_trace:
        return None

    try:
        payload = build_provenance_cache_payload(source_lookup=source_lookup, block_trace=block_trace)
        Path(cache_path).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        current_app.logger.warning("Failed to save provenance trace cache for %s", job_dir, exc_info=True)

    return _build_trace_from_provenance_blocks(block_trace)


def _build_paragraph_trace(
    job_dir: str,
    result_docx: str,
    log_path: str,
    entries: list[dict],
    titles_to_hide: list[str],
) -> list[dict[str, object]]:
    trace_dir = os.path.join(job_dir, "preview_trace")
    os.makedirs(trace_dir, exist_ok=True)
    trace_path = os.path.join(trace_dir, "paragraph_trace.json")
    hide_set = {_normalize_trace_text(title) for title in titles_to_hide if _normalize_trace_text(title)}

    try:
        if (
            os.path.isfile(trace_path)
            and os.path.getmtime(trace_path) >= os.path.getmtime(result_docx)
            and os.path.getmtime(trace_path) >= os.path.getmtime(log_path)
        ):
            with open(trace_path, "r", encoding="utf-8") as f:
                cached = json.load(f)
            if isinstance(cached, list):
                return cached
    except Exception:
        current_app.logger.warning("Failed to load cached paragraph trace for %s", job_dir)

    final_paragraphs = _extract_docx_trace_paragraphs(result_docx, hide_set=hide_set)
    candidates: list[dict[str, object]] = []

    for entry in entries:
        if not isinstance(entry, dict):
            continue
        if entry.get("status") == "error":
            continue
        if entry.get("type") in {
            "file_metadata",
            "template_merge",
            "copy_files",
            "copy_directory",
            "extract_specific_figure_from_word",
            "extract_specific_table_from_word",
        }:
            continue
        output_docx = entry.get("output_docx")
        if not output_docx or not os.path.isfile(str(output_docx)):
            continue

        source_file = _trace_source_label(entry)
        source_step = str(entry.get("type") or "")
        try:
            paragraphs = _extract_docx_trace_paragraphs(str(output_docx), hide_set=hide_set)
        except Exception:
            current_app.logger.warning("Failed to extract trace paragraphs from %s", output_docx, exc_info=True)
            continue

        for para in paragraphs:
            candidates.append(
                {
                    "source_file": source_file,
                    "source_step": source_step,
                    "source_paragraph_index": para["paragraph_index"],
                    "text": para["text"],
                    "normalized_text": para["normalized_text"],
                }
            )

    trace: list[dict[str, object]] = []
    pointer = 0
    for merged_index, para in enumerate(final_paragraphs):
        normalized = str(para["normalized_text"])
        matched: dict[str, object] | None = None
        matched_index: int | None = None

        for idx in range(pointer, len(candidates)):
            if candidates[idx]["normalized_text"] == normalized:
                matched = candidates[idx]
                matched_index = idx
                break

        if matched_index is not None:
            pointer = matched_index + 1

        trace.append(
            {
                "merged_paragraph_index": merged_index,
                "source_file": matched["source_file"] if matched else "未知來源",
                "source_paragraph_index": matched["source_paragraph_index"] if matched else None,
                "source_step": matched["source_step"] if matched else "",
                "match_status": "matched" if matched else "unmatched",
                "text": para["text"],
            }
        )

    try:
        with open(trace_path, "w", encoding="utf-8") as f:
            json.dump(trace, f, ensure_ascii=False, indent=2)
    except Exception:
        current_app.logger.warning("Failed to save paragraph trace cache for %s", job_dir, exc_info=True)

    return trace


def _build_trace_text_probes(normalized_text: str) -> list[str]:
    text = (normalized_text or "").strip()
    if not text:
        return []
    probes: list[str] = []
    if len(text) <= 180:
        probes.append(text)
    for size in (180, 120, 80, 48):
        if len(text) > size:
            probes.append(text[:size].strip())
    deduped: list[str] = []
    seen: set[str] = set()
    for probe in probes:
        cleaned = probe.strip()
        if len(cleaned) < 12 or cleaned in seen:
            continue
        seen.add(cleaned)
        deduped.append(cleaned)
    return deduped


def _select_page_sources_for_display(
    ordered_sources: list[tuple[str, int]],
    *,
    inherited_from_previous: bool = False,
    preserve_sources: set[str] | None = None,
) -> list[tuple[str, int]]:
    del inherited_from_previous
    del preserve_sources
    return [
        (source_file, count)
        for source_file, count in ordered_sources
        if str(source_file).strip() and int(count) > 0
    ]


def _order_page_sources_by_first_seen(
    sources: list[tuple[str, int]],
    first_seen_order: dict[str, int] | None = None,
) -> list[tuple[str, int]]:
    if not sources:
        return []

    first_seen = first_seen_order or {}
    default_order = len(first_seen) + 1
    return sorted(
        sources,
        key=lambda item: (
            first_seen.get(item[0], default_order),
            item[0],
        ),
    )


def _page_has_explicit_paragraph_sources(source_counts: dict[str, int]) -> bool:
    return any(count > 0 for count in source_counts.values())


def _score_probe_matches(page_text: str, probe_texts: list[str]) -> int:
    normalized_page_text = _normalize_trace_text(page_text)
    score = 0
    seen: set[str] = set()
    for raw_probe in probe_texts:
        for probe in _build_trace_text_probes(_normalize_trace_text(str(raw_probe))):
            if probe in seen:
                continue
            if probe in normalized_page_text:
                seen.add(probe)
                score += min(len(probe), 120)
    return score


def _extract_preview_page_sources(
    page_texts: list[str],
    source_lookup: dict[str, dict[str, object]] | None = None,
) -> list[list[str]]:
    unique_sources: list[str] = []
    seen_sources: set[str] = set()
    for meta in (source_lookup or {}).values():
        source_file = str(meta.get("source_file") or "").strip()
        if not source_file or source_file in seen_sources:
            continue
        seen_sources.add(source_file)
        unique_sources.append(source_file)

    label_patterns = [
        (source_file, _normalize_trace_text(f"{PROVENANCE_PREVIEW_LABEL_PREFIX}{source_file}"))
        for source_file in unique_sources
    ]

    page_sources: list[list[str]] = []
    for page_text in page_texts:
        normalized_page_text = _normalize_trace_text(page_text)
        ordered_hits: list[tuple[int, str]] = []
        for source_file, pattern in label_patterns:
            if not pattern:
                continue
            start = 0
            while True:
                position = normalized_page_text.find(pattern, start)
                if position < 0:
                    break
                ordered_hits.append((position, source_file))
                start = position + max(1, len(pattern))
        ordered_hits.sort(key=lambda item: (item[0], item[1]))
        page_sources.append([source_file for _, source_file in ordered_hits])

    return page_sources


def _merge_page_source_map_with_preview_labels(
    page_source_map: list[dict[str, object]],
    preview_page_sources: list[list[str]],
) -> list[dict[str, object]]:
    if not page_source_map:
        return []

    merged_map: list[dict[str, object]] = []
    for idx, bucket in enumerate(page_source_map):
        updated_bucket = dict(bucket)
        raw_label_sources = preview_page_sources[idx] if idx < len(preview_page_sources) else []
        label_sources: list[str] = []
        for source_file in raw_label_sources:
            source_key = str(source_file).strip()
            if not source_key:
                continue
            label_sources.append(source_key)

        updated_bucket["preview_label_sequence"] = label_sources[:]

        existing_sources = [
            str(item.get("source_file") or "").strip()
            for item in (updated_bucket.get("sources") or [])
            if str(item.get("source_file") or "").strip()
        ]

        if label_sources and not existing_sources:
            updated_bucket["sources"] = [
                {
                    "source_file": source_key,
                    "count": 1,
                    "inherited": False,
                    "from_preview_label": True,
                    "preview_segment_role": "label",
                }
                for source_key in label_sources
            ]
            updated_bucket["dominant_source"] = label_sources[0]
            updated_bucket["uses_preview_labels"] = True
            updated_bucket["preview_has_source_switch"] = len(label_sources) > 1
        merged_map.append(updated_bucket)

    return merged_map


def _has_provenance_result_blocks(
    paragraph_trace: list[dict[str, object]],
    object_trace_candidates: list[dict[str, object]] | None = None,
) -> bool:
    if any(item.get("result_block_index") is not None for item in paragraph_trace):
        return True
    return any(
        candidate.get("result_block_index") is not None
        for candidate in (object_trace_candidates or [])
    )


def _build_provenance_page_mapping_blocks(
    paragraph_trace: list[dict[str, object]],
    object_trace_candidates: list[dict[str, object]] | None = None,
) -> list[dict[str, object]]:
    blocks: list[dict[str, object]] = []

    for item in paragraph_trace:
        block_index = item.get("result_block_index")
        if block_index is None:
            continue
        text = _normalize_trace_text(str(item.get("text") or ""))
        probe_texts = [
            _normalize_trace_text(str(probe))
            for probe in (item.get("probe_texts") or [])
            if _normalize_trace_text(str(probe))
        ]
        if text and text not in probe_texts:
            probe_texts.insert(0, text)
        blocks.append(
            {
                "block_index": int(block_index),
                "block_type": str(item.get("content_type") or "paragraph"),
                "source_id": str(item.get("source_id") or "").strip(),
                "source_file": str(item.get("source_file") or "未知來源"),
                "source_step": str(item.get("source_step") or ""),
                "count_as_source": bool(item.get("count_as_source")),
                "text": str(item.get("text") or ""),
                "primary_probe_texts": probe_texts,
                "fallback_probe_texts": [],
                "allow_multi_page": False,
            }
        )

    for candidate in object_trace_candidates or []:
        block_index = candidate.get("result_block_index")
        if block_index is None:
            continue
        primary_probe_texts = [
            _normalize_trace_text(str(probe))
            for probe in (candidate.get("primary_probe_texts") or [])
            if _normalize_trace_text(str(probe))
        ]
        fallback_probe_texts = [
            _normalize_trace_text(str(probe))
            for probe in (candidate.get("fallback_probe_texts") or [])
            if _normalize_trace_text(str(probe))
        ]
        blocks.append(
            {
                "block_index": int(block_index),
                "block_type": str(candidate.get("content_type") or "object"),
                "source_id": str(candidate.get("source_id") or "").strip(),
                "source_file": str(candidate.get("source_file") or "未知來源"),
                "source_step": str(candidate.get("source_step") or ""),
                "count_as_source": bool(candidate.get("count_as_source")),
                "text": "",
                "primary_probe_texts": primary_probe_texts,
                "fallback_probe_texts": fallback_probe_texts,
                "allow_multi_page": bool(candidate.get("allow_multi_page")),
            }
        )

    return sorted(blocks, key=lambda item: int(item.get("block_index") or -1))


def _select_result_block_pages(
    page_texts: list[str],
    *,
    start_page_idx: int,
    primary_probe_texts: list[str],
    fallback_probe_texts: list[str],
    allow_multi_page: bool,
) -> list[int]:
    if not page_texts:
        return []

    bounded_start = max(0, min(start_page_idx, len(page_texts) - 1))

    def _score_pages(probes: list[str]) -> dict[int, int]:
        return {
            page_idx: _score_probe_matches(page_texts[page_idx], probes)
            for page_idx in range(bounded_start, len(page_texts))
        }

    def _pick_pages(scores: dict[int, int], *, permit_multi_page: bool) -> list[int]:
        positive_scores = {page_idx: score for page_idx, score in scores.items() if score > 0}
        if not positive_scores:
            return []
        best_page = max(positive_scores, key=lambda page_idx: (positive_scores[page_idx], -page_idx))
        max_score = positive_scores[best_page]
        if not permit_multi_page:
            return [best_page]

        min_score = max(1, (max_score + 2) // 3)
        selected = {page_idx for page_idx, score in positive_scores.items() if score >= min_score}
        cluster = [best_page]

        next_page = best_page - 1
        while next_page in selected:
            cluster.insert(0, next_page)
            next_page -= 1

        next_page = best_page + 1
        while next_page in selected:
            cluster.append(next_page)
            next_page += 1

        return cluster

    selected_pages = _pick_pages(_score_pages(primary_probe_texts), permit_multi_page=allow_multi_page)
    if selected_pages:
        return selected_pages

    selected_pages = _pick_pages(_score_pages(fallback_probe_texts), permit_multi_page=allow_multi_page)
    if selected_pages:
        return selected_pages

    return [bounded_start]


def _build_page_source_map_from_provenance_blocks(
    page_texts: list[str],
    paragraph_trace: list[dict[str, object]],
    object_trace_candidates: list[dict[str, object]] | None = None,
) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    page_buckets: list[dict[str, object]] = [
        {
            "page_number": idx + 1,
            "dominant_source": "",
            "sources": [],
        }
        for idx in range(len(page_texts))
    ]
    source_counts_by_page: list[dict[str, int]] = [dict() for _ in page_texts]
    source_first_seen_by_page: list[dict[str, int]] = [dict() for _ in page_texts]
    has_context_only_content_by_page = [False for _ in page_texts]
    preserved_sources_by_page: list[set[str]] = [set() for _ in page_texts]
    block_page_map: dict[int, list[int]] = {}
    current_page_idx = 0
    source_seen_sequence = 0

    for block in _build_provenance_page_mapping_blocks(paragraph_trace, object_trace_candidates):
        selected_pages = _select_result_block_pages(
            page_texts,
            start_page_idx=current_page_idx,
            primary_probe_texts=[str(item) for item in (block.get("primary_probe_texts") or []) if str(item).strip()],
            fallback_probe_texts=[str(item) for item in (block.get("fallback_probe_texts") or []) if str(item).strip()],
            allow_multi_page=bool(block.get("allow_multi_page")),
        )
        if not selected_pages:
            continue

        block_page_map[int(block["block_index"])] = selected_pages
        current_page_idx = selected_pages[-1]

        source_file = str(block.get("source_file") or "未知來源")
        source_step = str(block.get("source_step") or "")
        count_as_source = bool(
            block.get("count_as_source", str(block.get("source_id") or "").strip())
        )

        for page_idx in selected_pages:
            if count_as_source and source_file not in {"", "未知來源"}:
                source_counts = source_counts_by_page[page_idx]
                source_counts[source_file] = source_counts.get(source_file, 0) + 1
                source_first_seen = source_first_seen_by_page[page_idx]
                if source_file not in source_first_seen:
                    source_first_seen[source_file] = source_seen_sequence
                    source_seen_sequence += 1
                if source_step in {"extract_specific_table_from_word", "extract_specific_figure_from_word"}:
                    preserved_sources_by_page[page_idx].add(source_file)
            else:
                has_context_only_content_by_page[page_idx] = True

    annotated_trace: list[dict[str, object]] = []
    for item in paragraph_trace:
        trace_item = dict(item)
        block_index = trace_item.get("result_block_index")
        assigned_pages = block_page_map.get(int(block_index)) if block_index is not None else None
        trace_item["result_page"] = assigned_pages[0] + 1 if assigned_pages else None
        annotated_trace.append(trace_item)

    for idx, bucket in enumerate(page_buckets):
        source_counts = source_counts_by_page[idx]
        ordered_sources = sorted(source_counts.items(), key=lambda item: (-item[1], item[0]))
        display_sources = _select_page_sources_for_display(
            ordered_sources,
            preserve_sources=preserved_sources_by_page[idx],
        )
        display_sources = _order_page_sources_by_first_seen(
            display_sources,
            source_first_seen_by_page[idx],
        )
        dominant_source = ordered_sources[0][0] if ordered_sources else ""
        bucket["sources"] = [
            {
                "source_file": source_file,
                "count": count,
                "inherited": False,
            }
            for source_file, count in display_sources
        ]
        bucket["dominant_source"] = dominant_source

    return annotated_trace, page_buckets


def _select_object_candidate_pages(
    page_texts: list[str],
    source_counts_by_page: list[dict[str, int]],
    *,
    primary_probe_texts: list[str],
    fallback_probe_texts: list[str],
    allow_multi_page: bool,
) -> list[int]:
    available_pages = [
        page_idx
        for page_idx in range(len(page_texts))
        if not _page_has_explicit_paragraph_sources(source_counts_by_page[page_idx])
    ]
    if not available_pages:
        return []

    def _score_pages(probes: list[str]) -> dict[int, int]:
        return {
            page_idx: _score_probe_matches(page_texts[page_idx], probes)
            for page_idx in available_pages
        }

    def _pick_pages(scores: dict[int, int], *, permit_multi_page: bool) -> list[int]:
        positive_scores = {page_idx: score for page_idx, score in scores.items() if score > 0}
        if not positive_scores:
            return []
        best_page = max(positive_scores, key=lambda page_idx: (positive_scores[page_idx], -page_idx))
        max_score = positive_scores[best_page]
        if not permit_multi_page:
            return [best_page]

        min_score = max(1, (max_score + 2) // 3)
        selected = {page_idx for page_idx, score in positive_scores.items() if score >= min_score}
        cluster = [best_page]

        next_page = best_page - 1
        while next_page in selected:
            cluster.insert(0, next_page)
            next_page -= 1

        next_page = best_page + 1
        while next_page in selected:
            cluster.append(next_page)
            next_page += 1

        return cluster

    primary_scores = _score_pages(primary_probe_texts)
    selected_pages = _pick_pages(primary_scores, permit_multi_page=allow_multi_page)
    if selected_pages:
        return selected_pages

    fallback_scores = _score_pages(fallback_probe_texts)
    return _pick_pages(fallback_scores, permit_multi_page=allow_multi_page)


def _build_page_source_map(
    job_dir: str,
    result_pdf_path: str,
    paragraph_trace: list[dict[str, object]],
    object_trace_candidates: list[dict[str, object]] | None = None,
    source_lookup: dict[str, dict[str, object]] | None = None,
) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    if not result_pdf_path or not os.path.isfile(result_pdf_path):
        return paragraph_trace, []

    trace_dir = os.path.join(job_dir, "preview_trace")
    os.makedirs(trace_dir, exist_ok=True)
    page_map_path = os.path.join(trace_dir, "page_source_map.json")

    try:
        if os.path.isfile(page_map_path) and os.path.getmtime(page_map_path) >= os.path.getmtime(result_pdf_path):
            with open(page_map_path, "r", encoding="utf-8") as f:
                payload = json.load(f)
            if payload.get("version") != _PAGE_SOURCE_MAP_CACHE_VERSION:
                raise ValueError("stale page source map cache")
            cached_trace = payload.get("paragraph_trace")
            cached_pages = payload.get("page_source_map")
            if isinstance(cached_trace, list) and isinstance(cached_pages, list):
                return cached_trace, cached_pages
    except Exception:
        current_app.logger.warning("Failed to load cached page source map for %s", job_dir, exc_info=True)

    import fitz

    with fitz.open(result_pdf_path) as pdf:
        page_texts = [
            _normalize_trace_text(page.get_text("text"))
            for page in pdf
        ]
    preview_page_sources = _extract_preview_page_sources(page_texts, source_lookup)

    if _has_provenance_result_blocks(paragraph_trace, object_trace_candidates):
        annotated_trace, page_buckets = _build_page_source_map_from_provenance_blocks(
            page_texts,
            paragraph_trace,
            object_trace_candidates,
        )
        page_buckets = _merge_page_source_map_with_preview_labels(page_buckets, preview_page_sources)
        try:
            with open(page_map_path, "w", encoding="utf-8") as f:
                json.dump(
                    {
                        "version": _PAGE_SOURCE_MAP_CACHE_VERSION,
                        "paragraph_trace": annotated_trace,
                        "page_source_map": page_buckets,
                    },
                    f,
                    ensure_ascii=False,
                    indent=2,
                )
        except Exception:
            current_app.logger.warning("Failed to save page source map cache for %s", job_dir, exc_info=True)
        return annotated_trace, page_buckets

    page_buckets: list[dict[str, object]] = [
        {
            "page_number": idx + 1,
            "dominant_source": "",
            "sources": [],
        }
        for idx in range(len(page_texts))
    ]
    source_counts_by_page: list[dict[str, int]] = [dict() for _ in page_texts]
    source_first_seen_by_page: list[dict[str, int]] = [dict() for _ in page_texts]
    has_context_only_content_by_page = [False for _ in page_texts]
    preserved_sources_by_page: list[set[str]] = [set() for _ in page_texts]
    annotated_trace: list[dict[str, object]] = []
    current_page_idx = 0
    source_seen_sequence = 0

    for item in paragraph_trace:
        trace_item = dict(item)
        normalized = _normalize_trace_text(str(trace_item.get("text") or ""))
        assigned_page_idx = current_page_idx if page_texts else -1
        if normalized and page_texts:
            probes = _build_trace_text_probes(normalized)
            found_idx = None
            for page_idx in range(current_page_idx, len(page_texts)):
                page_text = page_texts[page_idx]
                if any(probe in page_text for probe in probes):
                    found_idx = page_idx
                    break
            if found_idx is not None:
                assigned_page_idx = found_idx
                current_page_idx = found_idx

        trace_item["result_page"] = assigned_page_idx + 1 if assigned_page_idx >= 0 else None
        annotated_trace.append(trace_item)

        if assigned_page_idx >= 0:
            count_as_source = bool(
                trace_item.get("count_as_source", str(trace_item.get("source_file") or "").strip() not in {"", "未知來源"})
            )
            source_file = str(trace_item.get("source_file") or "未知來源")
            source_step = str(trace_item.get("source_step") or "")
            if count_as_source and source_file not in {"", "未知來源"}:
                source_counts = source_counts_by_page[assigned_page_idx]
                source_counts[source_file] = source_counts.get(source_file, 0) + 1
                source_first_seen = source_first_seen_by_page[assigned_page_idx]
                if source_file not in source_first_seen:
                    source_first_seen[source_file] = source_seen_sequence
                    source_seen_sequence += 1
                if source_step in {"extract_specific_table_from_word", "extract_specific_figure_from_word"}:
                    preserved_sources_by_page[assigned_page_idx].add(source_file)
            else:
                has_context_only_content_by_page[assigned_page_idx] = True

    object_candidates = object_trace_candidates or []
    for candidate in object_candidates:
        if not page_texts:
            continue

        matched_pages = _select_object_candidate_pages(
            page_texts,
            source_counts_by_page,
            primary_probe_texts=[str(item) for item in (candidate.get("primary_probe_texts") or [])],
            fallback_probe_texts=[str(item) for item in (candidate.get("fallback_probe_texts") or [])],
            allow_multi_page=bool(candidate.get("allow_multi_page")),
        )
        if not matched_pages:
            continue

        source_file = str(candidate.get("source_file") or "未知來源")
        for page_idx in matched_pages:
            source_counts = source_counts_by_page[page_idx]
            source_counts[source_file] = source_counts.get(source_file, 0) + 1
            source_first_seen = source_first_seen_by_page[page_idx]
            if source_file not in source_first_seen:
                source_first_seen[source_file] = source_seen_sequence
                source_seen_sequence += 1

    for idx, bucket in enumerate(page_buckets):
        source_counts = source_counts_by_page[idx]
        ordered_sources = sorted(source_counts.items(), key=lambda item: (-item[1], item[0]))
        display_sources = _select_page_sources_for_display(
            ordered_sources,
            preserve_sources=preserved_sources_by_page[idx],
        )
        display_sources = _order_page_sources_by_first_seen(
            display_sources,
            source_first_seen_by_page[idx],
        )
        dominant_source = ordered_sources[0][0] if ordered_sources else ""
        bucket["sources"] = [
            {
                "source_file": source_file,
                "count": count,
                "inherited": False,
            }
            for source_file, count in display_sources
        ]
        bucket["dominant_source"] = dominant_source

    page_buckets = _merge_page_source_map_with_preview_labels(page_buckets, preview_page_sources)

    try:
        with open(page_map_path, "w", encoding="utf-8") as f:
            json.dump(
                {
                    "version": _PAGE_SOURCE_MAP_CACHE_VERSION,
                    "paragraph_trace": annotated_trace,
                    "page_source_map": page_buckets,
                },
                f,
                ensure_ascii=False,
                indent=2,
            )
    except Exception:
        current_app.logger.warning("Failed to save page source map cache for %s", job_dir, exc_info=True)

    return annotated_trace, page_buckets
