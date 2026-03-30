from __future__ import annotations

import json
import os
import shutil
import subprocess
import tempfile
import uuid
from datetime import datetime
import re
from pathlib import Path

from flask import Blueprint, abort, current_app, flash, jsonify, redirect, render_template, request, send_file, send_from_directory, url_for
from flask_login import current_user
from werkzeug.utils import secure_filename

from app.services.flow_service import (
    SKIP_DOCX_CLEANUP,
    clean_compare_html_content,
    collect_titles_to_hide,
    load_titles_to_hide_from_log,
    load_version_metadata,
    parse_template_paragraphs,
    remove_hidden_runs,
    remove_paragraphs_with_text,
    save_compare_output,
    save_version_metadata,
    sanitize_version_slug,
    translate_file,
)
from app.services.audit_service import record_audit

from app.services.task_service import (
    build_file_tree,
    delete_task_record,
    enforce_max_copy_size,
    ensure_windows_long_path,
    gather_available_files,
    list_dirs,
    list_files,
    list_tasks,
    record_task_in_db,
    task_name_exists,
)

from app.services.nas_service import get_configured_nas_roots, resolve_nas_path
from app.utils import normalize_docx_output_filename
from modules.auth_models import ROLE_ADMIN, user_has_role
from modules.docx_provenance import (
    PROVENANCE_PREVIEW_LABEL_PREFIX,
    build_provenance_cache_payload,
    create_provenance_preview_docx,
    extract_provenance_block_trace,
    load_cached_provenance_payload,
)
from modules.file_copier import copy_files

tasks_bp = Blueprint("tasks_bp", __name__, template_folder="templates")
_INVALID_UPLOAD_FILENAME_CHARS = '\\/:*?"<>|'
_WINDOWS_RESERVED_FILE_NAMES = {
    "CON",
    "PRN",
    "AUX",
    "NUL",
    "COM1",
    "COM2",
    "COM3",
    "COM4",
    "COM5",
    "COM6",
    "COM7",
    "COM8",
    "COM9",
    "LPT1",
    "LPT2",
    "LPT3",
    "LPT4",
    "LPT5",
    "LPT6",
    "LPT7",
    "LPT8",
    "LPT9",
}

_LIBREOFFICE_CANDIDATES = (
    "soffice",
    "libreoffice",
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
)
_PROVENANCE_PREVIEW_DOCX_CACHE_VERSION = 5
_PAGE_SOURCE_MAP_CACHE_VERSION = 14
_HTML_PREVIEW_CACHE_VERSION = 2


def _get_actor_info():
    if current_user and getattr(current_user, "is_authenticated", False):
        display_name = (getattr(current_user, "display_name", "") or "").strip()
        chinese_only = "".join(re.findall(r"[\u4e00-\u9fff\u3400-\u4dbf\uF900-\uFAFF]+", display_name))
        work_id = (getattr(current_user, "work_id", "") or "").strip()
        if chinese_only:
            label = f"{work_id} {chinese_only}" if work_id else chinese_only
        else:
            label = display_name or work_id
        return work_id, label
    return "", ""


def _safe_uploaded_filename(filename: str, default_stem: str = "upload") -> str:
    raw_name = os.path.basename((filename or "").replace("\\", "/")).strip()
    secured = secure_filename(raw_name)
    raw_stem, raw_ext = os.path.splitext(raw_name)
    secured_raw_stem = secure_filename(raw_stem) if raw_stem else ""
    if secured and (not raw_stem or secured_raw_stem):
        return secured

    cleaned = "".join(
        "_" if (ord(ch) < 32 or ch in _INVALID_UPLOAD_FILENAME_CHARS) else ch
        for ch in raw_name
    ).strip().strip(".")
    if cleaned in {"", ".", ".."}:
        cleaned = default_stem

    stem, ext = os.path.splitext(cleaned)
    if not stem:
        stem = default_stem
    if stem.upper() in _WINDOWS_RESERVED_FILE_NAMES:
        stem = f"_{stem}"
    return f"{stem}{ext}" if ext else stem


def _apply_last_edit(meta: dict) -> None:
    work_id, label = _get_actor_info()
    meta["last_edited"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    if label:
        meta["last_editor"] = label
    if work_id:
        meta["last_editor_work_id"] = work_id


def _find_libreoffice_binary() -> str | None:
    for candidate in _LIBREOFFICE_CANDIDATES:
        resolved = shutil.which(candidate) if not os.path.isabs(candidate) else candidate
        if resolved and os.path.isfile(resolved):
            return resolved
    return None


def _build_preview_pdf_name(source_path: str) -> str:
    stem = secure_filename(Path(source_path).stem) or "preview"
    digest = uuid.uuid5(uuid.NAMESPACE_URL, os.path.abspath(source_path)).hex[:10]
    return f"{stem}_{digest}.pdf"


def _ensure_pdf_preview(source_path: str, job_dir: str, subdir: str) -> tuple[str | None, str | None]:
    if not source_path or not os.path.isfile(source_path):
        return None, "找不到要預覽的文件"

    output_dir = os.path.join(job_dir, subdir)
    os.makedirs(output_dir, exist_ok=True)
    pdf_name = _build_preview_pdf_name(source_path)
    pdf_rel = os.path.join(subdir, pdf_name).replace("\\", "/")
    pdf_path = os.path.join(job_dir, pdf_rel)

    try:
        if os.path.exists(pdf_path) and os.path.getmtime(pdf_path) >= os.path.getmtime(source_path):
            return pdf_rel, None
    except OSError:
        pass

    libreoffice_bin = _find_libreoffice_binary()
    if not libreoffice_bin:
        return None, "找不到 LibreOffice，無法建立 PDF 預覽"

    try:
        with tempfile.TemporaryDirectory(prefix="lo_pdf_") as temp_dir:
            result = subprocess.run(
                [
                    libreoffice_bin,
                    "--headless",
                    "--convert-to",
                    "pdf:writer_pdf_Export",
                    "--outdir",
                    temp_dir,
                    source_path,
                ],
                capture_output=True,
                text=True,
                timeout=180,
                check=False,
            )
            converted_pdf = os.path.join(temp_dir, f"{Path(source_path).stem}.pdf")
            if result.returncode != 0 or not os.path.exists(converted_pdf):
                stdout = (result.stdout or "").strip()
                stderr = (result.stderr or "").strip()
                current_app.logger.warning(
                    "LibreOffice PDF conversion failed for %s (code=%s, stdout=%r, stderr=%r)",
                    source_path,
                    result.returncode,
                    stdout,
                    stderr,
                )
                return None, "LibreOffice 轉 PDF 失敗"
            shutil.copyfile(converted_pdf, pdf_path)
            return pdf_rel, None
    except subprocess.TimeoutExpired:
        current_app.logger.warning("LibreOffice PDF conversion timed out for %s", source_path)
        return None, "LibreOffice 轉 PDF 逾時"
    except Exception:
        current_app.logger.exception("Unexpected error while converting %s to PDF preview", source_path)
        return None, "建立 PDF 預覽時發生錯誤"


def _ensure_html_preview(source_path: str, job_dir: str, subdir: str, base_name: str) -> tuple[str | None, str | None]:
    if not source_path or not os.path.isfile(source_path):
        return None, "找不到要預覽的文件"

    output_dir = os.path.join(job_dir, subdir)
    os.makedirs(output_dir, exist_ok=True)
    html_name = f"{base_name}.html"
    html_rel = os.path.join(subdir, html_name).replace("\\", "/")
    html_path = os.path.join(job_dir, html_rel)
    meta_path = os.path.join(output_dir, "_meta.json")

    try:
        if os.path.exists(html_path) and os.path.exists(meta_path):
            with open(meta_path, "r", encoding="utf-8") as meta_file:
                meta = json.load(meta_file)
            if (
                meta.get("version") == _HTML_PREVIEW_CACHE_VERSION
                and os.path.getmtime(html_path) >= os.path.getmtime(source_path)
            ):
                return html_rel, None
    except OSError:
        pass
    except (ValueError, TypeError, json.JSONDecodeError):
        pass

    libreoffice_bin = _find_libreoffice_binary()
    if not libreoffice_bin:
        return None, "找不到 LibreOffice，無法建立 HTML 預覽"

    try:
        with tempfile.TemporaryDirectory(prefix="lo_html_") as temp_dir:
            result = subprocess.run(
                [
                    libreoffice_bin,
                    "--headless",
                    "--convert-to",
                    "html",
                    "--outdir",
                    temp_dir,
                    source_path,
                ],
                capture_output=True,
                text=True,
                timeout=180,
            )
            if result.returncode != 0:
                current_app.logger.warning(
                    "LibreOffice HTML conversion failed for %s (code=%s, stdout=%s, stderr=%s)",
                    source_path,
                    result.returncode,
                    result.stdout,
                    result.stderr,
                )
                return None, "LibreOffice 轉 HTML 失敗"

            converted_html = os.path.join(temp_dir, f"{Path(source_path).stem}.html")
            if not os.path.isfile(converted_html):
                current_app.logger.warning(
                    "LibreOffice HTML conversion did not produce %s for %s",
                    converted_html,
                    source_path,
                )
                return None, "找不到 LibreOffice 產生的 HTML 預覽"

            for entry in os.listdir(output_dir):
                entry_path = os.path.join(output_dir, entry)
                if os.path.isdir(entry_path):
                    shutil.rmtree(entry_path, ignore_errors=True)
                else:
                    try:
                        os.remove(entry_path)
                    except FileNotFoundError:
                        pass

            for entry in os.listdir(temp_dir):
                src_entry = os.path.join(temp_dir, entry)
                dst_entry = os.path.join(output_dir, entry)
                if os.path.isdir(src_entry):
                    shutil.copytree(src_entry, dst_entry, dirs_exist_ok=True)
                else:
                    shutil.copy2(src_entry, dst_entry)

            generated_html = os.path.join(output_dir, f"{Path(source_path).stem}.html")
            if generated_html != html_path and os.path.exists(generated_html):
                shutil.move(generated_html, html_path)

            _normalize_html_preview_alignment(html_path)

            with open(meta_path, "w", encoding="utf-8") as meta_file:
                json.dump({"version": _HTML_PREVIEW_CACHE_VERSION}, meta_file, ensure_ascii=False, indent=2)
        return html_rel, None
    except subprocess.TimeoutExpired:
        current_app.logger.warning("LibreOffice HTML conversion timed out for %s", source_path)
        return None, "LibreOffice 轉 HTML 逾時"
    except Exception:
        current_app.logger.exception("Unexpected error while converting %s to HTML preview", source_path)
        return None, "建立 HTML 預覽時發生錯誤"


def _normalize_html_preview_alignment(html_path: str) -> None:
    if not html_path or not os.path.isfile(html_path):
        return

    try:
        with open(html_path, "r", encoding="utf-8") as html_file:
            html = html_file.read()
    except OSError:
        return

    align_pattern = re.compile(
        r"<(?P<tag>[a-zA-Z0-9]+)(?P<before>[^>]*?)\salign=\"(?P<align>[^\"]+)\"(?P<after>[^>]*?)>",
        re.IGNORECASE,
    )

    def _replace_align(match: re.Match[str]) -> str:
        tag = match.group("tag")
        before = match.group("before") or ""
        align = (match.group("align") or "").strip().lower()
        after = match.group("after") or ""
        if not align:
            return match.group(0)

        attrs = f"{before}{after}"
        style_match = re.search(r'style=\"([^\"]*)\"', attrs, re.IGNORECASE)
        if style_match:
            style_value = style_match.group(1)
            if "text-align" in style_value.lower():
                return match.group(0)
            updated_style = f'{style_value.rstrip("; ")}; text-align: {align}'
            updated_attrs = attrs.replace(style_match.group(0), f'style="{updated_style}"', 1)
            return f"<{tag}{updated_attrs}>"

        return f'<{tag}{before} align="{align}" style="text-align: {align}"{after}>'

    normalized_html = align_pattern.sub(_replace_align, html)
    if normalized_html == html:
        return

    with open(html_path, "w", encoding="utf-8") as html_file:
        html_file.write(normalized_html)


def _ensure_provenance_preview_docx(
    result_docx: str,
    log_path: str,
    job_dir: str,
    source_lookup: dict[str, dict[str, object]],
) -> tuple[str | None, str | None]:
    if not result_docx or not os.path.isfile(result_docx):
        return None, "找不到結果文件"
    if not source_lookup:
        return None, None

    preview_rel = os.path.join("preview_trace", "provenance_preview.docx").replace("\\", "/")
    preview_path = os.path.join(job_dir, preview_rel)
    preview_meta_path = os.path.join(job_dir, "preview_trace", "provenance_preview.meta.json")

    try:
        preview_meta = {}
        if os.path.isfile(preview_meta_path):
            with open(preview_meta_path, "r", encoding="utf-8") as f:
                preview_meta = json.load(f) or {}
        if (
            os.path.isfile(preview_path)
            and os.path.getmtime(preview_path) >= os.path.getmtime(result_docx)
            and (not os.path.isfile(log_path) or os.path.getmtime(preview_path) >= os.path.getmtime(log_path))
            and int(preview_meta.get("version") or 0) == _PROVENANCE_PREVIEW_DOCX_CACHE_VERSION
        ):
            return preview_rel, None
    except OSError:
        pass
    except Exception:
        current_app.logger.warning("Failed to load provenance preview cache metadata for %s", job_dir, exc_info=True)

    try:
        if create_provenance_preview_docx(result_docx, preview_path, source_lookup):
            try:
                Path(preview_meta_path).write_text(
                    json.dumps({"version": _PROVENANCE_PREVIEW_DOCX_CACHE_VERSION}, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                )
            except Exception:
                current_app.logger.warning("Failed to save provenance preview cache metadata for %s", job_dir, exc_info=True)
            return preview_rel, None
        return None, "建立來源標記預覽失敗"
    except Exception:
        current_app.logger.exception("Unexpected error while building provenance preview for %s", result_docx)
        return None, "建立來源標記預覽時發生錯誤"


def _normalize_trace_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\xa0", " ")).strip()


def _format_source_file_label(path_raw: str) -> str:
    path = os.path.abspath(str(path_raw))
    file_path = Path(path)
    name = file_path.name or str(path_raw)
    parent_parts = [part for part in file_path.parts[-3:-1] if part not in {"", os.path.sep}]
    context = "/".join(parent_parts)
    if context:
        return f"{name} ({context})"
    return name


def _trace_source_label(entry: dict) -> str:
    params = entry.get("params") or {}
    input_file = params.get("input_file")
    if input_file:
        return _format_source_file_label(str(input_file))

    stype = str(entry.get("type") or "").strip()
    if stype == "insert_text":
        return "流程文字"
    if stype in {"insert_roman_heading", "insert_numbered_heading", "insert_bulleted_heading"}:
        return "流程標題"
    if stype == "template_merge":
        template_file = entry.get("template_file")
        if template_file:
            return _format_source_file_label(str(template_file))
    return stype or "未知來源"


def _extract_docx_trace_paragraphs(
    docx_path: str,
    *,
    hide_set: set[str] | None = None,
) -> list[dict[str, object]]:
    from docx import Document as DocxDocument

    items: list[dict[str, object]] = []
    hide_values = hide_set or set()
    doc = DocxDocument(docx_path)
    for idx, para in enumerate(doc.paragraphs):
        raw_text = (para.text or "").strip()
        normalized = _normalize_trace_text(raw_text)
        if not normalized or normalized in hide_values:
            continue
        items.append(
            {
                "paragraph_index": idx,
                "text": raw_text,
                "normalized_text": normalized,
            }
        )
    return items


def _extract_docx_table_texts(
    docx_path: str,
    *,
    hide_set: set[str] | None = None,
) -> list[str]:
    from docx import Document as DocxDocument

    hide_values = hide_set or set()
    doc = DocxDocument(docx_path)
    texts: list[str] = []
    seen: set[str] = set()

    for table in doc.tables:
        for row in table.rows:
            row_parts: list[str] = []
            for cell in row.cells:
                cell_text = _normalize_trace_text(cell.text or "")
                if not cell_text or cell_text in hide_values:
                    continue
                row_parts.append(cell_text)
                if len(cell_text) >= 20 and cell_text not in seen:
                    seen.add(cell_text)
                    texts.append(cell_text)
            row_text = _normalize_trace_text(" ".join(row_parts))
            if len(row_text) >= 24 and row_text not in seen:
                seen.add(row_text)
                texts.append(row_text)

    return texts


def _build_object_trace_candidates(
    entries: list[dict],
    titles_to_hide: list[str],
) -> list[dict[str, object]]:
    hide_set = {_normalize_trace_text(title) for title in titles_to_hide if _normalize_trace_text(title)}
    candidates: list[dict[str, object]] = []

    for entry in entries:
        if not isinstance(entry, dict):
            continue
        if entry.get("status") == "error":
            continue

        stype = str(entry.get("type") or "")
        if stype not in {"extract_specific_figure_from_word", "extract_specific_table_from_word"}:
            continue

        params = entry.get("params") or {}
        source_file = _trace_source_label(entry)
        output_docx = entry.get("output_docx")
        primary_probe_texts: list[str] = []
        fallback_probe_texts: list[str] = []

        if output_docx and os.path.isfile(str(output_docx)):
            try:
                paragraphs = _extract_docx_trace_paragraphs(str(output_docx), hide_set=hide_set)
                primary_probe_texts.extend(
                    [str(item["normalized_text"]) for item in paragraphs if item.get("normalized_text")]
                )
                if stype == "extract_specific_table_from_word":
                    fallback_probe_texts.extend(_extract_docx_table_texts(str(output_docx), hide_set=hide_set))
            except Exception:
                current_app.logger.warning("Failed to extract object trace paragraphs from %s", output_docx, exc_info=True)

        for part in (
            params.get("target_caption_label"),
            params.get("target_figure_title"),
            params.get("target_table_title"),
            params.get("target_chapter_title"),
            params.get("target_subtitle"),
        ):
            synthetic_text = _normalize_trace_text(str(part).strip())
            if synthetic_text:
                primary_probe_texts.append(synthetic_text)

        def _dedupe_probes(raw_probes: list[str]) -> list[str]:
            deduped: list[str] = []
            seen: set[str] = set()
            for probe in raw_probes:
                cleaned = _normalize_trace_text(probe)
                if not cleaned or cleaned in seen:
                    continue
                seen.add(cleaned)
                deduped.append(cleaned)
            return deduped

        candidates.append(
            {
                "source_file": source_file,
                "source_step": stype,
                "primary_probe_texts": _dedupe_probes(primary_probe_texts),
                "fallback_probe_texts": _dedupe_probes(fallback_probe_texts),
                "allow_multi_page": stype == "extract_specific_table_from_word",
            }
        )

    return candidates


def _build_compare_source_label(entry: dict) -> str:
    params = entry.get("params") or {}
    base = _trace_source_label(entry)
    stype = str(entry.get("type") or "")

    if stype == "extract_specific_figure_from_word":
        label = params.get("target_caption_label") or params.get("target_figure_title") or "圖片"
        return f"{base} 圖片 {label}".strip()
    if stype == "extract_specific_table_from_word":
        label = params.get("target_caption_label") or params.get("target_table_title") or "表格"
        return f"{base} 表格 {label}".strip()
    return base


def _collect_provenance_probe_texts(entry: dict) -> list[str]:
    params = entry.get("params") or {}
    texts: list[str] = []
    for raw_value in (
        params.get("target_caption_label"),
        params.get("target_figure_title"),
        params.get("target_table_title"),
        params.get("target_chapter_title"),
        params.get("target_title_section"),
        params.get("target_subtitle"),
        params.get("subheading_text"),
    ):
        normalized = _normalize_trace_text(str(raw_value or ""))
        if normalized:
            texts.append(normalized)
    deduped: list[str] = []
    seen: set[str] = set()
    for text in texts:
        if text in seen:
            continue
        seen.add(text)
        deduped.append(text)
    return deduped


def _build_provenance_source_lookup(entries: list[dict]) -> dict[str, dict[str, object]]:
    lookup: dict[str, dict[str, object]] = {}
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        provenance = entry.get("provenance")
        if not isinstance(provenance, dict):
            continue
        source_id = str(provenance.get("source_id") or "").strip()
        if not source_id:
            continue
        lookup[source_id] = {
            "source_id": source_id,
            "source_file": _trace_source_label(entry),
            "source_step": str(entry.get("type") or ""),
            "content_type": str(provenance.get("content_type") or ""),
            "bookmark_start": str(provenance.get("bookmark_start") or ""),
            "bookmark_end": str(provenance.get("bookmark_end") or ""),
            "bookmark_id": provenance.get("bookmark_id"),
            "fragment_path": str(provenance.get("fragment_path") or entry.get("output_docx") or ""),
            "fragment_order": provenance.get("fragment_order"),
            "template_index": provenance.get("template_index", entry.get("template_index")),
            "template_mode": provenance.get("template_mode", entry.get("template_mode")),
            "primary_probe_texts": _collect_provenance_probe_texts(entry),
        }
    return lookup


def _build_trace_from_provenance_blocks(
    block_trace: list[dict[str, object]],
) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    paragraph_trace: list[dict[str, object]] = []
    object_candidates: list[dict[str, object]] = []

    for block in block_trace:
        block_type = str(block.get("block_type") or "")
        source_file = str(block.get("source_file") or "未知來源")
        source_step = str(block.get("source_step") or "")
        content_type = str(block.get("content_type") or "")
        probe_texts = [str(item) for item in (block.get("probe_texts") or []) if str(item).strip()]
        if block_type == "paragraph":
            text = str(block.get("text") or "").strip()
            if not text:
                continue
            source_id = str(block.get("source_id") or "").strip()
            paragraph_trace.append(
                {
                    "merged_paragraph_index": len(paragraph_trace),
                    "source_id": source_id,
                    "source_file": source_file,
                    "source_paragraph_index": block.get("block_index"),
                    "result_block_index": block.get("block_index"),
                    "source_step": source_step,
                    "content_type": content_type or "paragraph",
                    "match_status": "provenance" if source_id else "context",
                    "count_as_source": bool(source_id),
                    "text": text,
                    "probe_texts": probe_texts[:],
                }
            )
            continue

        if not probe_texts:
            continue
        object_candidates.append(
            {
                "source_file": source_file,
                "source_step": source_step,
                "source_id": str(block.get("source_id") or "").strip(),
                "content_type": content_type,
                "count_as_source": bool(block.get("source_id")),
                "result_block_index": block.get("block_index"),
                "primary_probe_texts": probe_texts[:6],
                "fallback_probe_texts": probe_texts[6:],
                "allow_multi_page": content_type == "table",
            }
        )

    return paragraph_trace, object_candidates


def _build_provenance_trace(
    job_dir: str,
    result_docx: str,
    log_path: str,
    entries: list[dict],
    titles_to_hide: list[str],
) -> tuple[list[dict[str, object]], list[dict[str, object]]] | None:
    trace_dir = os.path.join(job_dir, "preview_trace")
    os.makedirs(trace_dir, exist_ok=True)
    cache_path = os.path.join(trace_dir, "provenance_map.json")
    source_lookup = _build_provenance_source_lookup(entries)
    if not source_lookup:
        return None

    cached_payload = None
    try:
        if (
            os.path.isfile(cache_path)
            and os.path.getmtime(cache_path) >= os.path.getmtime(result_docx)
            and os.path.getmtime(cache_path) >= os.path.getmtime(log_path)
        ):
            cached_payload = load_cached_provenance_payload(cache_path)
    except Exception:
        current_app.logger.warning("Failed to load cached provenance trace for %s", job_dir, exc_info=True)

    if cached_payload:
        block_trace = cached_payload.get("block_trace") or []
        if isinstance(block_trace, list):
            return _build_trace_from_provenance_blocks(block_trace)

    hide_set = {_normalize_trace_text(title) for title in titles_to_hide if _normalize_trace_text(title)}
    block_trace = extract_provenance_block_trace(result_docx, source_lookup, hide_set=hide_set)
    if not block_trace:
        return None

    try:
        payload = build_provenance_cache_payload(source_lookup=source_lookup, block_trace=block_trace)
        Path(cache_path).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        current_app.logger.warning("Failed to save provenance trace cache for %s", job_dir, exc_info=True)

    return _build_trace_from_provenance_blocks(block_trace)


def _build_paragraph_trace(
    job_dir: str,
    result_docx: str,
    log_path: str,
    entries: list[dict],
    titles_to_hide: list[str],
) -> list[dict[str, object]]:
    trace_dir = os.path.join(job_dir, "preview_trace")
    os.makedirs(trace_dir, exist_ok=True)
    trace_path = os.path.join(trace_dir, "paragraph_trace.json")
    hide_set = {_normalize_trace_text(title) for title in titles_to_hide if _normalize_trace_text(title)}

    try:
        if (
            os.path.isfile(trace_path)
            and os.path.getmtime(trace_path) >= os.path.getmtime(result_docx)
            and os.path.getmtime(trace_path) >= os.path.getmtime(log_path)
        ):
            with open(trace_path, "r", encoding="utf-8") as f:
                cached = json.load(f)
            if isinstance(cached, list):
                return cached
    except Exception:
        current_app.logger.warning("Failed to load cached paragraph trace for %s", job_dir)

    final_paragraphs = _extract_docx_trace_paragraphs(result_docx, hide_set=hide_set)
    candidates: list[dict[str, object]] = []

    for entry in entries:
        if not isinstance(entry, dict):
            continue
        if entry.get("status") == "error":
            continue
        if entry.get("type") in {"file_metadata", "template_merge", "copy_files", "copy_directory"}:
            continue
        output_docx = entry.get("output_docx")
        if not output_docx or not os.path.isfile(str(output_docx)):
            continue

        source_file = _trace_source_label(entry)
        source_step = str(entry.get("type") or "")
        try:
            paragraphs = _extract_docx_trace_paragraphs(str(output_docx), hide_set=hide_set)
        except Exception:
            current_app.logger.warning("Failed to extract trace paragraphs from %s", output_docx, exc_info=True)
            continue

        for para in paragraphs:
            candidates.append(
                {
                    "source_file": source_file,
                    "source_step": source_step,
                    "source_paragraph_index": para["paragraph_index"],
                    "text": para["text"],
                    "normalized_text": para["normalized_text"],
                }
            )

    trace: list[dict[str, object]] = []
    pointer = 0
    for merged_index, para in enumerate(final_paragraphs):
        normalized = str(para["normalized_text"])
        matched: dict[str, object] | None = None
        matched_index: int | None = None

        for idx in range(pointer, len(candidates)):
            if candidates[idx]["normalized_text"] == normalized:
                matched = candidates[idx]
                matched_index = idx
                break

        if matched_index is not None:
            pointer = matched_index + 1

        trace.append(
            {
                "merged_paragraph_index": merged_index,
                "source_file": matched["source_file"] if matched else "未知來源",
                "source_paragraph_index": matched["source_paragraph_index"] if matched else None,
                "source_step": matched["source_step"] if matched else "",
                "match_status": "matched" if matched else "unmatched",
                "text": para["text"],
            }
        )

    try:
        with open(trace_path, "w", encoding="utf-8") as f:
            json.dump(trace, f, ensure_ascii=False, indent=2)
    except Exception:
        current_app.logger.warning("Failed to save paragraph trace cache for %s", job_dir, exc_info=True)

    return trace


def _build_trace_text_probes(normalized_text: str) -> list[str]:
    text = (normalized_text or "").strip()
    if not text:
        return []
    probes: list[str] = []
    if len(text) <= 180:
        probes.append(text)
    for size in (180, 120, 80, 48):
        if len(text) > size:
            probes.append(text[:size].strip())
    deduped: list[str] = []
    seen: set[str] = set()
    for probe in probes:
        cleaned = probe.strip()
        if len(cleaned) < 12 or cleaned in seen:
            continue
        seen.add(cleaned)
        deduped.append(cleaned)
    return deduped


def _select_page_sources_for_display(
    ordered_sources: list[tuple[str, int]],
    *,
    inherited_from_previous: bool = False,
    preserve_sources: set[str] | None = None,
) -> list[tuple[str, int]]:
    del inherited_from_previous
    del preserve_sources
    return [
        (source_file, count)
        for source_file, count in ordered_sources
        if str(source_file).strip() and int(count) > 0
    ]


def _order_page_sources_by_first_seen(
    sources: list[tuple[str, int]],
    first_seen_order: dict[str, int] | None = None,
) -> list[tuple[str, int]]:
    if not sources:
        return []

    first_seen = first_seen_order or {}
    default_order = len(first_seen) + 1
    return sorted(
        sources,
        key=lambda item: (
            first_seen.get(item[0], default_order),
            item[0],
        ),
    )


def _page_has_explicit_paragraph_sources(source_counts: dict[str, int]) -> bool:
    return any(count > 0 for count in source_counts.values())


def _score_probe_matches(page_text: str, probe_texts: list[str]) -> int:
    normalized_page_text = _normalize_trace_text(page_text)
    score = 0
    seen: set[str] = set()
    for raw_probe in probe_texts:
        for probe in _build_trace_text_probes(_normalize_trace_text(str(raw_probe))):
            if probe in seen:
                continue
            if probe in normalized_page_text:
                seen.add(probe)
                score += min(len(probe), 120)
    return score


def _extract_preview_page_sources(
    page_texts: list[str],
    source_lookup: dict[str, dict[str, object]] | None = None,
) -> list[list[str]]:
    unique_sources: list[str] = []
    seen_sources: set[str] = set()
    for meta in (source_lookup or {}).values():
        source_file = str(meta.get("source_file") or "").strip()
        if not source_file or source_file in seen_sources:
            continue
        seen_sources.add(source_file)
        unique_sources.append(source_file)

    label_patterns = [
        (source_file, _normalize_trace_text(f"{PROVENANCE_PREVIEW_LABEL_PREFIX}{source_file}"))
        for source_file in unique_sources
    ]

    page_sources: list[list[str]] = []
    for page_text in page_texts:
        normalized_page_text = _normalize_trace_text(page_text)
        ordered_hits: list[tuple[int, str]] = []
        for source_file, pattern in label_patterns:
            if not pattern:
                continue
            start = 0
            while True:
                position = normalized_page_text.find(pattern, start)
                if position < 0:
                    break
                ordered_hits.append((position, source_file))
                start = position + max(1, len(pattern))
        ordered_hits.sort(key=lambda item: (item[0], item[1]))
        page_sources.append([source_file for _, source_file in ordered_hits])

    return page_sources


def _merge_page_source_map_with_preview_labels(
    page_source_map: list[dict[str, object]],
    preview_page_sources: list[list[str]],
) -> list[dict[str, object]]:
    if not page_source_map:
        return []

    merged_map: list[dict[str, object]] = []
    for idx, bucket in enumerate(page_source_map):
        updated_bucket = dict(bucket)
        raw_label_sources = preview_page_sources[idx] if idx < len(preview_page_sources) else []
        label_sources: list[str] = []
        for source_file in raw_label_sources:
            source_key = str(source_file).strip()
            if not source_key:
                continue
            label_sources.append(source_key)

        updated_bucket["preview_label_sequence"] = label_sources[:]

        existing_sources = [
            str(item.get("source_file") or "").strip()
            for item in (updated_bucket.get("sources") or [])
            if str(item.get("source_file") or "").strip()
        ]

        if label_sources and not existing_sources:
            updated_bucket["sources"] = [
                {
                    "source_file": source_key,
                    "count": 1,
                    "inherited": False,
                    "from_preview_label": True,
                    "preview_segment_role": "label",
                }
                for source_key in label_sources
            ]
            updated_bucket["dominant_source"] = label_sources[0]
            updated_bucket["uses_preview_labels"] = True
            updated_bucket["preview_has_source_switch"] = len(label_sources) > 1
        merged_map.append(updated_bucket)

    return merged_map


def _has_provenance_result_blocks(
    paragraph_trace: list[dict[str, object]],
    object_trace_candidates: list[dict[str, object]] | None = None,
) -> bool:
    if any(item.get("result_block_index") is not None for item in paragraph_trace):
        return True
    return any(
        candidate.get("result_block_index") is not None
        for candidate in (object_trace_candidates or [])
    )


def _build_provenance_page_mapping_blocks(
    paragraph_trace: list[dict[str, object]],
    object_trace_candidates: list[dict[str, object]] | None = None,
) -> list[dict[str, object]]:
    blocks: list[dict[str, object]] = []

    for item in paragraph_trace:
        block_index = item.get("result_block_index")
        if block_index is None:
            continue
        text = _normalize_trace_text(str(item.get("text") or ""))
        probe_texts = [
            _normalize_trace_text(str(probe))
            for probe in (item.get("probe_texts") or [])
            if _normalize_trace_text(str(probe))
        ]
        if text and text not in probe_texts:
            probe_texts.insert(0, text)
        blocks.append(
            {
                "block_index": int(block_index),
                "block_type": str(item.get("content_type") or "paragraph"),
                "source_id": str(item.get("source_id") or "").strip(),
                "source_file": str(item.get("source_file") or "未知來源"),
                "source_step": str(item.get("source_step") or ""),
                "count_as_source": bool(item.get("count_as_source")),
                "text": str(item.get("text") or ""),
                "primary_probe_texts": probe_texts,
                "fallback_probe_texts": [],
                "allow_multi_page": False,
            }
        )

    for candidate in object_trace_candidates or []:
        block_index = candidate.get("result_block_index")
        if block_index is None:
            continue
        primary_probe_texts = [
            _normalize_trace_text(str(probe))
            for probe in (candidate.get("primary_probe_texts") or [])
            if _normalize_trace_text(str(probe))
        ]
        fallback_probe_texts = [
            _normalize_trace_text(str(probe))
            for probe in (candidate.get("fallback_probe_texts") or [])
            if _normalize_trace_text(str(probe))
        ]
        blocks.append(
            {
                "block_index": int(block_index),
                "block_type": str(candidate.get("content_type") or "object"),
                "source_id": str(candidate.get("source_id") or "").strip(),
                "source_file": str(candidate.get("source_file") or "未知來源"),
                "source_step": str(candidate.get("source_step") or ""),
                "count_as_source": bool(candidate.get("count_as_source")),
                "text": "",
                "primary_probe_texts": primary_probe_texts,
                "fallback_probe_texts": fallback_probe_texts,
                "allow_multi_page": bool(candidate.get("allow_multi_page")),
            }
        )

    return sorted(blocks, key=lambda item: int(item.get("block_index") or -1))


def _select_result_block_pages(
    page_texts: list[str],
    *,
    start_page_idx: int,
    primary_probe_texts: list[str],
    fallback_probe_texts: list[str],
    allow_multi_page: bool,
) -> list[int]:
    if not page_texts:
        return []

    bounded_start = max(0, min(start_page_idx, len(page_texts) - 1))

    def _score_pages(probes: list[str]) -> dict[int, int]:
        return {
            page_idx: _score_probe_matches(page_texts[page_idx], probes)
            for page_idx in range(bounded_start, len(page_texts))
        }

    def _pick_pages(scores: dict[int, int], *, permit_multi_page: bool) -> list[int]:
        positive_scores = {page_idx: score for page_idx, score in scores.items() if score > 0}
        if not positive_scores:
            return []
        best_page = max(positive_scores, key=lambda page_idx: (positive_scores[page_idx], -page_idx))
        max_score = positive_scores[best_page]
        if not permit_multi_page:
            return [best_page]

        min_score = max(1, (max_score + 2) // 3)
        selected = {page_idx for page_idx, score in positive_scores.items() if score >= min_score}
        cluster = [best_page]

        next_page = best_page - 1
        while next_page in selected:
            cluster.insert(0, next_page)
            next_page -= 1

        next_page = best_page + 1
        while next_page in selected:
            cluster.append(next_page)
            next_page += 1

        return cluster

    selected_pages = _pick_pages(_score_pages(primary_probe_texts), permit_multi_page=allow_multi_page)
    if selected_pages:
        return selected_pages

    selected_pages = _pick_pages(_score_pages(fallback_probe_texts), permit_multi_page=allow_multi_page)
    if selected_pages:
        return selected_pages

    return [bounded_start]


def _build_page_source_map_from_provenance_blocks(
    page_texts: list[str],
    paragraph_trace: list[dict[str, object]],
    object_trace_candidates: list[dict[str, object]] | None = None,
) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    page_buckets: list[dict[str, object]] = [
        {
            "page_number": idx + 1,
            "dominant_source": "",
            "sources": [],
        }
        for idx in range(len(page_texts))
    ]
    source_counts_by_page: list[dict[str, int]] = [dict() for _ in page_texts]
    source_first_seen_by_page: list[dict[str, int]] = [dict() for _ in page_texts]
    has_context_only_content_by_page = [False for _ in page_texts]
    preserved_sources_by_page: list[set[str]] = [set() for _ in page_texts]
    block_page_map: dict[int, list[int]] = {}
    current_page_idx = 0
    source_seen_sequence = 0

    for block in _build_provenance_page_mapping_blocks(paragraph_trace, object_trace_candidates):
        selected_pages = _select_result_block_pages(
            page_texts,
            start_page_idx=current_page_idx,
            primary_probe_texts=[str(item) for item in (block.get("primary_probe_texts") or []) if str(item).strip()],
            fallback_probe_texts=[str(item) for item in (block.get("fallback_probe_texts") or []) if str(item).strip()],
            allow_multi_page=bool(block.get("allow_multi_page")),
        )
        if not selected_pages:
            continue

        block_page_map[int(block["block_index"])] = selected_pages
        current_page_idx = selected_pages[-1]

        source_file = str(block.get("source_file") or "未知來源")
        source_step = str(block.get("source_step") or "")
        count_as_source = bool(
            block.get("count_as_source", str(block.get("source_id") or "").strip())
        )

        for page_idx in selected_pages:
            if count_as_source and source_file not in {"", "未知來源"}:
                source_counts = source_counts_by_page[page_idx]
                source_counts[source_file] = source_counts.get(source_file, 0) + 1
                source_first_seen = source_first_seen_by_page[page_idx]
                if source_file not in source_first_seen:
                    source_first_seen[source_file] = source_seen_sequence
                    source_seen_sequence += 1
                if source_step in {"extract_specific_table_from_word", "extract_specific_figure_from_word"}:
                    preserved_sources_by_page[page_idx].add(source_file)
            else:
                has_context_only_content_by_page[page_idx] = True

    annotated_trace: list[dict[str, object]] = []
    for item in paragraph_trace:
        trace_item = dict(item)
        block_index = trace_item.get("result_block_index")
        assigned_pages = block_page_map.get(int(block_index)) if block_index is not None else None
        trace_item["result_page"] = assigned_pages[0] + 1 if assigned_pages else None
        annotated_trace.append(trace_item)

    for idx, bucket in enumerate(page_buckets):
        source_counts = source_counts_by_page[idx]
        ordered_sources = sorted(source_counts.items(), key=lambda item: (-item[1], item[0]))
        display_sources = _select_page_sources_for_display(
            ordered_sources,
            preserve_sources=preserved_sources_by_page[idx],
        )
        display_sources = _order_page_sources_by_first_seen(
            display_sources,
            source_first_seen_by_page[idx],
        )
        dominant_source = ordered_sources[0][0] if ordered_sources else ""
        bucket["sources"] = [
            {
                "source_file": source_file,
                "count": count,
                "inherited": False,
            }
            for source_file, count in display_sources
        ]
        bucket["dominant_source"] = dominant_source

    return annotated_trace, page_buckets


def _select_object_candidate_pages(
    page_texts: list[str],
    source_counts_by_page: list[dict[str, int]],
    *,
    primary_probe_texts: list[str],
    fallback_probe_texts: list[str],
    allow_multi_page: bool,
) -> list[int]:
    available_pages = [
        page_idx
        for page_idx in range(len(page_texts))
        if not _page_has_explicit_paragraph_sources(source_counts_by_page[page_idx])
    ]
    if not available_pages:
        return []

    def _score_pages(probes: list[str]) -> dict[int, int]:
        return {
            page_idx: _score_probe_matches(page_texts[page_idx], probes)
            for page_idx in available_pages
        }

    def _pick_pages(scores: dict[int, int], *, permit_multi_page: bool) -> list[int]:
        positive_scores = {page_idx: score for page_idx, score in scores.items() if score > 0}
        if not positive_scores:
            return []
        best_page = max(positive_scores, key=lambda page_idx: (positive_scores[page_idx], -page_idx))
        max_score = positive_scores[best_page]
        if not permit_multi_page:
            return [best_page]

        min_score = max(1, (max_score + 2) // 3)
        selected = {page_idx for page_idx, score in positive_scores.items() if score >= min_score}
        cluster = [best_page]

        next_page = best_page - 1
        while next_page in selected:
            cluster.insert(0, next_page)
            next_page -= 1

        next_page = best_page + 1
        while next_page in selected:
            cluster.append(next_page)
            next_page += 1

        return cluster

    primary_scores = _score_pages(primary_probe_texts)
    selected_pages = _pick_pages(primary_scores, permit_multi_page=allow_multi_page)
    if selected_pages:
        return selected_pages

    fallback_scores = _score_pages(fallback_probe_texts)
    return _pick_pages(fallback_scores, permit_multi_page=allow_multi_page)


def _build_page_source_map(
    job_dir: str,
    result_pdf_path: str,
    paragraph_trace: list[dict[str, object]],
    object_trace_candidates: list[dict[str, object]] | None = None,
    source_lookup: dict[str, dict[str, object]] | None = None,
) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    if not result_pdf_path or not os.path.isfile(result_pdf_path):
        return paragraph_trace, []

    trace_dir = os.path.join(job_dir, "preview_trace")
    os.makedirs(trace_dir, exist_ok=True)
    page_map_path = os.path.join(trace_dir, "page_source_map.json")

    try:
        if os.path.isfile(page_map_path) and os.path.getmtime(page_map_path) >= os.path.getmtime(result_pdf_path):
            with open(page_map_path, "r", encoding="utf-8") as f:
                payload = json.load(f)
            if payload.get("version") != _PAGE_SOURCE_MAP_CACHE_VERSION:
                raise ValueError("stale page source map cache")
            cached_trace = payload.get("paragraph_trace")
            cached_pages = payload.get("page_source_map")
            if isinstance(cached_trace, list) and isinstance(cached_pages, list):
                return cached_trace, cached_pages
    except Exception:
        current_app.logger.warning("Failed to load cached page source map for %s", job_dir, exc_info=True)

    import fitz

    with fitz.open(result_pdf_path) as pdf:
        page_texts = [
            _normalize_trace_text(page.get_text("text"))
            for page in pdf
        ]
    preview_page_sources = _extract_preview_page_sources(page_texts, source_lookup)

    if _has_provenance_result_blocks(paragraph_trace, object_trace_candidates):
        annotated_trace, page_buckets = _build_page_source_map_from_provenance_blocks(
            page_texts,
            paragraph_trace,
            object_trace_candidates,
        )
        page_buckets = _merge_page_source_map_with_preview_labels(page_buckets, preview_page_sources)
        try:
            with open(page_map_path, "w", encoding="utf-8") as f:
                json.dump(
                    {
                        "version": _PAGE_SOURCE_MAP_CACHE_VERSION,
                        "paragraph_trace": annotated_trace,
                        "page_source_map": page_buckets,
                    },
                    f,
                    ensure_ascii=False,
                    indent=2,
                )
        except Exception:
            current_app.logger.warning("Failed to save page source map cache for %s", job_dir, exc_info=True)
        return annotated_trace, page_buckets

    page_buckets: list[dict[str, object]] = [
        {
            "page_number": idx + 1,
            "dominant_source": "",
            "sources": [],
        }
        for idx in range(len(page_texts))
    ]
    source_counts_by_page: list[dict[str, int]] = [dict() for _ in page_texts]
    source_first_seen_by_page: list[dict[str, int]] = [dict() for _ in page_texts]
    has_context_only_content_by_page = [False for _ in page_texts]
    preserved_sources_by_page: list[set[str]] = [set() for _ in page_texts]
    annotated_trace: list[dict[str, object]] = []
    current_page_idx = 0
    source_seen_sequence = 0

    for item in paragraph_trace:
        trace_item = dict(item)
        normalized = _normalize_trace_text(str(trace_item.get("text") or ""))
        assigned_page_idx = current_page_idx if page_texts else -1
        if normalized and page_texts:
            probes = _build_trace_text_probes(normalized)
            found_idx = None
            for page_idx in range(current_page_idx, len(page_texts)):
                page_text = page_texts[page_idx]
                if any(probe in page_text for probe in probes):
                    found_idx = page_idx
                    break
            if found_idx is not None:
                assigned_page_idx = found_idx
                current_page_idx = found_idx

        trace_item["result_page"] = assigned_page_idx + 1 if assigned_page_idx >= 0 else None
        annotated_trace.append(trace_item)

        if assigned_page_idx >= 0:
            count_as_source = bool(
                trace_item.get("count_as_source", str(trace_item.get("source_file") or "").strip() not in {"", "未知來源"})
            )
            source_file = str(trace_item.get("source_file") or "未知來源")
            source_step = str(trace_item.get("source_step") or "")
            if count_as_source and source_file not in {"", "未知來源"}:
                source_counts = source_counts_by_page[assigned_page_idx]
                source_counts[source_file] = source_counts.get(source_file, 0) + 1
                source_first_seen = source_first_seen_by_page[assigned_page_idx]
                if source_file not in source_first_seen:
                    source_first_seen[source_file] = source_seen_sequence
                    source_seen_sequence += 1
                if source_step in {"extract_specific_table_from_word", "extract_specific_figure_from_word"}:
                    preserved_sources_by_page[assigned_page_idx].add(source_file)
            else:
                has_context_only_content_by_page[assigned_page_idx] = True

    object_candidates = object_trace_candidates or []
    for candidate in object_candidates:
        if not page_texts:
            continue

        matched_pages = _select_object_candidate_pages(
            page_texts,
            source_counts_by_page,
            primary_probe_texts=[str(item) for item in (candidate.get("primary_probe_texts") or [])],
            fallback_probe_texts=[str(item) for item in (candidate.get("fallback_probe_texts") or [])],
            allow_multi_page=bool(candidate.get("allow_multi_page")),
        )
        if not matched_pages:
            continue

        source_file = str(candidate.get("source_file") or "未知來源")
        for page_idx in matched_pages:
            source_counts = source_counts_by_page[page_idx]
            source_counts[source_file] = source_counts.get(source_file, 0) + 1
            source_first_seen = source_first_seen_by_page[page_idx]
            if source_file not in source_first_seen:
                source_first_seen[source_file] = source_seen_sequence
                source_seen_sequence += 1

    for idx, bucket in enumerate(page_buckets):
        source_counts = source_counts_by_page[idx]
        ordered_sources = sorted(source_counts.items(), key=lambda item: (-item[1], item[0]))
        display_sources = _select_page_sources_for_display(
            ordered_sources,
            preserve_sources=preserved_sources_by_page[idx],
        )
        display_sources = _order_page_sources_by_first_seen(
            display_sources,
            source_first_seen_by_page[idx],
        )
        dominant_source = ordered_sources[0][0] if ordered_sources else ""
        bucket["sources"] = [
            {
                "source_file": source_file,
                "count": count,
                "inherited": False,
            }
            for source_file, count in display_sources
        ]
        bucket["dominant_source"] = dominant_source

    page_buckets = _merge_page_source_map_with_preview_labels(page_buckets, preview_page_sources)

    try:
        with open(page_map_path, "w", encoding="utf-8") as f:
            json.dump(
                {
                    "version": _PAGE_SOURCE_MAP_CACHE_VERSION,
                    "paragraph_trace": annotated_trace,
                    "page_source_map": page_buckets,
                },
                f,
                ensure_ascii=False,
                indent=2,
            )
    except Exception:
        current_app.logger.warning("Failed to save page source map cache for %s", job_dir, exc_info=True)

    return annotated_trace, page_buckets


def _get_creator_work_id(meta: dict) -> str:
    creator_work_id = (meta.get("creator_work_id") or "").strip()
    if creator_work_id:
        return creator_work_id
    creator = (meta.get("creator") or "").strip()
    if creator:
        return creator.split()[0]
    return ""


def _can_delete_task(meta: dict) -> bool:
    if not current_app.config.get("AUTH_ENABLED", True):
        return True
    if not current_user or not getattr(current_user, "is_authenticated", False):
        return False
    if user_has_role(current_user.id, ROLE_ADMIN):
        return True
    creator_work_id = _get_creator_work_id(meta)
    return bool(creator_work_id) and current_user.work_id == creator_work_id

def _load_task_context(task_id: str) -> dict:
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    meta_path = os.path.join(tdir, "meta.json")
    task = {"id": task_id}
    if os.path.exists(meta_path):
        try:
            with open(meta_path, "r", encoding="utf-8") as f:
                meta = json.load(f)
            task.update(
                {
                    "name": meta.get("name", task_id),
                    "description": meta.get("description", ""),
                    "creator": meta.get("creator", "") or "",
                    "nas_path": meta.get("nas_path", "") or "",
                }
            )
        except Exception:
            pass
    return task

@tasks_bp.route("/tasks/<task_id>/copy-files", methods=["GET", "POST"], endpoint="task_copy_files")
def task_copy_files(task_id):
    base = os.path.join(current_app.config["TASK_FOLDER"], task_id, "files")
    if not os.path.isdir(base):
        abort(404)

    def _safe_path(rel: str) -> str:
        norm = os.path.normpath(rel)
        if not rel or os.path.isabs(norm) or norm.startswith(".."):
            raise ValueError("資料夾名稱不合法")
        return os.path.join(base, norm)

    message = ""
    if request.method == "POST":
        action = request.form.get("action")
        if action == "create_dir":
            new_rel = request.form.get("new_dir", "").strip()
            try:
                os.makedirs(_safe_path(new_rel), exist_ok=True)
                message = f"已建立資料夾 {os.path.normpath(new_rel)}"
            except ValueError:
                message = "資料夾名稱不合法"
        else:
            source_rel = request.form.get("source_dir", "").strip()
            dest_rel = request.form.get("dest_dir", "").strip()
            keywords_raw = request.form.get("keywords", "")
            keywords = [k.strip() for k in keywords_raw.split(",") if k.strip()]
            if not source_rel or not dest_rel or not keywords:
                message = "請完整輸入資料"
            else:
                try:
                    src = _safe_path(source_rel)
                    dest = _safe_path(dest_rel)
                    copied = copy_files(src, dest, keywords)
                    message = f"已複製 {len(copied)} 個檔案"
                except ValueError:
                    message = "資料夾名稱不合法"
                except Exception as e:
                    message = str(e)
    dirs = list_dirs(base)
    dirs.insert(0, ".")
    return render_template(
        "tasks/copy_files.html",
        dirs=dirs,
        message=message,
        task_id=task_id,
        task=_load_task_context(task_id),
    )

@tasks_bp.route("/tasks/<task_id>/mapping", methods=["GET", "POST"], endpoint="task_mapping")
def task_mapping(task_id):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    if not os.path.isdir(tdir):
        abort(404)
    files_dir = os.path.join(tdir, "files")
    out_dir = os.path.join(tdir, "mapping_job")
    messages = []
    outputs = []
    log_file = None
    zip_file = None
    log_file_name = None
    zip_file_name = None
    step_runs = []
    last_mapping_marker = os.path.join(tdir, "mapping_last.txt")
    validation_state_path = os.path.join(tdir, "mapping_validation_state.json")
    last_mapping_file = None
    validation_state = {
        "mapping_file": "",
        "mapping_display_name": "",
        "reference_ok": False,
        "extract_ok": False,
        "run_id": "",
    }
    current_run_id = None
    current_mapping_display_name = ""

    # 如果是頁面跳轉/重新整理 (GET)，則清掉之前的暫存紀錄與檔案
    if request.method == "GET":
        if os.path.isfile(last_mapping_marker):
            try:
                cached_name = Path(last_mapping_marker).read_text(encoding="utf-8").strip()
                cached_path = os.path.join(tdir, cached_name)
                if cached_name and os.path.isfile(cached_path):
                    os.remove(cached_path)
                os.remove(last_mapping_marker)
            except Exception:
                pass
        try:
            if os.path.isfile(validation_state_path):
                os.remove(validation_state_path)
        except Exception:
            pass
    else:
        # POST 請求時才讀取暫存紀錄
        if os.path.isfile(last_mapping_marker):
            try:
                cached_name = Path(last_mapping_marker).read_text(encoding="utf-8").strip()
                cached_path = os.path.join(tdir, cached_name)
                if cached_name and os.path.isfile(cached_path):
                    last_mapping_file = cached_name
            except Exception:
                last_mapping_file = None
        if os.path.isfile(validation_state_path):
            try:
                loaded_state = json.loads(Path(validation_state_path).read_text(encoding="utf-8"))
                if isinstance(loaded_state, dict):
                    validation_state.update(
                        {
                            "mapping_file": str(loaded_state.get("mapping_file") or ""),
                            "mapping_display_name": str(loaded_state.get("mapping_display_name") or ""),
                            "reference_ok": bool(loaded_state.get("reference_ok")),
                            "extract_ok": bool(loaded_state.get("extract_ok")),
                            "run_id": str(loaded_state.get("run_id") or ""),
                        }
                    )
                    current_mapping_display_name = validation_state.get("mapping_display_name") or ""
            except Exception:
                pass

    def _format_step_label(entry: dict) -> tuple[str, str]:
        stype = entry.get("type") or ""
        params = entry.get("params") or {}

        def _boolish(value, default: bool = False) -> bool:
            if value in (None, ""):
                return default
            return str(value).strip().lower() in {"1", "true", "yes", "y", "on"}

        def _base(path: str) -> str:
            if not path: return "?"
            name = os.path.basename(path)
            # 移除 "Section 1_", "Section 2_" 等前綴
            import re
            name = re.sub(r"^Section\s+\d+_", "", name)
            return name

        row_no = params.get("mapping_row")
        row_prefix = f"(Row {row_no}) " if row_no not in (None, "", "None") else ""
        preset_action = (params.get("mapping_action_label") or "").strip()
        preset_detail = (params.get("mapping_detail_label") or "").strip()
        if preset_action:
            return f"{row_prefix}{preset_action}", preset_detail
        if stype == "extract_word_chapter":
            src = _base(params.get("input_file", ""))
            chapter_start = (params.get("target_chapter_section") or "").strip()
            chapter_end = (params.get("explicit_end_number") or "").strip()
            chapter = f"{chapter_start}-{chapter_end}" if chapter_start and chapter_end else chapter_start
            title = (params.get("target_chapter_title") or params.get("target_title_section") or "").strip()
            sub = (params.get("target_subtitle") or params.get("subheading_text") or "").strip()
            
            # 組合章節與標題: "1.1.1 General description"
            main_header = f"{chapter} {title}".strip()
            
            parts = [src]
            if main_header:
                parts.append(main_header)
            if sub:
                parts.append(sub)
            if _boolish(params.get("hide_chapter_title"), default=False):
                parts.append("不含標題")
            
            return f"{row_prefix}擷取章節", " | ".join(parts)
        if stype == "extract_word_all_content":
            src = _base(params.get("input_file", ""))
            return f"{row_prefix}擷取全文", src.strip()
        if stype == "extract_pdf_pages_as_images":
            src = _base(params.get("input_file", ""))
            pages = params.get("pages")
            parts = [src.strip()]
            if pages:
                parts.append(f"pages={pages}")
            return f"{row_prefix}擷取 PDF 圖片", " | ".join(p for p in parts if p)
        if stype == "extract_specific_table_from_word":
            src = _base(params.get("input_file", ""))
            section = (params.get("target_chapter_section") or "").strip()
            title = (params.get("target_chapter_title") or params.get("target_title_section") or "").strip()
            main_header = f"{section} {title}".strip()
            label = (
                params.get("target_caption_label")
                or params.get("target_table_label", "")
                or params.get("target_figure_label", "")
            ).strip()
            table_title = (params.get("target_table_title") or "").strip()
            table_index = str(params.get("target_table_index") or "").strip()
            parts = [src]
            if main_header:
                parts.append(main_header)
            if label:
                parts.append(label)
            if table_title:
                parts.append(f"title={table_title}")
            if table_index:
                parts.append(f"index={table_index}")
            if not _boolish(params.get("include_caption"), default=True):
                parts.append("不含標題")
            return f"{row_prefix}擷取表格", " | ".join(parts)
        if stype == "extract_specific_figure_from_word":
            src = _base(params.get("input_file", ""))
            section = (params.get("target_chapter_section") or "").strip()
            title = (params.get("target_chapter_title") or params.get("target_title_section") or "").strip()
            main_header = f"{section} {title}".strip()
            label = (
                params.get("target_caption_label")
                or params.get("target_figure_label", "")
                or params.get("target_table_label", "")
            ).strip()
            figure_title = (params.get("target_figure_title") or "").strip()
            figure_index = str(params.get("target_figure_index") or "").strip()
            parts = [src]
            if main_header:
                parts.append(main_header)
            if label:
                parts.append(label)
            if figure_title:
                parts.append(f"title={figure_title}")
            if figure_index:
                parts.append(f"index={figure_index}")
            if not _boolish(params.get("include_caption"), default=True):
                parts.append("不含標題")
            return f"{row_prefix}擷取圖片", " | ".join(parts)
        if stype == "insert_text":
            text_val = (params.get("text") or "").strip()
            return f"{row_prefix}插入文字", text_val
        if stype == "copy_file":
            src = _base(params.get("source", ""))
            dest = (params.get("destination") or "").strip().replace("\\", "/")
            target_name = (params.get("target_name") or "").strip()
            parts = [src]
            if target_name:
                parts.append(f"目標名稱={target_name}")
            if dest:
                parts.append(dest)
            return f"{row_prefix}複製檔案", " | ".join(p for p in parts if p)
        if stype == "copy_folder":
            src = _base(params.get("source", ""))
            dest = (params.get("destination") or "").strip().replace("\\", "/")
            target_name = (params.get("target_name") or "").strip()
            parts = [src]
            if target_name:
                parts.append(f"目標名稱={target_name}")
            if dest:
                parts.append(dest)
            return f"{row_prefix}複製資料夾", " | ".join(p for p in parts if p)
        if stype == "template_merge":
            tpl = _base(entry.get("template_file", ""))
            return f"{row_prefix}模版合併", tpl.strip()
        return f"{row_prefix}{stype or '步驟'}", ""

    def _truncate_detail(text: str, limit: int = 160) -> tuple[str, bool]:
        if len(text) <= limit:
            return text, False
        trimmed = text[: max(0, limit - 1)].rstrip()
        return f"{trimmed}…", True
    if request.method == "POST":
        action = request.form.get("action") or "run"
        mapping_path = None
        uploaded_new_mapping = False
        if action == "run_cached":
            if not last_mapping_file:
                messages.append("找不到上次檢查的檔案，請重新上傳。")
            else:
                mapping_path = os.path.join(tdir, last_mapping_file)
        else:
            f = request.files.get("mapping_file")
            if f and f.filename:
                display_name = os.path.basename((f.filename or "").replace("\\", "/")).strip()
                filename = _safe_uploaded_filename(
                    f.filename,
                    default_stem=f"mapping_{uuid.uuid4().hex[:8]}",
                )
                mapping_path = os.path.join(tdir, filename)
                f.save(mapping_path)
                uploaded_new_mapping = True
                current_mapping_display_name = display_name or filename
                try:
                    Path(last_mapping_marker).write_text(filename, encoding="utf-8")
                    last_mapping_file = filename
                    validation_state = {
                        "mapping_file": filename,
                        "mapping_display_name": current_mapping_display_name,
                        "reference_ok": False,
                        "extract_ok": False,
                        "run_id": "",
                    }
                except Exception:
                    pass
            elif last_mapping_file:
                mapping_path = os.path.join(tdir, last_mapping_file)
                current_mapping_display_name = current_mapping_display_name or last_mapping_file
            else:
                messages.append("請選擇檔案")

        if action == "check_extract" and not validation_state.get("reference_ok"):
            messages.append("請先通過檢查引用文件。")
            mapping_path = None
        if action == "run_cached" and not validation_state.get("extract_ok"):
            messages.append("請先通過檢查擷取參數。")
            mapping_path = None
        if mapping_path:
            try:
                from modules.mapping_processor import process_mapping_excel
                existing_run_id = str(validation_state.get("run_id") or "").strip()
                current_run_id = existing_run_id if (existing_run_id and not uploaded_new_mapping) else uuid.uuid4().hex[:8]
                run_out_dir = os.path.join(out_dir, current_run_id)
                result = process_mapping_excel(
                    mapping_path,
                    files_dir,
                    run_out_dir,
                    log_dir=run_out_dir,
                    validate_only=(action == "check"),
                    validate_extract_only=(action == "check_extract"),
                )
                messages = result["logs"]
                outputs = result["outputs"]
                log_file_raw = result.get("log_file")
                zip_file_raw = result.get("zip_file")
                log_file_name = log_file_raw
                zip_file_name = zip_file_raw
                log_file = f"{current_run_id}/{log_file_raw}" if log_file_raw else None
                zip_file = f"{current_run_id}/{zip_file_raw}" if zip_file_raw else None
                current_has_error = any("ERROR" in (m or "") for m in messages)
                current_mapping_name = os.path.basename(mapping_path)
                current_mapping_display_name = current_mapping_display_name or validation_state.get("mapping_display_name") or current_mapping_name
                if action == "check":
                    validation_state = {
                        "mapping_file": current_mapping_name,
                        "mapping_display_name": current_mapping_display_name,
                        "reference_ok": not current_has_error,
                        "extract_ok": False,
                        "run_id": current_run_id,
                    }
                elif action == "check_extract":
                    validation_state = {
                        "mapping_file": current_mapping_name,
                        "mapping_display_name": current_mapping_display_name,
                        "reference_ok": bool(validation_state.get("reference_ok")),
                        "extract_ok": not current_has_error,
                        "run_id": current_run_id,
                    }
                elif action == "run_cached":
                    validation_state = {
                        "mapping_file": current_mapping_name,
                        "mapping_display_name": current_mapping_display_name,
                        "reference_ok": bool(validation_state.get("reference_ok")),
                        "extract_ok": bool(validation_state.get("extract_ok")),
                        "run_id": current_run_id,
                    }
                if action in {"check", "check_extract"}:
                    try:
                        Path(validation_state_path).write_text(
                            json.dumps(validation_state, ensure_ascii=False),
                            encoding="utf-8",
                        )
                    except Exception:
                        pass
                elif action == "run_cached":
                    try:
                        Path(validation_state_path).write_text(
                            json.dumps(validation_state, ensure_ascii=False),
                            encoding="utf-8",
                        )
                    except Exception:
                        pass
            except Exception as e:
                messages = [str(e)]
    if log_file:
        log_candidates = [
            os.path.join(out_dir, log_file),
        ]
        log_path = next((p for p in log_candidates if os.path.isfile(p)), None)
        if log_path:
            try:
                with open(log_path, "r", encoding="utf-8") as f:
                    log_data = json.load(f)
                for run in log_data.get("runs", []):
                    for entry in run.get("workflow_log", []):
                        if "step" not in entry:
                            continue
                        action, detail = _format_step_label(entry)
                        row_no = (entry.get("params") or {}).get("mapping_row")
                        detail_short, detail_long = _truncate_detail(detail) if detail else ("", False)
                        step_runs.append(
                            {
                                "action": action,
                                "detail": detail,
                                "detail_short": detail_short,
                                "detail_long": detail_long,
                                "row_no": row_no,
                                "status": entry.get("status") or "ok",
                                "error": entry.get("error") or "",
                            }
                        )
                if step_runs:
                    messages = [m for m in messages if not (m or "").startswith("WF_ERROR:")]
            except Exception as e:
                messages.append(f"ERROR: failed to read log file ({e})")
    has_error = any("ERROR" in (m or "") for m in messages) or any(
        step.get("status") == "error" for step in step_runs
    )
    warning_messages = [m for m in messages if (m or "").startswith("WARN:") or (m or "").startswith("WARNING:")]
    has_warning = bool(warning_messages)
    warning_confirm = None
    if has_warning:
        trimmed = []
        for m in warning_messages[:3]:
            trimmed.append(m.replace("WARN:", "").replace("WARNING:", "").strip())
        warning_confirm = "Warnings found. Run anyway?\n" + "\n".join(trimmed)

    error_messages = [m for m in messages if (m or "").startswith("ERROR:")]
    if error_messages:
        def _norm_error_text(text: str) -> str:
            return re.sub(r"\s+", " ", (text or "").strip())

        existing_row_errors: dict[int | None, set[str]] = {}
        for step in step_runs:
            if step.get("status") != "error":
                continue
            row_no = step.get("row_no")
            bucket = existing_row_errors.setdefault(row_no, set())
            for candidate in (step.get("error"), step.get("detail")):
                normalized = _norm_error_text(str(candidate or ""))
                if normalized:
                    bucket.add(normalized)

        error_steps = []
        for msg in error_messages:
            raw = (msg or "").replace("ERROR:", "", 1).strip()
            raw = re.sub(r"^Row\s+\d+\s*:\s*", "", raw, flags=re.IGNORECASE)
            action = raw
            detail = ""
            error_text = raw
            row_match = re.search(r"Row\s+(\d+)", msg or "", re.IGNORECASE)
            row_prefix = f"(Row {row_match.group(1)}) " if row_match else ""
            if "::" in raw:
                parts = [p.strip() for p in raw.split("::", 2)]
                if len(parts) >= 2:
                    base_action = parts[0] or action
                    if base_action.startswith("(Row "):
                        action = base_action
                    else:
                        action = f"{row_prefix}{base_action}".strip()
                    detail = parts[1]
                if len(parts) == 3:
                    error_text = parts[2]
            elif ":" in raw:
                head, tail = raw.split(":", 1)
                base_action = head.strip() or raw
                if base_action.startswith("(Row "):
                    action = base_action
                else:
                    action = f"{row_prefix}{base_action}".strip()
                detail = tail.strip()
            display_detail = detail or error_text
            parsed_row_no = int(row_match.group(1)) if row_match else None
            norm_error_text = _norm_error_text(error_text)
            norm_display_detail = _norm_error_text(display_detail)
            existing_bucket = existing_row_errors.get(parsed_row_no, set())
            if norm_error_text in existing_bucket or norm_display_detail in existing_bucket:
                continue

            detail_short, detail_long = _truncate_detail(display_detail)
            error_steps.append(
                {
                    "action": action,
                    "detail": display_detail,
                    "detail_short": detail_short,
                    "detail_long": detail_long,
                    "row_no": parsed_row_no,
                    "status": "error",
                    "error": error_text,
                }
            )
        if error_steps:
            step_runs = error_steps + step_runs
        error_messages = []
    if step_runs:
        step_runs = sorted(
            step_runs,
            key=lambda s: (s.get("row_no") is None, s.get("row_no") or 10**9),
        )
        row_has_error = {}
        for step in step_runs:
            row_no = step.get("row_no")
            if row_no is None:
                continue
            if step.get("status") == "error":
                row_has_error[row_no] = True
            else:
                row_has_error.setdefault(row_no, False)
        if row_has_error:
            step_runs = [
                step
                for step in step_runs
                if step.get("row_no") is None
                or not row_has_error.get(step.get("row_no"))
                or step.get("status") == "error"
            ]
    step_ok_count = sum(1 for step in step_runs if step.get("status") != "error")
    step_error_count = sum(1 for step in step_runs if step.get("status") == "error")
    rel_outputs = []
    for p in outputs:
        rel = os.path.relpath(p, out_dir)
        rel_outputs.append(rel.replace("\\", "/"))
    if request.method == "POST" and (request.form.get("action") == "run_cached") and current_run_id:
        run_dir = os.path.join(out_dir, current_run_id)
        run_outputs = []
        run_prefix = f"{current_run_id}/"
        for rel in rel_outputs:
            run_outputs.append(rel[len(run_prefix):] if rel.startswith(run_prefix) else rel)
        first_error = ""
        for step in step_runs:
            if step.get("status") == "error":
                first_error = str(step.get("error") or step.get("detail") or "").strip()
                if first_error:
                    break
        if not first_error:
            for msg in messages:
                if "ERROR" in (msg or ""):
                    first_error = str(msg).strip()
                    break
        status_text = "failed" if has_error else "completed"
        work_id, actor_label = _get_actor_info()
        _write_mapping_run_meta(
            run_dir,
            {
                "record_type": "mapping_run",
                "run_id": current_run_id,
                "mapping_file": validation_state.get("mapping_file") or last_mapping_file or "",
                "mapping_display_name": current_mapping_display_name or validation_state.get("mapping_display_name") or validation_state.get("mapping_file") or last_mapping_file or "",
                "status": status_text,
                "started_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "completed_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "reference_ok": bool(validation_state.get("reference_ok")),
                "extract_ok": bool(validation_state.get("extract_ok")),
                "outputs": run_outputs,
                "output_count": len(run_outputs),
                "zip_file": zip_file_name or "",
                "log_file": log_file_name or "",
                "error": first_error,
                "actor_work_id": work_id,
                "actor_label": actor_label,
            },
        )
    return render_template(
        "tasks/mapping.html",
        task_id=task_id,
        task=_load_task_context(task_id),
        messages=messages,
        outputs=rel_outputs,
        log_file=log_file,
        zip_file=zip_file,
        has_error=has_error,
        has_warning=has_warning,
        warning_confirm=warning_confirm,
        step_runs=step_runs,
        step_ok_count=step_ok_count,
        step_error_count=step_error_count,
        error_messages=error_messages,
        last_mapping_file=last_mapping_file,
        current_mapping_display_name=current_mapping_display_name or last_mapping_file,
        allow_check_extract=bool(
            last_mapping_file
            and validation_state.get("mapping_file") == last_mapping_file
            and validation_state.get("reference_ok")
        ),
        allow_direct_run=bool(
            last_mapping_file
            and validation_state.get("mapping_file") == last_mapping_file
            and validation_state.get("extract_ok")
            and request.method == "POST"
            and (request.form.get("action") == "check_extract")
            and not has_error
        ),
    )


@tasks_bp.get("/tasks/<task_id>/mapping/example", endpoint="task_download_mapping_example")
def task_download_mapping_example(task_id):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    if not os.path.isdir(tdir):
        abort(404)

    static_dir = current_app.static_folder or ""
    rel_sample = "samples/mapping_example.xlsx"
    sample_path = os.path.join(static_dir, rel_sample)
    checked_paths: list[str] = [sample_path]

    if os.path.isfile(sample_path):
        return send_from_directory(
            static_dir,
            rel_sample,
            as_attachment=True,
            download_name="mapping_example.xlsx",
        )

    project_root = Path(current_app.root_path).parent
    fallback_candidates = [
        project_root / "Mapping.xlsx",
        project_root / "mapping.xlsx",
        project_root / "MAPPING.xlsx",
    ]
    for candidate in fallback_candidates:
        checked_paths.append(str(candidate))
        if candidate.is_file():
            return send_file(
                str(candidate),
                as_attachment=True,
                download_name="mapping_example.xlsx",
            )

    current_app.logger.error(
        "Mapping example file not found. checked_paths=%s",
        checked_paths,
    )
    return (
        "找不到 Mapping 範例檔案。請確認 static/samples/mapping_example.xlsx 或專案根目錄 Mapping.xlsx 是否存在。",
        404,
    )

def _send_mapping_output_file(task_id: str, filename: str):
    safe_name = (filename or "").replace("\\", "/").strip("/")
    if not safe_name:
        abort(404)

    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    mapping_job_dir = os.path.join(tdir, "mapping_job")
    legacy_out_dir = os.path.join(current_app.config["OUTPUT_FOLDER"], task_id)

    for base_dir in (mapping_job_dir, legacy_out_dir):
        file_path = os.path.join(base_dir, safe_name)
        if os.path.isfile(file_path):
            return send_from_directory(base_dir, safe_name, as_attachment=True)
    abort(404)


def _write_mapping_run_meta(run_dir: str, payload: dict) -> None:
    try:
        os.makedirs(run_dir, exist_ok=True)
        meta_path = os.path.join(run_dir, "meta.json")
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
    except Exception:
        current_app.logger.exception("Failed to write mapping run meta")


@tasks_bp.get("/tasks/<task_id>/output/download", endpoint="task_download_output_query")
def task_download_output_query(task_id):
    return _send_mapping_output_file(task_id, request.args.get("filename", ""))


@tasks_bp.get("/tasks/<task_id>/output/<path:filename>", endpoint="task_download_output")
def task_download_output(task_id, filename):
    return _send_mapping_output_file(task_id, filename)

@tasks_bp.get("/", endpoint="tasks")
def tasks():
    task_list_all = list_tasks()
    pin_scope_key, _ = _get_actor_info()
    
    # Pagination
    page = request.args.get("page", 1, type=int)
    per_page = 10
    total_count = len(task_list_all)
    total_pages = (total_count + per_page - 1) // per_page
    start = (page - 1) * per_page
    task_list = task_list_all[start : start + per_page]
    
    pagination = {
        "page": page,
        "total_count": total_count,
        "total_pages": total_pages,
        "has_prev": page > 1,
        "has_next": page < total_pages
    }
    
    for t in task_list:
        meta = {
            "creator_work_id": t.get("creator_work_id", ""),
            "creator": t.get("creator", ""),
        }
        t["can_delete"] = _can_delete_task(meta)
    return render_template(
        "tasks/tasks.html",
        tasks=task_list,
        pagination=pagination,
        all_task_ids=[(t.get("id") or "").strip() for t in task_list_all if (t.get("id") or "").strip()],
        pin_scope_key=pin_scope_key or "anonymous",
        allowed_nas_roots=get_configured_nas_roots(),
    )

@tasks_bp.post("/tasks", endpoint="create_task")
def create_task():
    def _fail(message: str):
        flash(message, "danger")
        return redirect(url_for("tasks_bp.tasks"))

    nas_path = request.form.get("nas_path", "")
    try:
        nas_root_index = request.form.get("nas_root_index", "").strip()
        resolved_path = resolve_nas_path(
            nas_path,
            allow_recursive=current_app.config.get("NAS_ALLOW_RECURSIVE", True),
            root_index=nas_root_index or None,
        )
        if not os.path.isdir(resolved_path):
            return _fail("指定的 NAS 路徑不是資料夾")
        enforce_max_copy_size(resolved_path)
    except ValueError as exc:
        return _fail(str(exc))
    except FileNotFoundError as exc:
        return _fail(str(exc))
    task_name = request.form.get("task_name", "").strip() or "未命名任務"
    task_desc = request.form.get("task_desc", "").strip()
    if task_name_exists(task_name):
        return _fail("任務名稱已存在")
    tid = str(uuid.uuid4())[:8]
    tdir = os.path.join(current_app.config["TASK_FOLDER"], tid)
    files_dir = os.path.join(tdir, "files")
    os.makedirs(files_dir, exist_ok=True)
    src_dir = ensure_windows_long_path(resolved_path)
    dest_dir = ensure_windows_long_path(files_dir)
    try:
        shutil.copytree(src_dir, dest_dir, dirs_exist_ok=True)
    except PermissionError:
        shutil.rmtree(tdir, ignore_errors=True)
        return _fail("沒有足夠的權限讀取或複製指定路徑")
    except shutil.Error as exc:
        current_app.logger.exception("複製 NAS 目錄失敗")
        shutil.rmtree(tdir, ignore_errors=True)
        detail = ""
        if exc.args and isinstance(exc.args[0], list) and exc.args[0]:
            first_error = exc.args[0][0]
            if len(first_error) >= 3:
                detail = f"：{first_error[2]}"
        return _fail(f"複製 NAS 目錄時發生錯誤{detail or ''}，請稍後再試")
    except Exception:
        current_app.logger.exception("複製 NAS 目錄失敗")
        shutil.rmtree(tdir, ignore_errors=True)
        return _fail("複製 NAS 目錄時發生錯誤，請稍後再試")
    work_id, creator = _get_actor_info()
    display_nas_path = resolved_path
    if nas_root_index:
        try:
            roots = get_configured_nas_roots()
            idx = int(nas_root_index)
            if 0 <= idx < len(roots) and not os.path.isabs(nas_path):
                root = roots[idx]
                sep = "\\" if "\\" in root else "/"
                root_clean = re.sub(r"[\\/]+$", "", root)
                rel = re.sub(r"^[./\\\\]+", "", nas_path).replace("/", sep)
                display_nas_path = f"{root_clean}{sep}{rel}" if rel else root_clean
        except (ValueError, TypeError):
            pass

    created_at = datetime.now()
    meta_payload = {
        "name": task_name,
        "description": task_desc,
        "created": created_at.strftime("%Y-%m-%d %H:%M"),
        "nas_path": display_nas_path,
    }
    if creator:
        meta_payload["creator"] = creator
    if work_id:
        meta_payload["creator_work_id"] = work_id
    if creator:
        meta_payload["last_editor"] = creator
    if work_id:
        meta_payload["last_editor_work_id"] = work_id
    meta_payload["last_edited"] = created_at.strftime("%Y-%m-%d %H:%M")
    with open(os.path.join(tdir, "meta.json"), "w", encoding="utf-8") as meta:
        json.dump(
            meta_payload,
            meta,
            ensure_ascii=False,
            indent=2,
        )
    record_task_in_db(
        tid,
        name=task_name,
        description=task_desc,
        creator=creator or None,
        nas_path=display_nas_path or None,
        created_at=created_at,
    )
    record_audit(
        action="task_create",
        actor={"work_id": work_id, "label": creator},
        detail={"task_id": tid, "task_name": task_name, "nas_path": display_nas_path},
        task_id=tid,
    )
    return redirect(url_for("tasks_bp.tasks"))

@tasks_bp.post("/tasks/<task_id>/copy", endpoint="copy_task")
def copy_task(task_id):
    def _fail(message: str):
        flash(message, "danger")
        return redirect(url_for("tasks_bp.tasks"))

    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    if not os.path.isdir(tdir):
        return _fail("找不到任務資料夾")

    new_name = request.form.get("name", "").strip()
    if not new_name:
        return _fail("缺少任務名稱")
    if task_name_exists(new_name):
        return _fail("任務名稱已存在")

    meta_path = os.path.join(tdir, "meta.json")
    meta = {}
    if os.path.exists(meta_path):
        with open(meta_path, "r", encoding="utf-8") as f:
            meta = json.load(f)
    source_nas_path = (meta.get("nas_path", "") or "").strip()

    requested_nas_path = request.form.get("nas_path")
    requested_root_index = request.form.get("nas_root_index", "").strip()
    target_nas_path = source_nas_path
    if requested_nas_path is not None:
        raw_nas_path = requested_nas_path.strip()
        if not raw_nas_path:
            target_nas_path = ""
        elif raw_nas_path == source_nas_path:
            target_nas_path = source_nas_path
        else:
            try:
                if os.path.isabs(raw_nas_path):
                    resolved_path = os.path.abspath(raw_nas_path)
                    roots = get_configured_nas_roots()
                    if roots and not requested_root_index:
                        for idx, root in enumerate(roots):
                            root_abs = os.path.abspath(root)
                            try:
                                if os.path.commonpath([root_abs, resolved_path]) == root_abs:
                                    rel = os.path.relpath(resolved_path, root_abs).replace("\\", "/")
                                    raw_nas_path = "." if rel == "." else rel
                                    requested_root_index = str(idx)
                                    break
                            except ValueError:
                                continue
                    if roots:
                        allowed = False
                        for root in roots:
                            root_abs = os.path.abspath(root)
                            try:
                                if os.path.commonpath([root_abs, resolved_path]) == root_abs:
                                    allowed = True
                                    break
                            except ValueError:
                                continue
                        if not allowed:
                            return _fail("NAS 路徑不在允許的根目錄內。")
                    if not os.path.isdir(resolved_path):
                        return _fail("NAS 路徑不存在或不是資料夾。")
                    target_nas_path = resolved_path
                else:
                    resolved_path = resolve_nas_path(
                        raw_nas_path,
                        allow_recursive=current_app.config.get("NAS_ALLOW_RECURSIVE", True),
                        root_index=requested_root_index or None,
                    )
                    if not os.path.isdir(resolved_path):
                        return _fail("NAS 路徑不存在或不是資料夾。")
                    target_nas_path = resolved_path
                    if requested_root_index:
                        roots = get_configured_nas_roots()
                        try:
                            idx = int(requested_root_index)
                            if 0 <= idx < len(roots):
                                root_clean = roots[idx].rstrip("/\\")
                                sep = "\\" if "\\" in root_clean else "/"
                                rel = re.sub(r"^[./\\]+", "", raw_nas_path).replace("/", sep)
                                target_nas_path = f"{root_clean}{sep}{rel}" if rel else root_clean
                        except ValueError:
                            pass
            except ValueError as exc:
                return _fail(str(exc))
            except FileNotFoundError as exc:
                return _fail(str(exc))

    created_at = datetime.now()
    work_id, creator = _get_actor_info()

    new_id = str(uuid.uuid4())[:8]
    new_dir = os.path.join(current_app.config["TASK_FOLDER"], new_id)
    os.makedirs(new_dir, exist_ok=False)
    try:
        for subdir in ("files", "flows"):
            src = os.path.join(tdir, subdir)
            dest = os.path.join(new_dir, subdir)
            if os.path.isdir(src):
                shutil.copytree(ensure_windows_long_path(src), ensure_windows_long_path(dest))
            elif subdir == "files":
                os.makedirs(dest, exist_ok=True)
    except Exception:
        current_app.logger.exception("複製任務資料夾失敗")
        shutil.rmtree(new_dir, ignore_errors=True)
        return _fail("複製任務資料夾失敗，請稍後再試")

    new_meta = {
        "name": new_name,
        "description": meta.get("description", ""),
        "nas_path": target_nas_path,
        "created": created_at.strftime("%Y-%m-%d %H:%M"),
        "last_edited": created_at.strftime("%Y-%m-%d %H:%M"),
    }
    if creator:
        new_meta["creator"] = creator
        new_meta["last_editor"] = creator
    if work_id:
        new_meta["creator_work_id"] = work_id
        new_meta["last_editor_work_id"] = work_id
    with open(os.path.join(new_dir, "meta.json"), "w", encoding="utf-8") as f:
        json.dump(new_meta, f, ensure_ascii=False, indent=2)

    record_task_in_db(
        new_id,
        name=new_name,
        description=new_meta.get("description") or None,
        creator=creator or None,
        nas_path=new_meta.get("nas_path") or None,
        created_at=created_at,
    )
    record_audit(
        action="task_copy",
        actor={"work_id": work_id, "label": creator},
        detail={
            "task_id": new_id,
            "task_name": new_name,
            "nas_path": new_meta.get("nas_path"),
            "source_nas_path": source_nas_path,
            "source_task_id": task_id
        },
        task_id=new_id,
    )
    flash("已複製任務", "success")
    return redirect(url_for("tasks_bp.tasks"))

@tasks_bp.post("/tasks/<task_id>/delete", endpoint="delete_task")
def delete_task(task_id):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    meta_path = os.path.join(tdir, "meta.json")
    meta = {}
    if os.path.exists(meta_path):
        with open(meta_path, "r", encoding="utf-8") as f:
            meta = json.load(f)
    if not _can_delete_task(meta):
        abort(403)
    work_id, label = _get_actor_info()
    record_audit(
        action="task_delete",
        actor={"work_id": work_id, "label": label},
        detail={"task_id": task_id, "task_name": meta.get("name", "")},
        task_id=task_id,
    )
    if os.path.isdir(tdir):
        import shutil
        shutil.rmtree(tdir)
    delete_task_record(task_id)
    return redirect(url_for("tasks_bp.tasks"))

@tasks_bp.get("/tasks/<task_id>/nas-diff", endpoint="task_nas_diff")
def task_nas_diff(task_id):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    files_dir = os.path.join(tdir, "files")
    meta_path = os.path.join(tdir, "meta.json")
    if not os.path.isdir(files_dir) or not os.path.exists(meta_path):
        return jsonify({"ok": False, "error": "Task not found"}), 404

    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)
    nas_path = (meta.get("nas_path") or "").strip()
    if not nas_path:
        return jsonify({"ok": True, "diff": None, "message": "尚未設定 NAS 路徑"}), 200
    if not os.path.isdir(nas_path):
        return jsonify({"ok": True, "diff": None, "message": "NAS 路徑不存在或不是資料夾"}), 200

    try:
        def _list_empty_dirs(base: str) -> set[str]:
            empties: set[str] = set()
            for root, dirs, files in os.walk(base):
                if dirs or files:
                    continue
                rel = os.path.relpath(root, base)
                if rel == ".":
                    continue
                empties.add(rel.replace("\\", "/") + "/")
            return empties

        task_files_map = {p.replace("\\", "/"): os.path.join(files_dir, p) for p in list_files(files_dir)}
        nas_files_map = {p.replace("\\", "/"): os.path.join(nas_path, p) for p in list_files(nas_path)}
        task_entries = set(task_files_map.keys()) | _list_empty_dirs(files_dir)
        nas_entries = set(nas_files_map.keys()) | _list_empty_dirs(nas_path)

        added = sorted(nas_entries - task_entries)
        removed = sorted(task_entries - nas_entries)
        updated = []

        # Check for modified files among common files
        common_files = set(task_files_map.keys()) & set(nas_files_map.keys())
        for rel in common_files:
            try:
                t_stat = os.stat(task_files_map[rel])
                n_stat = os.stat(nas_files_map[rel])
                # 同步邏輯使用：size 不同或 NAS 檔案較新
                if n_stat.st_size != t_stat.st_size or int(n_stat.st_mtime) > int(t_stat.st_mtime):
                    updated.append(rel)
            except Exception:
                continue
        updated.sort()

        if not added and not removed and not updated:
            return jsonify({"ok": True, "diff": None, "message": "未偵測到變更"}), 200

        limit = 5
        diff = {
            "added": added[:limit],
            "removed": removed[:limit],
            "updated": updated[:limit],
            "added_count": len(added),
            "removed_count": len(removed),
            "updated_count": len(updated),
            "limit": limit,
        }
        return jsonify({"ok": True, "diff": diff}), 200
    except Exception:
        current_app.logger.exception("Failed to compare NAS files")
        return jsonify({"ok": False, "error": "Failed to compare NAS files"}), 500


@tasks_bp.post("/tasks/<task_id>/sync-nas", endpoint="sync_task_nas")
def sync_task_nas(task_id):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    files_dir = os.path.join(tdir, "files")
    meta_path = os.path.join(tdir, "meta.json")
    if not os.path.exists(meta_path):
        abort(404)

    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)
    nas_path = (meta.get("nas_path") or "").strip()
    if not nas_path:
        flash("尚未設定 NAS 路徑，無法更新。", "warning")
        return redirect(url_for("tasks_bp.task_detail", task_id=task_id))

    abs_path = os.path.abspath(nas_path)
    roots = get_configured_nas_roots()
    if roots:
        allowed = False
        for root in roots:
            root_abs = os.path.abspath(root)
            try:
                if os.path.commonpath([root_abs, abs_path]) == root_abs:
                    allowed = True
                    break
            except ValueError:
                continue
        if not allowed:
            flash("NAS 路徑不在允許的根目錄內。", "danger")
            return redirect(url_for("tasks_bp.task_detail", task_id=task_id))

    if not os.path.isdir(abs_path):
        flash("NAS 路徑不存在或不是資料夾。", "danger")
        return redirect(url_for("tasks_bp.task_detail", task_id=task_id))

    try:
        enforce_max_copy_size(abs_path)
    except ValueError as exc:
        flash(str(exc), "danger")
        return redirect(url_for("tasks_bp.task_detail", task_id=task_id))

    try:
        src_dir = ensure_windows_long_path(abs_path)
        dst_dir = ensure_windows_long_path(files_dir)
        os.makedirs(dst_dir, exist_ok=True)
        copied = 0
        updated = 0
        deleted = 0
        created_dirs = 0
        deleted_dirs = 0
        for root, dirs, files in os.walk(src_dir):
            rel = os.path.relpath(root, src_dir)
            dest_root = dst_dir if rel == "." else os.path.join(dst_dir, rel)
            if not os.path.exists(dest_root):
                os.makedirs(dest_root, exist_ok=True)
                if rel != ".":
                    created_dirs += 1
            else:
                os.makedirs(dest_root, exist_ok=True)
            for fname in files:
                src_file = os.path.join(root, fname)
                dst_file = os.path.join(dest_root, fname)
                try:
                    if not os.path.exists(dst_file):
                        shutil.copy2(src_file, dst_file)
                        copied += 1
                        continue
                    src_stat = os.stat(src_file)
                    dst_stat = os.stat(dst_file)
                    if src_stat.st_size != dst_stat.st_size or int(src_stat.st_mtime) > int(dst_stat.st_mtime):
                        shutil.copy2(src_file, dst_file)
                        updated += 1
                except FileNotFoundError:
                    continue
        for root, dirs, files in os.walk(dst_dir, topdown=False):
            rel = os.path.relpath(root, dst_dir)
            src_root = src_dir if rel == "." else os.path.join(src_dir, rel)
            for fname in files:
                dst_file = os.path.join(root, fname)
                src_file = os.path.join(src_root, fname)
                if not os.path.exists(src_file):
                    try:
                        os.remove(dst_file)
                        deleted += 1
                    except FileNotFoundError:
                        continue
            if rel != "." and not os.path.exists(src_root):
                try:
                    shutil.rmtree(root)
                    deleted_dirs += 1
                except FileNotFoundError:
                    pass
        _apply_last_edit(meta)
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)
        total_added = copied + created_dirs
        total_deleted = deleted + deleted_dirs
        flash(f"已更新 NAS 內容（新增 {total_added}、更新 {updated}、刪除 {total_deleted}）。", "success")
        work_id, label = _get_actor_info()
        record_audit(
            action="nas_sync",
            actor={"work_id": work_id, "label": label},
            detail={
                "task_id": task_id,
                "nas_path": nas_path,
                "copied": copied,
                "updated": updated,
                "deleted": deleted,
                "created_dirs": created_dirs,
                "deleted_dirs": deleted_dirs,
            },
            task_id=task_id,
        )
    except PermissionError:
        flash("沒有足夠的權限讀取或複製指定路徑。", "danger")
    except Exception:
        current_app.logger.exception("更新 NAS 文件失敗")
        flash("更新 NAS 文件失敗，請稍後再試。", "danger")

    return redirect(url_for("tasks_bp.task_detail", task_id=task_id))


@tasks_bp.post("/tasks/<task_id>/rename", endpoint="rename_task")
def rename_task(task_id):
    new_name = request.form.get("name", "").strip()
    if not new_name:
        return "缺少名稱", 400
    if task_name_exists(new_name, exclude_id=task_id):
        return "任務名稱已存在", 400
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    if not os.path.isdir(tdir):
        abort(404)
    meta_path = os.path.join(tdir, "meta.json")
    meta = {}
    if os.path.exists(meta_path):
        with open(meta_path, "r", encoding="utf-8") as f:
            meta = json.load(f)
    meta["name"] = new_name
    if "created" not in meta:
        meta["created"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)
    record_task_in_db(task_id, name=new_name)
    work_id, label = _get_actor_info()
    record_audit(
        action="task_rename",
        actor={"work_id": work_id, "label": label},
        detail={"task_id": task_id, "name": new_name},
        task_id=task_id,
    )
    return redirect(url_for("tasks_bp.tasks"))

@tasks_bp.post("/tasks/<task_id>/description", endpoint="update_task_description")
def update_task_description(task_id):
    new_desc = request.form.get("description", "").strip()
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    if not os.path.isdir(tdir):
        abort(404)
    meta_path = os.path.join(tdir, "meta.json")
    meta = {}
    if os.path.exists(meta_path):
        with open(meta_path, "r", encoding="utf-8") as f:
            meta = json.load(f)
    meta["description"] = new_desc
    if "name" not in meta:
        meta["name"] = task_id
    if "created" not in meta:
        meta["created"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)
    record_task_in_db(task_id, description=new_desc)
    work_id, label = _get_actor_info()
    record_audit(
        action="task_update_description",
        actor={"work_id": work_id, "label": label},
        detail={"task_id": task_id, "description": new_desc},
        task_id=task_id,
    )
    return redirect(url_for("tasks_bp.tasks"))

@tasks_bp.get("/tasks/<task_id>", endpoint="task_detail")
def task_detail(task_id):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    files_dir = os.path.join(tdir, "files")
    if not os.path.isdir(files_dir):
        abort(404)
    meta_path = os.path.join(tdir, "meta.json")
    name = task_id
    description = ""
    creator = ""
    nas_path = ""
    if os.path.exists(meta_path):
        with open(meta_path, "r", encoding="utf-8") as f:
            meta = json.load(f)
            name = meta.get("name", task_id)
            description = meta.get("description", "")
            creator = meta.get("creator", "") or ""
            nas_path = meta.get("nas_path", "") or ""
    nas_diff = None
    if nas_path and os.path.isdir(nas_path):
        try:
            task_files = {p.replace("\\", "/") for p in list_files(files_dir)}
            nas_files = {p.replace("\\", "/") for p in list_files(nas_path)}
            added = sorted(nas_files - task_files)
            removed = sorted(task_files - nas_files)
            if added or removed:
                limit = 5
                nas_diff = {
                    "added": added[:limit],
                    "removed": removed[:limit],
                    "added_count": len(added),
                    "removed_count": len(removed),
                    "limit": limit,
                }
        except Exception:
            current_app.logger.exception("Failed to compare NAS files")
    tree = build_file_tree(files_dir)
    return render_template(
        "tasks/task_detail.html",
        task={"id": task_id, "name": name, "description": description, "creator": creator, "nas_path": nas_path},
        nas_diff=nas_diff,
        files_tree=tree,
    )

@tasks_bp.post("/tasks/<task_id>/templates/parse", endpoint="parse_template_doc")
def parse_template_doc(task_id):
    """Upload or parse an existing template docx and return paragraph metadata."""
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    files_dir = os.path.join(tdir, "files")
    if not os.path.isdir(files_dir):
        abort(404)

    upload = request.files.get("template_file")
    template_rel = ""
    existing = request.form.get("template_path", "").strip()

    if upload and upload.filename:
        if not upload.filename.lower().endswith(".docx"):
            return jsonify({"ok": False, "error": "僅支援 .docx 模板"}), 400
        safe_name = deduplicate_name(files_dir, secure_filename(upload.filename))
        save_path = os.path.join(files_dir, safe_name)
        upload.save(save_path)
        template_rel = safe_name
    elif existing:
        normalized = os.path.normpath(existing)
        if normalized.startswith("..") or os.path.isabs(normalized):
            return jsonify({"ok": False, "error": "無效的檔案路徑"}), 400
        template_rel = normalized
    else:
        return jsonify({"ok": False, "error": "請選擇或上傳模板檔案"}), 400

    template_abs = os.path.join(files_dir, template_rel)
    if not os.path.isfile(template_abs):
        return jsonify({"ok": False, "error": "找不到模板檔案"}), 404

    try:
        paragraphs = parse_template_paragraphs(template_abs)
    except Exception as e:
        current_app.logger.exception("Failed to parse template docx")
        return jsonify({"ok": False, "error": f"解析模板失敗: {e}"}), 400

    return jsonify(
        {
            "ok": True,
            "template_file": template_rel.replace("\\", "/"),
            "paragraphs": paragraphs,
        }
    )

@tasks_bp.get("/tasks/<task_id>/result/<job_id>", endpoint="task_result")
def task_result(task_id, job_id):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    docx_path = os.path.join(job_dir, "result.docx")
    if not os.path.exists(docx_path):
        return "Job not found or failed.", 404
    log_json_path = os.path.join(job_dir, "log.json")
    log_entries = []
    overall_status = "ok"
    if os.path.exists(log_json_path):
        with open(log_json_path, "r", encoding="utf-8") as f:
            log_entries = json.load(f)
        if any(e.get("status") == "error" for e in log_entries):
            overall_status = "error"
    return render_template(
        "tasks/run.html",
        task=_load_task_context(task_id),
        job_id=job_id,
        docx_path=url_for("tasks_bp.task_download", task_id=task_id, job_id=job_id, kind="docx"),
        log_path=url_for("tasks_bp.task_download", task_id=task_id, job_id=job_id, kind="log"),
        translate_path=url_for("tasks_bp.task_translate", task_id=task_id, job_id=job_id),
        compare_path=url_for("tasks_bp.task_compare", task_id=task_id, job_id=job_id),
        back_link=url_for("flows_bp.flow_builder", task_id=task_id),
        log_entries=log_entries,
        overall_status=overall_status,
    )

@tasks_bp.get("/tasks/<task_id>/translate/<job_id>", endpoint="task_translate")
def task_translate(task_id, job_id):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    src = os.path.join(job_dir, "result.docx")
    if not os.path.exists(src):
        abort(404)
    out_docx = os.path.join(job_dir, "translated.docx")
    if not os.path.exists(out_docx):
        tmp_md = os.path.join(job_dir, "translated.md")
        translate_file(src, tmp_md)
        import docx
        doc = docx.Document()
        with open(tmp_md, "r", encoding="utf-8") as f:
            for line in f.read().splitlines():
                doc.add_paragraph(line)
        doc.save(out_docx)
    return send_file(
        out_docx,
        as_attachment=True,
        download_name=f"translated_{job_id}.docx",
    )

@tasks_bp.get("/tasks/<task_id>/compare/<job_id>", endpoint="task_compare")
def task_compare(task_id, job_id):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    docx_path = os.path.join(job_dir, "result.docx")
    log_path = os.path.join(job_dir, "log.json")
    if not os.path.exists(docx_path) or not os.path.exists(log_path):
        abort(404)

    with open(log_path, "r", encoding="utf-8") as f:
        entries = json.load(f)
    titles_to_hide = collect_titles_to_hide(entries)
    preview_messages = []
    source_lookup = _build_provenance_source_lookup(entries)
    preview_docx_path = docx_path
    preview_docx_rel, preview_docx_error = _ensure_provenance_preview_docx(
        docx_path,
        log_path,
        job_dir,
        source_lookup,
    )
    if preview_docx_error:
        preview_messages.append(f"來源標記預覽建立失敗: {preview_docx_error}")
    elif preview_docx_rel:
        preview_docx_path = os.path.join(job_dir, preview_docx_rel)

    result_pdf_rel, result_pdf_error = _ensure_pdf_preview(preview_docx_path, job_dir, "preview_pdf")
    if result_pdf_error:
        preview_messages.append(f"結果文件預覽失敗: {result_pdf_error}")
    result_html_rel, result_html_error = _ensure_html_preview(
        preview_docx_path,
        job_dir,
        "preview_html",
        "provenance_preview",
    )
    if result_html_error:
        preview_messages.append(f"HTML 預覽建立失敗: {result_html_error}")

    chapter_sources = {}
    source_urls = {}
    converted_docx = {}
    current = None
    for entry in entries:
        stype = entry.get("type")
        params = entry.get("params", {})
        if stype == "insert_roman_heading":
            current = params.get("text", "")
            chapter_sources.setdefault(current, [])
        elif stype == "extract_pdf_chapter_to_table":
            pdf_dir = os.path.join(job_dir, "pdfs_extracted")
            pdfs = []
            if os.path.isdir(pdf_dir):
                for fn in sorted(os.listdir(pdf_dir)):
                    if fn.lower().endswith(".pdf"):
                        pdfs.append(fn)
                        rel = os.path.join("pdfs_extracted", fn)
                        source_urls[fn] = url_for("tasks_bp.task_view_file", task_id=task_id, job_id=job_id, filename=rel
                        )
            chapter_sources.setdefault(current or "未分類", []).extend(pdfs)
        elif stype == "extract_word_chapter":
            infile = params.get("input_file", "")
            base = os.path.basename(infile)
            source_label = _trace_source_label(entry)
            sec_start = params.get("target_chapter_section", "")
            sec_end = params.get("explicit_end_number", "")
            sec = f"{sec_start}-{sec_end}" if sec_start and sec_end else sec_start
            use_title = str(params.get("use_chapter_title", params.get("target_title", ""))).lower() in ["1", "true", "yes", "on"]
            title = params.get("target_chapter_title") or params.get("target_title_section", "")
            info = source_label
            if sec:
                info += f" 章節 {sec}"
            if title:
                info += f" 標題 {title}"
            chapter_sources.setdefault(current or "未分類", []).append(info)
            source_key = os.path.abspath(infile) if infile else ""
            if source_key and source_key not in converted_docx and os.path.exists(infile):
                pdf_rel, pdf_error = _ensure_pdf_preview(infile, job_dir, "source_pdf")
                if pdf_rel:
                    converted_docx[source_key] = pdf_rel
                elif pdf_error:
                    preview_messages.append(f"{base} 預覽失敗: {pdf_error}")
            if source_key in converted_docx:
                source_urls[info] = url_for(
                    "tasks_bp.task_view_file",
                    task_id=task_id,
                    job_id=job_id,
                    filename=converted_docx[source_key],
                )
                source_urls.setdefault(
                    source_label,
                    url_for(
                        "tasks_bp.task_view_file",
                        task_id=task_id,
                        job_id=job_id,
                        filename=converted_docx[source_key],
                    ),
                )
        elif stype == "extract_word_all_content":
            infile = params.get("input_file", "")
            base = os.path.basename(infile)
            source_label = _trace_source_label(entry)
            chapter_sources.setdefault(current or "未分類", []).append(source_label)
            source_key = os.path.abspath(infile) if infile else ""
            if source_key and source_key not in converted_docx and os.path.exists(infile):
                pdf_rel, pdf_error = _ensure_pdf_preview(infile, job_dir, "source_pdf")
                if pdf_rel:
                    converted_docx[source_key] = pdf_rel
                elif pdf_error:
                    preview_messages.append(f"{base} 預覽失敗: {pdf_error}")
            if source_key in converted_docx:
                source_urls[source_label] = url_for(
                    "tasks_bp.task_view_file",
                    task_id=task_id,
                    job_id=job_id,
                    filename=converted_docx[source_key],
                )
        elif stype == "extract_pdf_pages_as_images":
            infile = params.get("input_file", "")
            base = os.path.basename(infile)
            source_label = _trace_source_label(entry)
            chapter_sources.setdefault(current or "未分類", []).append(source_label)
            pdf_rel, pdf_error = _ensure_pdf_preview(infile, job_dir, "source_pdf")
            if pdf_rel:
                source_urls.setdefault(
                    source_label,
                    url_for(
                        "tasks_bp.task_view_file",
                        task_id=task_id,
                        job_id=job_id,
                        filename=pdf_rel,
                    ),
                )
            elif pdf_error:
                preview_messages.append(f"{base} 預覽失敗: {pdf_error}")
        elif stype in {"extract_specific_figure_from_word", "extract_specific_table_from_word"}:
            infile = params.get("input_file", "")
            base = os.path.basename(infile)
            source_label = _trace_source_label(entry)
            info = _build_compare_source_label(entry)
            chapter_sources.setdefault(current or "未分類", []).append(info)
            source_key = os.path.abspath(infile) if infile else ""
            if source_key and source_key not in converted_docx and os.path.exists(infile):
                pdf_rel, pdf_error = _ensure_pdf_preview(infile, job_dir, "source_pdf")
                if pdf_rel:
                    converted_docx[source_key] = pdf_rel
                elif pdf_error:
                    preview_messages.append(f"{base} 預覽失敗: {pdf_error}")
            if source_key in converted_docx:
                source_url = url_for(
                    "tasks_bp.task_view_file",
                    task_id=task_id,
                    job_id=job_id,
                    filename=converted_docx[source_key],
                )
                source_urls[info] = source_url
                source_urls.setdefault(source_label, source_url)

    chapters = list(chapter_sources.keys())
    provenance_trace = _build_provenance_trace(job_dir, docx_path, log_path, entries, titles_to_hide)
    if provenance_trace:
        paragraph_trace, object_trace_candidates = provenance_trace
    else:
        paragraph_trace = _build_paragraph_trace(job_dir, docx_path, log_path, entries, titles_to_hide)
        object_trace_candidates = _build_object_trace_candidates(entries, titles_to_hide)
    result_pdf_abs = os.path.join(job_dir, result_pdf_rel) if result_pdf_rel else ""
    paragraph_trace, page_source_map = _build_page_source_map(
        job_dir,
        result_pdf_abs,
        paragraph_trace,
        object_trace_candidates,
        source_lookup,
    )
    return render_template(
        "tasks/compare.html",
        task=_load_task_context(task_id),
        preview_url=url_for(
            "tasks_bp.task_view_file",
            task_id=task_id,
            job_id=job_id,
            filename=result_pdf_rel,
        ) if result_pdf_rel else "",
        html_preview_url=url_for(
            "tasks_bp.task_view_file",
            task_id=task_id,
            job_id=job_id,
            filename=result_html_rel,
        ) if result_html_rel else "",
        chapters=chapters,
        chapter_sources=chapter_sources,
        source_urls=source_urls,
        titles_to_hide=titles_to_hide,
        paragraph_trace=paragraph_trace,
        page_source_map=page_source_map,
        preview_messages=list(dict.fromkeys(preview_messages)),
        back_link=url_for("tasks_bp.task_result", task_id=task_id, job_id=job_id),
        download_url=url_for("tasks_bp.task_download", task_id=task_id, job_id=job_id, kind="docx"),
    )

@tasks_bp.post("/tasks/<task_id>/compare/<job_id>/save", endpoint="task_compare_save")
def task_compare_save(task_id, job_id):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    titles_to_hide = load_titles_to_hide_from_log(job_dir)
    html_content = request.form.get("html")
    if not html_content:
        data = request.get_json(silent=True) or {}
        html_content = data.get("html", "")
    if not html_content:
        return "缺少內容", 400
    html_content = clean_compare_html_content(html_content)
    save_compare_output(job_dir, html_content, titles_to_hide)
    return "OK"

@tasks_bp.post("/tasks/<task_id>/compare/<job_id>/save-as", endpoint="task_compare_save_as")
def task_compare_save_as(task_id, job_id):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    titles_to_hide = load_titles_to_hide_from_log(job_dir)
    payload = request.get_json(silent=True) or {}
    html_content = payload.get("html")
    name = payload.get("name") or ""
    if not html_content:
        html_content = request.form.get("html")
        name = request.form.get("name") or name
    if not html_content:
        return jsonify({"error": "缺少內容"}), 400
    version_name = (name or "").strip()
    if not version_name:
        return jsonify({"error": "缺少版本名稱"}), 400
    html_content = clean_compare_html_content(html_content)
    versions_dir = os.path.join(job_dir, "versions")
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    unique_suffix = uuid.uuid4().hex[:6]
    version_id = f"{timestamp}_{unique_suffix}"
    slug = sanitize_version_slug(version_name)
    base_name = f"{version_id}_{slug}" if slug else version_id
    save_compare_output(
        job_dir,
        html_content,
        titles_to_hide,
        base_name=base_name,
        subdir="versions",
    )
    metadata = load_version_metadata(versions_dir)
    versions = metadata.get("versions", [])
    versions = [v for v in versions if v.get("id") != version_id]
    created_ts = datetime.now()
    versions.append(
        {
            "id": version_id,
            "name": version_name,
            "slug": slug,
            "base_name": base_name,
            "created_at": created_ts.isoformat(timespec="seconds"),
        }
    )
    versions.sort(key=lambda v: v.get("created_at", ""), reverse=True)
    metadata["versions"] = versions
    save_version_metadata(versions_dir, metadata)
    version_payload = {
        "id": version_id,
        "name": version_name,
        "created_at_display": created_ts.strftime("%Y-%m-%d %H:%M:%S"),
        "html_url": url_for("tasks_bp.task_view_file",
            task_id=task_id,
            job_id=job_id,
            filename=f"versions/{base_name}.html",
        ),
        "docx_url": url_for("tasks_bp.task_download_version",
            task_id=task_id,
            job_id=job_id,
            version_id=version_id,
        ),
        "restore_url": url_for("tasks_bp.task_compare_restore_version",
            task_id=task_id,
            job_id=job_id,
            version_id=version_id,
        ),
        "delete_url": url_for("tasks_bp.task_compare_delete_version",
            task_id=task_id,
            job_id=job_id,
            version_id=version_id,
        ),
    }
    return jsonify({"status": "ok", "version": version_payload})

@tasks_bp.get("/tasks/<task_id>/view/<job_id>/<path:filename>", endpoint="task_view_file")
def task_view_file(task_id, job_id, filename):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    safe_filename = filename.replace("\\", "/")
    file_path = os.path.join(job_dir, safe_filename)
    if not os.path.isfile(file_path):
        abort(404)
    return send_from_directory(job_dir, safe_filename)

@tasks_bp.post("/tasks/<task_id>/compare/<job_id>/restore/<version_id>", endpoint="task_compare_restore_version")
def task_compare_restore_version(task_id, job_id, version_id):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    versions_dir = os.path.join(job_dir, "versions")
    metadata = load_version_metadata(versions_dir)
    versions = metadata.get("versions", [])
    version = next((v for v in versions if v.get("id") == version_id), None)
    if not version:
        return jsonify({"error": "找不到指定版本"}), 404
    base_name = version.get("base_name")
    if not base_name:
        return jsonify({"error": "版本資料不完整"}), 404
    html_src = os.path.join(versions_dir, f"{base_name}.html")
    docx_src = os.path.join(versions_dir, f"{base_name}.docx")
    if not os.path.exists(html_src) or not os.path.exists(docx_src):
        return jsonify({"error": "版本檔案不存在"}), 404
    shutil.copyfile(html_src, os.path.join(job_dir, "result.html"))
    shutil.copyfile(docx_src, os.path.join(job_dir, "result.docx"))
    return jsonify({"status": "ok"})

@tasks_bp.post("/tasks/<task_id>/compare/<job_id>/delete/<version_id>", endpoint="task_compare_delete_version")
def task_compare_delete_version(task_id, job_id, version_id):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    versions_dir = os.path.join(job_dir, "versions")
    metadata = load_version_metadata(versions_dir)
    versions = metadata.get("versions", [])
    version = next((v for v in versions if v.get("id") == version_id), None)
    if not version:
        return jsonify({"error": "找不到指定版本"}), 404
    metadata["versions"] = [v for v in versions if v.get("id") != version_id]
    save_version_metadata(versions_dir, metadata)
    base_name = version.get("base_name")
    if base_name:
        for ext in ("html", "docx"):
            path = os.path.join(versions_dir, f"{base_name}.{ext}")
            try:
                if os.path.exists(path):
                    os.remove(path)
            except OSError:
                pass
    return jsonify({"status": "ok"})

@tasks_bp.get("/tasks/<task_id>/download/<job_id>/version/<version_id>", endpoint="task_download_version")
def task_download_version(task_id, job_id, version_id):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    versions_dir = os.path.join(job_dir, "versions")
    metadata = load_version_metadata(versions_dir)
    versions = metadata.get("versions", [])
    version = next((v for v in versions if v.get("id") == version_id), None)
    if not version:
        abort(404)
    base_name = version.get("base_name")
    if not base_name:
        abort(404)
    docx_src = os.path.join(versions_dir, f"{base_name}.docx")
    if not os.path.exists(docx_src):
        abort(404)
    slug = version.get("slug") or version_id
    download_name = f"{slug}_{version_id}.docx"
    return send_file(docx_src, as_attachment=True, download_name=download_name)

@tasks_bp.get("/tasks/<task_id>/download/<job_id>/<kind>", endpoint="task_download")
def task_download(task_id, job_id, kind):
    tdir = os.path.join(current_app.config["TASK_FOLDER"], task_id)
    job_dir = os.path.join(tdir, "jobs", job_id)
    if kind == "docx":
        result_path = os.path.join(job_dir, "result.docx")
        if not os.path.exists(result_path):
            abort(404)
        titles_to_remove = []
        log_path = os.path.join(job_dir, "log.json")
        if os.path.exists(log_path):
            try:
                with open(log_path, "r", encoding="utf-8") as f:
                    entries = json.load(f)
                titles_to_remove = collect_titles_to_hide(entries)
            except Exception:
                titles_to_remove = []

        download_path = os.path.join(job_dir, "result_download.docx")
        shutil.copyfile(result_path, download_path)
        if titles_to_remove:
            remove_paragraphs_with_text(download_path, titles_to_remove)
        if not SKIP_DOCX_CLEANUP:
            remove_hidden_runs(download_path)
        download_name = f"result_{job_id}.docx"
        meta_path = os.path.join(job_dir, "meta.json")
        if os.path.exists(meta_path):
            try:
                with open(meta_path, "r", encoding="utf-8") as f:
                    meta = json.load(f)
                if isinstance(meta, dict):
                    candidate_name, candidate_error = normalize_docx_output_filename(
                        meta.get("output_filename"),
                        default="",
                    )
                    if not candidate_error and candidate_name:
                        download_name = candidate_name
            except Exception:
                pass
        return send_file(
            download_path,
            as_attachment=True,
            download_name=download_name,
        )
    elif kind == "log":
        return send_file(
            os.path.join(job_dir, "log.json"),
            as_attachment=True,
            download_name=f"log_{job_id}.json",
        )
    abort(404)
