import os
import re
import queue
import tempfile
from typing import Iterable, Optional
from uuid import uuid4
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from spire.doc import *
from spire.doc.common import *
from modules.chapter_section_parse import (
    parse_chapter_section_expression as _parse_chapter_section_expression,
)
from modules.extract_word_all_content import extract_body_with_options
from modules.extract_word_chapter import extract_section_docx_xml
from modules.extract_specific_figure_xml import extract_specific_figure_from_word_xml
from modules.extract_specific_table_xml import extract_specific_table_from_word_xml


def set_run_font_eastasia(run, eastasia_name: str):
    """補設定東亞字型，避免中文顯示為預設字型。"""
    if eastasia_name:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), eastasia_name)


def _build_output_docx_path(input_file: str, suffix: str) -> str:
    base_dir = os.path.dirname(os.path.abspath(input_file))
    base_name = os.path.splitext(os.path.basename(input_file))[0]
    safe_suffix = re.sub(r"[^A-Za-z0-9_.-]+", "_", suffix) or "extract"
    candidate = os.path.join(base_dir, f"{base_name}_{safe_suffix}.docx")
    if os.path.exists(candidate):
        candidate = os.path.join(
            tempfile.gettempdir(),
            f"{base_name}_{safe_suffix}_{uuid4().hex[:8]}.docx",
        )
    return candidate


def _append_docx_to_section(docx_path: str, output_doc: Document, section=None):
    target_section = section or output_doc.AddSection()
    temp_doc = Document()
    temp_doc.LoadFromFile(docx_path)
    try:
        for s_idx in range(temp_doc.Sections.Count):
            src_section = temp_doc.Sections.get_Item(s_idx)
            body = src_section.Body
            for i in range(body.ChildObjects.Count):
                cloned = body.ChildObjects.get_Item(i).Clone()
                target_section.Body.ChildObjects.Add(cloned)
    finally:
        temp_doc.Close()
    return target_section


def _read_first_paragraph_text(docx_path: str) -> str:
    try:
        doc = DocxDocument(docx_path)
    except Exception:
        return ""
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            return text
    return ""


def _get_paragraph_text(paragraph: Paragraph) -> str:
    text = paragraph.ListText + " " if paragraph.ListText else ""
    for j in range(paragraph.ChildObjects.Count):
        sub = paragraph.ChildObjects.get_Item(j)
        if sub.DocumentObjectType == DocumentObjectType.TextRange:
            text += sub.Text or ""
    return text


def _get_paragraph_text_stripped(paragraph: Paragraph) -> str:
    return _get_paragraph_text(paragraph).strip()


def _hide_titles_in_section(section, titles: list[str], start_index: int = 0):
    if not section or not titles:
        return
    targets = {_normalize_text(t) for t in titles if _normalize_text(t)}
    if not targets:
        return
    child_objects = section.Body.ChildObjects
    for idx in range(start_index, child_objects.Count):
        obj = child_objects.get_Item(idx)
        if isinstance(obj, Paragraph):
            para_text = _normalize_text(_get_paragraph_text(obj))
            if para_text in targets:
                for j in range(obj.ChildObjects.Count):
                    sub = obj.ChildObjects.get_Item(j)
                    if sub.DocumentObjectType == DocumentObjectType.TextRange:
                        sub.CharacterFormat.Hidden = True


def extract_pdf_chapter_to_table(pdf_folder_path: str, target_section: str, output_doc=None, section=None):
    upper_ratio = 0.1
    lower_ratio = 0.9

    stop_pattern = re.compile(

        r"^\s*(?:\d+\.\d+\.\d+|\d+\.\d+|[A-Z]\.|圖\s*\d+|Fig\.?\s*\d+|Figure\s+\d+)",
        re.IGNORECASE | re.MULTILINE
    )
    section_pattern = re.compile(rf"^\s*\d*\.?\s*{re.escape(target_section)}:?", re.IGNORECASE | re.MULTILINE)
    english_pattern = re.compile(r'^[\x00-\x7F]+$')

    is_standalone = False
    is_docx = isinstance(output_doc, DocxDocument)
    if output_doc is None:
        output_doc = DocxDocument()
        is_docx = True
        is_standalone = True

    if is_docx:
        table = output_doc.add_table(rows=1, cols=2)
        header_row = table.rows[0]
        header_row.cells[0].text = "Packaging test report No."
        header_row.cells[1].text = "Rationale for Test Article Selection"
    else:
        if section is None:
            section = output_doc.AddSection()
        table = section.AddTable(True)
        table.ResetCells(1, 2)
        row = table.Rows.get_Item(0)
        cell1 = row.Cells.get_Item(0)
        cell2 = row.Cells.get_Item(1)
        cell1.AddParagraph().AppendText("Packaging test report No.")
        cell2.AddParagraph().AppendText("Rationale for Test Article Selection")
        bg_color = Color.FromRgb(0xBA, 0xE0, 0xD2)
        cell1.CellFormat.BackColor = bg_color
        cell2.CellFormat.BackColor = bg_color

    for filename in os.listdir(pdf_folder_path):
        if not filename.lower().endswith(".pdf"):
            continue
        pdf_path = os.path.join(pdf_folder_path, filename)
        doc_pdf = fitz.open(pdf_path)
        all_text = []
        for page in doc_pdf:
            width, height = page.rect.width, page.rect.height
            capture_rect = fitz.Rect(0, height * upper_ratio, width, height * lower_ratio)
            blocks = page.get_text("blocks", clip=capture_rect)
            all_text.extend([block[4].strip() for block in blocks if block[4].strip()])
        doc_pdf.close()

        full_text = "\n".join(all_text)
        capture_mode = False
        section_lines = []
        for line in full_text.splitlines():
            if section_pattern.match(line):
                capture_mode = True
                if english_pattern.match(line):
                    section_lines.append(line)
            elif capture_mode and stop_pattern.match(line):
                break
            elif capture_mode and english_pattern.match(line):
                section_lines.append(line)

        extracted_text = " ".join(section_lines).strip()
        if extracted_text:
            match = re.search(r"(UOC|United)", extracted_text, re.IGNORECASE)
            if match:
                extracted_text = extracted_text[:match.end()]
            if not extracted_text.endswith("."):
                extracted_text += "."
        else:
            extracted_text = "（未找到英文內容）"

        table_filename = filename.split(' ')[0]
        if is_docx:
            new_row = table.add_row()
            new_row.cells[0].text = table_filename
            new_row.cells[1].text = extracted_text
        else:
            new_row = TableRow(output_doc)
            cell1 = TableCell(output_doc)
            cell2 = TableCell(output_doc)
            new_row.Cells.Add(cell1)
            new_row.Cells.Add(cell2)
            cell1.AddParagraph().AppendText(table_filename)
            cell2.AddParagraph().AppendText(extracted_text)
            table.Rows.Add(new_row)

    if is_standalone:
        if is_docx:
            output_doc.save("pdf_chapter_table.docx")
        else:
            output_doc.SaveToFile("pdf_chapter_table.docx", FileFormat.Docx)
    print(f"已將PDF章節 {target_section} 擷取至表格")


def extract_word_all_content(
    input_file: str,
    output_image_path: str | None = None,
    output_doc=None,
    section=None,
    *,
    output_docx_path: str | None = None,
    ignore_toc: bool = True,
    ignore_header_footer: bool = True,
):
    if not os.path.isfile(input_file):
        raise FileNotFoundError(f"input file not found: {input_file}")

    out_path = output_docx_path or _build_output_docx_path(input_file, "body")
    extract_body_with_options(
        input_docx=input_file,
        output_docx=out_path,
        ignore_toc_and_before=ignore_toc,
        ignore_header_footer=ignore_header_footer,
    )

    appended_section = None
    if output_doc is not None:
        target_section = section or output_doc.AddSection()
        _append_docx_to_section(out_path, output_doc, target_section)
        appended_section = target_section

    return {"output_docx": out_path, "section": appended_section}


def _normalize_text(value: str) -> str:
    return " ".join(value.split()) if value else ""


def _detect_image_extension(image_bytes: bytes, default_ext: str = "png") -> str:
    """Return a lowercase file extension based on the image signature."""

    signatures = {
        b"\x89PNG\r\n\x1a\n": "png",
        b"\xff\xd8\xff": "jpg",
        b"GIF87a": "gif",
        b"GIF89a": "gif",
        b"BM": "bmp",
        b"II*\x00": "tif",
        b"MM\x00*": "tif",
        b"\x00\x00\x01\x00": "ico",
    }
    for header, ext in signatures.items():
        if image_bytes.startswith(header):
            return ext
    return default_ext


def _save_picture_with_original_format(
    picture: DocPicture, image_dir: str, image_count: list[int]
) -> str:
    """Persist ``picture.ImageBytes`` using the detected image format.

    Parameters
    ----------
    picture : DocPicture
        The picture to save.
    image_dir : str
        Target directory for the exported image.
    image_count : list[int]
        Mutable counter used to build incrementing file names.

    Returns
    -------
    str
        The saved filename (not the absolute path).
    """

    image_bytes = picture.ImageBytes
    ext = _detect_image_extension(image_bytes)
    file_name = f"Image-{image_count[0]}.{ext}"
    img_path = os.path.join(image_dir, file_name)
    with open(img_path, "wb") as img:
        img.write(image_bytes)
    image_count[0] += 1
    return file_name


def is_heading_paragraph(paragraph: Paragraph) -> bool:
    """Return True when the paragraph uses a Heading style."""
    style_name = getattr(paragraph, "StyleName", "") or ""
    return "heading" in style_name.lower()


def extract_word_chapter(
    input_file: str,
    target_chapter_section: str,
    use_chapter_title: bool = False,
    target_chapter_title: str = "",
    output_image_path: str | None = None,
    output_doc=None,
    section=None,
    *,
    explicit_end_title: str | None = None,
    explicit_end_number: str | None = None,
    target_subtitle: str | None = None,
    subheading_strict_match: bool = True,
    hide_chapter_title: bool = False,
    ignore_header_footer: bool = True,
    ignore_toc: bool = True,
    output_docx_path: str | None = None,
    llm_boundary_fallback: bool | None = None,
    llm_boundary_model_id: str | None = None,
):
    if not os.path.isfile(input_file):
        raise FileNotFoundError(f"input file not found: {input_file}")

    raw_section = str(target_chapter_section or "").strip()
    start_section = raw_section
    end_section = (explicit_end_number or "").strip()
    heading_text = target_chapter_title.strip()
    end_title = (explicit_end_title or "").strip()

    # Support direct-call range syntax like "1.1.1-1.1.3" (or with title suffix),
    # including trailing dots like "1. 測試1.1".
    parsed_start, parsed_end, parsed_title = _parse_chapter_section_expression(raw_section)
    if parsed_start:
        start_section = parsed_start
        if not end_section and parsed_end:
            end_section = parsed_end
        if not heading_text and parsed_title:
            heading_text = parsed_title

    if not heading_text and use_chapter_title:
        heading_text = start_section

    # Allow combined end marker in one field:
    # "1.1.3 Accessories not included but necessary for use"
    end_match = re.match(r"^(\d+(?:\.\d+)*\.?)(?:\s+(.+))?$", end_title)
    if end_match:
        if not end_section:
            end_section = end_match.group(1).rstrip(".")
        if end_match.group(2):
            end_title = end_match.group(2).strip()

    start_heading = heading_text or start_section
    out_path = output_docx_path or _build_output_docx_path(input_file, f"section_{start_heading}")

    # Only trim to a subheading when the caller explicitly requests one.
    subheading_to_use = target_subtitle if target_subtitle else None

    extract_section_docx_xml(
        input_docx=input_file,
        output_docx=out_path,
        start_heading_text=start_heading,
        start_number=start_section,
        explicit_end_title=(end_title or None),
        explicit_end_number=(end_section or None),
        ignore_header_footer=ignore_header_footer,
        ignore_toc=ignore_toc,
        subheading_text=subheading_to_use,
        subheading_strict_match=subheading_strict_match,
        subheading_debug=False,
        llm_boundary_fallback=llm_boundary_fallback,
        llm_boundary_model_id=llm_boundary_model_id,
        strict_heading_number_match=True,
    )

    captured_title = _read_first_paragraph_text(out_path) or start_heading
    if hide_chapter_title and captured_title:
        hide_paragraphs_with_text(out_path, [captured_title])

    appended_section = None
    if output_doc is not None:
        target_section = section or output_doc.AddSection()
        start_idx = target_section.Body.ChildObjects.Count
        _append_docx_to_section(out_path, output_doc, target_section)
        if hide_chapter_title and captured_title:
            _hide_titles_in_section(target_section, [captured_title], start_index=start_idx)
        appended_section = target_section

    result = {
        "captured_titles": [captured_title] if (hide_chapter_title and captured_title) else [],
        "output_docx": out_path,
    }
    return result


def is_inline_subtitle_spire(paragraph: Paragraph) -> bool:
    """判斷這個 Spire 段落是不是像 'Device trade name' 這種小標題。
    規則：
    - 段落有文字
    - StyleName == 'Normal'
    - 所有有文字的 TextRange 都是粗體
    """
    # 1. 取得整段文字
    full_text = ""
    for i in range(paragraph.ChildObjects.Count):
        sub = paragraph.ChildObjects.get_Item(i)
        if sub.DocumentObjectType == DocumentObjectType.TextRange:
            full_text += sub.Text or ""
    if not full_text.strip():
        return False

    # 2. 只處理 Normal 樣式
    style_name = getattr(paragraph, "StyleName", None)
    if style_name != "Normal":
        return False

    # 3. 所有有文字的 run 都要是粗體
    has_text_run = False
    for i in range(paragraph.ChildObjects.Count):
        sub = paragraph.ChildObjects.get_Item(i)
        if sub.DocumentObjectType == DocumentObjectType.TextRange:
            tr: TextRange = sub
            text = (tr.Text or "").strip()
            if not text:
                continue
            has_text_run = True
            # Bold 不是 True 就視為不是小標題
            if tr.CharacterFormat.Bold is not True:
                return False

    return has_text_run

def extract_word_subsection(
    input_file: str,
    outer_chapter_section: str,   # 例如 "1.1.1"
    subsection_title: str,        # 例如 "Principles of operation and mode of action"
    output_image_path: str = "images",
    output_doc=None,
    section=None,):
    """
    從 Word 檔中擷取「指定章節內的某一個小節」內容。

    範圍：
      1. 先進入 outer_chapter_section 章節，例如 1.1.1
      2. 在章節內找到 subsection_title 這一行（Normal + 粗體）
      3. 從小節標題「下一行開始」擷取內容（文字、圖片、表格）
      4. 遇到下一個 inline 小標題（Normal + 粗體）或者下一個章節 (1.1.2...) 就停止
    """

    os.makedirs(output_image_path, exist_ok=True)

    # 章節起始：例如 "1.1.1"
    chapter_pattern = re.compile(rf"^\s*{re.escape(outer_chapter_section)}(\s|$)", re.IGNORECASE)

    # 章節停止：以前綴 "1.1" 找下一個章節 (1.1.2, 1.1.3 ...)
    stop_prefix = outer_chapter_section.rsplit(".", 1)[0]
    chapter_stop_pattern = re.compile(rf"^\s*{re.escape(stop_prefix)}(\.\d+)?(\s|$)", re.IGNORECASE)


    # 小節標題：整行比對文字
    subsection_pattern = re.compile(rf"^\s*{re.escape(subsection_title.strip())}\s*$", re.IGNORECASE)

    input_doc = Document()
    input_doc.LoadFromFile(input_file)

    is_standalone_doc = False
    if output_doc is None or section is None:
        output_doc = Document()
        section = output_doc.AddSection()
        is_standalone_doc = True

    nodes = queue.Queue()
    nodes.put(input_doc)

    image_count = [1]

    in_chapter = False
    in_subsection = False
    skip_first_line_after_title = False
    done = False

    def add_table_to_section(sec, table):
        """將表格複製到輸出文件，避免跨頁斷行問題。"""
        try:
            cloned = table.Clone()
            cloned.TableFormat.IsBreakAcrossPages = False
            for i in range(cloned.Rows.Count):
                cloned.Rows.get_Item(i).RowFormat.IsBreakAcrossPages = False
            sec.Tables.Add(cloned)
            sep_para = sec.AddParagraph()
            sep_para.AppendText("\u200B")
        except Exception as e:
            print("處理表格錯誤:", e)

    while nodes.qsize() > 0 and not done:
        node = nodes.get()
        for i in range(node.ChildObjects.Count):
            child = node.ChildObjects.get_Item(i)

            if isinstance(child, Paragraph):
                # 略過目錄樣式
                if "toc" in child.StyleName.lower() or "目錄" in child.StyleName.lower():
                    continue

                # 組出原段落文字（含 ListText）
                paragraph_text_stripped = _get_paragraph_text_stripped(child)

                # 1) 判斷是否進入指定章節 1.1.1
                if not in_chapter:
                    if chapter_pattern.match(paragraph_text_stripped):
                        in_chapter = True
                    continue

                # 2) 章節內，先看是否遇到下一個章節（1.1.2...），整個擷取結束
                if paragraph_text_stripped and chapter_stop_pattern.match(paragraph_text_stripped):
                    in_chapter = False
                    in_subsection = False
                    done = True
                    break

                # 3) 在 1.1.1 內，尚未進入小節 → 尋找目標小節標題
                if in_chapter and not in_subsection:
                    if is_inline_subtitle_spire(child) and _normalize_text(paragraph_text_stripped) == _normalize_text(subsection_title):
                        in_subsection = True
                        skip_first_line_after_title = True
                    continue

                # 4) 已在小節範圍內
                if in_subsection:
                    # 4-1) 遇到下一個 inline 小標題（Normal + 全粗體，文字不同）→ 結束小節
                    if is_inline_subtitle_spire(child):
                        if _normalize_text(paragraph_text_stripped) != _normalize_text(subsection_title):
                            print("STOP BY NEXT SUBTITLE")
                            print(f"-> text: {paragraph_text_stripped}")
                            print(f"-> style: {child.StyleName}")
                            in_subsection = False
                            done = True
                            break
                        else:
                            # 再遇到同標題，保守跳過
                            continue

                    # 4-2) 小節標題下一行開始才真正輸出內容
                    paragraph_alignment = getattr(child.Format, "HorizontalAlignment", None)

                    # 建立一個文字＋圖片標記的字串
                    text_with_markers = ""
                    for j in range(child.ChildObjects.Count):
                        sub = child.ChildObjects.get_Item(j)
                        if sub.DocumentObjectType == DocumentObjectType.TextRange:
                            text_with_markers += sub.Text or ""
                        elif sub.DocumentObjectType == DocumentObjectType.Picture and isinstance(sub, DocPicture):
                            img_name = _save_picture_with_original_format(
                                sub, output_image_path, image_count
                            )
                            text_with_markers += f"[Image: {img_name}|{sub.Width}|{sub.Height}]"

                    # 小節標題下一行的第一個內容段落也要輸出
                    if skip_first_line_after_title:
                        skip_first_line_after_title = False

                    if text_with_markers.strip():
                        para = section.AddParagraph()
                        # 拆解 [Image: ...] 標記，把圖片塞回去
                        for part in re.split(r'(\[Image:.+?\])', text_with_markers):
                            if part.startswith("[Image:"):
                                try:
                                    content = part[7:-1]
                                    img_name, width, height = content.split("|")
                                    width = float(width)
                                    height = float(height)
                                except ValueError:
                                    img_name = part[7:-1].strip()
                                    width = height = None
                                img_path = os.path.join(output_image_path, img_name.strip())
                                if os.path.isfile(img_path):
                                    pic = para.AppendPicture(img_path)
                                    if width and height:
                                        pic.Width = width
                                        pic.Height = height
                                    if paragraph_alignment is not None:
                                        para.Format.HorizontalAlignment = paragraph_alignment
                            else:
                                if part:
                                    para.AppendText(part)

            elif isinstance(child, Table):
                # 表格只有在小節範圍內才複製
                if in_subsection:
                    add_table_to_section(section, child)
            elif isinstance(child, Section):
                nodes.put(child.Body)
            elif isinstance(child, ICompositeObject):
                doc_type = getattr(child, "DocumentObjectType", None)
                # if doc_type in (DocumentObjectType.Header, DocumentObjectType.Footer):
                #     continue
                if _is_header_or_footer(doc_type):
                    continue
                nodes.put(child)

    if is_standalone_doc:
        output_doc.SaveToFile("word_subsection_result.docx", FileFormat.Docx)
        input_doc.Close()
    else:
        input_doc.Close()

    print(f"已擷取 {outer_chapter_section} 中小節「{subsection_title}」的文字與圖片")

def center_table_figure_paragraphs(input_file: str) -> bool:
    pattern = re.compile(r'^\s*(Table|Figure)\s+', re.IGNORECASE)
    doc = Document()
    try:
        doc.LoadFromFile(input_file)
    except Exception as e:
        print(f"錯誤：無法加載文件 {input_file}: {str(e)}")
        return False

    nodes = queue.Queue()
    nodes.put(doc)

    while nodes.qsize() > 0:
        node = nodes.get()
        for i in range(node.ChildObjects.Count):
            child = node.ChildObjects.get_Item(i)
            if isinstance(child, Paragraph):
                if "toc" in child.StyleName.lower() or "目錄" in child.StyleName.lower():
                    continue
                paragraph_text = _get_paragraph_text_stripped(child)
                if pattern.match(paragraph_text):
                    child.Format.HorizontalAlignment = HorizontalAlignment.Center
            elif isinstance(child, ICompositeObject):
                nodes.put(child)

    try:
        doc.SaveToFile(input_file, FileFormat.Docx)
        print(f"{input_file}，以將表格標題或圖片標題置中")
        return True
    except Exception as e:
        print(f"錯誤：保存文件 {input_file} 時出錯: {str(e)}")
        return False
    finally:
        doc.Close()


def _iter_paragraphs(parent):
    """Yield paragraphs in parent recursively, including those in tables.""" 
    if hasattr(parent, "paragraphs"):
        for p in parent.paragraphs:
            yield p
    if hasattr(parent, "tables"):
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_paragraphs(cell)

def remove_hidden_runs(
    input_file: str,
    preserve_texts: Optional[Iterable[str]] = None,
) -> bool:
    """Clear text in hidden runs without removing XML nodes."""
    try:
        doc = DocxDocument(input_file)
        preserve_set = {
            _normalize_text(t)
            for t in (preserve_texts or [])
            if isinstance(t, str) and _normalize_text(t)
        }
        for para in list(_iter_paragraphs(doc)):
            normalized_para_text = _normalize_text(para.text)
            if preserve_set and normalized_para_text in preserve_set:
                continue
            has_image = bool(para._element.xpath('.//w:drawing | .//w:pict'))
            if has_image:
                continue
            in_table = False
            parent = para._element.getparent()
            while parent is not None:
                if parent.tag == qn('w:tc'):
                    in_table = True
                    break
                parent = parent.getparent()
            if in_table:
                continue
            for run in para.runs:
                if not run.font.hidden:
                    continue
                for text_node in run._element.iter(qn('w:t')):
                    text_node.text = ""
                for text_node in run._element.iter(qn('w:instrText')):
                    text_node.text = ""
        doc.save(input_file)
        return True
    except Exception as e:
        print(f"錯誤：移除隱藏文字 {input_file} 時出錯: {str(e)}")
        return False


def hide_paragraphs_with_text(
    input_file: str,
    texts_to_hide: Iterable[str],
) -> bool:
    """Mark paragraphs whose text matches any provided strings as hidden."""
    cleaned = [
        t.strip()
        for t in texts_to_hide
        if isinstance(t, str) and t.strip()
    ]
    if not cleaned:
        return True
    targets = {_normalize_text(t) for t in cleaned}
    try:
        doc = DocxDocument(input_file)
        for para in _iter_paragraphs(doc):
            if _normalize_text(para.text) in targets:
                for run in para.runs:
                    run.font.hidden = True
        doc.save(input_file)
        return True
    except Exception as e:
        print(f"錯誤：隱藏段落於 {input_file} 時出錯: {str(e)}")
        return False


def remove_paragraphs_with_text(
    input_file: str,
    texts_to_remove: Iterable[str],
) -> bool:
    """Remove paragraphs whose text matches any provided strings.

    Paragraphs inside table cells keep an empty placeholder if they would be
    the last remaining paragraph to avoid corrupting the table structure.
    """

    cleaned = [
        t.strip()
        for t in texts_to_remove
        if isinstance(t, str) and t.strip()
    ]
    if not cleaned:
        return True

    targets = {_normalize_text(t) for t in cleaned}
    try:
        doc = DocxDocument(input_file)
        for para in list(_iter_paragraphs(doc)):
            if _normalize_text(para.text) not in targets:
                continue

            parent = para._element.getparent()
            if parent is not None and parent.tag == qn('w:tc'):
                paragraph_count = len(parent.findall(qn('w:p')))
                if paragraph_count <= 1:
                    for run in list(para.runs):
                        para._element.remove(run._element)
                    continue

            element = para._element
            container = element.getparent()
            if container is not None:
                container.remove(element)

        doc.save(input_file)
        return True
    except Exception as e:
        print(f"錯誤：移除段落於 {input_file} 時出錯: {str(e)}")
        return False


def apply_basic_style(
    input_file: str,
    western_font: str = "Times New Roman",
    east_asian_font: str = "新細明體",
    font_size: int = 12,
    line_spacing: float = 1.5,
    space_before: int = 6,
    space_after: int = 6,
) -> bool:
    """為整份文件套用基本字型與行距設定。"""
    try:
        doc = DocxDocument(input_file)
        for para in _iter_paragraphs(doc):
            pf = para.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = line_spacing
            pf.space_before = Pt(space_before)
            pf.space_after = Pt(space_after)
            for run in para.runs:
                run.font.name = western_font
                set_run_font_eastasia(run, east_asian_font)
                run.font.size = Pt(font_size)
        doc.save(input_file)
        return True
    except Exception as e:
        print(f"錯誤：套用樣式至 {input_file} 時出錯: {str(e)}")
        return False
    

def _is_header_or_footer(doc_type) -> bool:
    header = getattr(DocumentObjectType, "Header", None)
    footer = getattr(DocumentObjectType, "Footer", None)
    header_footer = getattr(DocumentObjectType, "HeaderFooter", None)
    return doc_type in (header, footer, header_footer)


def extract_specific_figure_from_word(
    input_file: str,
    target_chapter_section: str,   # 例如 "2.1.1"
    target_caption_label: str,      # 例如 "Figure 1."
    target_subtitle: str | None = None,  # 可選，有就填，沒有就 None
    target_chapter_title: str | None = None,
    target_figure_title: str | None = None,
    target_figure_index: int | str | None = None,
    output_image_path: str = "figure_output",
    output_docx_path: str | None = None,
    output_doc=None,
    section=None,
    *,
    allow_table_figure_container: bool = False,
    include_caption: bool = True,
    ignore_header_footer: bool = True,
    save_output: bool = True,
    return_reason: bool = False,
) -> bool | dict:
    """
    Compatibility wrapper: route figure extraction to the XML-based implementation.
    """
    _ = output_image_path  # Kept for backward-compatible call sites.
    if not os.path.isfile(input_file):
        raise FileNotFoundError(f"input file not found: {input_file}")

    append_to_spire_doc = output_doc is not None
    effective_output_docx_path = output_docx_path
    created_temp_output = False

    if append_to_spire_doc and not effective_output_docx_path:
        suffix = f"figure_{target_caption_label or 'extract'}"
        effective_output_docx_path = _build_output_docx_path(input_file, suffix)
        created_temp_output = True

    xml_save_output = save_output or append_to_spire_doc
    result = extract_specific_figure_from_word_xml(
        input_file=input_file,
        output_docx_path=effective_output_docx_path,
        target_chapter_section=target_chapter_section,
        target_caption_label=target_caption_label,
        target_subtitle=target_subtitle,
        target_chapter_title=target_chapter_title,
        target_figure_title=target_figure_title,
        target_figure_index=target_figure_index,
        allow_table_figure_container=allow_table_figure_container,
        include_caption=include_caption,
        ignore_header_footer=ignore_header_footer,
        save_output=xml_save_output,
        return_reason=return_reason,
    )

    ok = bool(result.get("ok")) if return_reason and isinstance(result, dict) else bool(result)
    if append_to_spire_doc and ok and effective_output_docx_path and os.path.isfile(effective_output_docx_path):
        _append_docx_to_section(effective_output_docx_path, output_doc, section)

    if created_temp_output and effective_output_docx_path and os.path.isfile(effective_output_docx_path):
        try:
            os.remove(effective_output_docx_path)
        except OSError:
            pass

    return result if return_reason else ok


def extract_specific_table_from_word(
    input_file: str,
    output_docx_path: str | None,      # 另存新檔的路徑，例如 "check_result.docx"
    target_chapter_section: str,   # 章節編號，例如 "2.1.1"
    target_caption_label: str,       # 表格標題開頭，例如 "Table 1."
    target_subtitle: str | None = None,
    target_chapter_title: str | None = None,
    *,
    target_table_title: str | None = None,
    target_table_index: int | str | None = None,
    include_caption: bool = True,
    ignore_header_footer: bool = True,
    save_output: bool = True,      # 是否存檔；如不存則僅回傳找到與否
    return_reason: bool = False,
) -> bool | dict:
    """
    Compatibility wrapper: route table extraction to the XML-based implementation.
    """
    return extract_specific_table_from_word_xml(
        input_file=input_file,
        output_docx_path=output_docx_path,
        target_chapter_section=target_chapter_section,
        target_caption_label=target_caption_label,
        target_subtitle=target_subtitle,
        target_chapter_title=target_chapter_title,
        target_table_title=target_table_title,
        target_table_index=target_table_index,
        include_caption=include_caption,
        ignore_header_footer=ignore_header_footer,
        save_output=save_output,
        return_reason=return_reason,
    )
