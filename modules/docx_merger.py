import warnings
warnings.filterwarnings("ignore", category=UserWarning, module='docxcompose')
from docx import Document
from docx.oxml.ns import qn
from docxcompose.composer import Composer
from typing import List


def _first_body_block_tag(doc: Document) -> str | None:
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag in {qn("w:p"), qn("w:tbl")}:
            return child.tag
    return None


def _last_body_block_tag(doc: Document) -> str | None:
    body = doc.element.body
    blocks = [child for child in body.iterchildren() if child.tag in {qn("w:p"), qn("w:tbl")}]
    if not blocks:
        return None
    return blocks[-1].tag


def merge_word_docs(file_list: List[str], output_path: str) -> None:
    """
    Merge multiple DOCX files in order using docxcompose.

    Parameters
    ----------
    file_list : list[str]
        File paths to merge in order.
    output_path : str
        Destination path for the merged DOCX.
    """
    if not file_list:
        return

    master_doc = Document(file_list[0])
    composer = Composer(master_doc)

    for path in file_list[1:]:
        append_doc = Document(path)
        # Prevent Word from visually merging two adjacent tables from
        # neighboring fragments into one continuous table.
        if _last_body_block_tag(master_doc) == qn("w:tbl") and _first_body_block_tag(append_doc) == qn("w:tbl"):
            master_doc.add_paragraph("")
        composer.append(append_doc)

    composer.save(output_path)
