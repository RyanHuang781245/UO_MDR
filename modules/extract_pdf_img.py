from __future__ import annotations

import argparse
import os
import tempfile

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_pdf_pages_to_folder(
    pdf_path: str,
    output_dir: str,
    *,
    zoom: float = 3.0,
) -> dict[str, object]:
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")

    os.makedirs(output_dir, exist_ok=True)
    matrix = fitz.Matrix(zoom, zoom)
    page_count = 0

    with fitz.open(pdf_path) as pdf:
        for page_index in range(len(pdf)):
            pix = pdf[page_index].get_pixmap(matrix=matrix, alpha=False)
            output_path = os.path.join(output_dir, f"page_{page_index + 1}.png")
            pix.save(output_path)
            page_count += 1

    return {"output_dir": output_dir, "pages": page_count}


def extract_pdf_pages_to_docx(
    input_pdf: str,
    output_docx: str,
    *,
    zoom: float = 3.0,
) -> dict[str, object]:
    if not os.path.isfile(input_pdf):
        raise FileNotFoundError(f"PDF file not found: {input_pdf}")

    os.makedirs(os.path.dirname(output_docx) or ".", exist_ok=True)
    doc = DocxDocument()
    section = doc.sections[0]
    available_width = section.page_width - section.left_margin - section.right_margin
    page_count = 0

    with tempfile.TemporaryDirectory(prefix="pdf_pages_") as temp_dir:
        with fitz.open(input_pdf) as pdf:
            matrix = fitz.Matrix(zoom, zoom)
            for page_index in range(len(pdf)):
                image_path = os.path.join(temp_dir, f"page_{page_index + 1}.png")
                pix = pdf[page_index].get_pixmap(matrix=matrix, alpha=False)
                pix.save(image_path)

                if page_count > 0:
                    doc.add_page_break()

                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run().add_picture(image_path, width=available_width)
                page_count += 1

    if page_count == 0:
        raise RuntimeError("No pages extracted from PDF")

    doc.save(output_docx)
    return {"output_docx": output_docx, "pages": page_count}


def main() -> None:
    parser = argparse.ArgumentParser(description="Export PDF pages as PNG images.")
    parser.add_argument("pdf_path", help="Path to the input PDF file")
    parser.add_argument("output_dir", help="Directory to store rendered page images")
    parser.add_argument("--zoom", type=float, default=2.0, help="Render zoom factor")
    args = parser.parse_args()

    result = export_pdf_pages_to_folder(args.pdf_path, args.output_dir, zoom=args.zoom)
    print(f"完成，共輸出 {result['pages']} 頁到 {result['output_dir']}")


if __name__ == "__main__":
    main()
