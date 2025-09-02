import re
from typing import Dict
from docx import Document


def _replace_run_text(paragraph, old: str, new: str) -> bool:
    """Replace first occurrence of old in the paragraph runs with new.
    Returns True if replaced."""
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return True
    return False


def renumber_figures_tables(docx_path: str) -> None:
    """Renumber figure and table captions and update references in-place.

    Parameters
    ----------
    docx_path: str
        Path to the Word document to process. The file is modified in-place.
    """
    doc = Document(docx_path)

    fig_map: Dict[str, str] = {}
    table_map: Dict[str, str] = {}
    fig_counter = 1
    table_counter = 1
    caption_pattern = re.compile(r'^(Figure|Fig\.?|Table)\s*(\d+)', re.IGNORECASE)

    # First pass: renumber captions and build mapping of old->new numbers
    for para in doc.paragraphs:
        match = caption_pattern.match(para.text.strip())
        if not match:
            continue
        label, old_num = match.group(1), match.group(2)
        if label.lower().startswith('fig'):
            fig_map[old_num] = str(fig_counter)
            _replace_run_text(para, old_num, str(fig_counter))
            fig_counter += 1
        else:
            table_map[old_num] = str(table_counter)
            _replace_run_text(para, old_num, str(table_counter))
            table_counter += 1

    # Second pass: update in-text references using the mappings
    ref_pattern = re.compile(r'(Figure|Fig\.?|Table)\s*(\d+)', re.IGNORECASE)
    for para in doc.paragraphs:
        for run in para.runs:
            def _repl(m: re.Match) -> str:
                label, num = m.group(1), m.group(2)
                if re.match(r'Fig\.?|Figure', label, re.IGNORECASE):
                    new_num = fig_map.get(num)
                else:
                    new_num = table_map.get(num)
                return f"{label} {new_num}" if new_num else m.group(0)

            new_text = ref_pattern.sub(_repl, run.text)
            run.text = new_text

    doc.save(docx_path)
