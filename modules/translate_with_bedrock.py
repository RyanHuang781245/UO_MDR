import os
import time
import argparse
from typing import List, Optional, Tuple

import boto3
from botocore.exceptions import ClientError, EndpointConnectionError

# ======== 可調參數 ========
# 重要：把這裡換成你的 Inference Profile ID（或支援 on-demand 的基礎模型 ID）
MODEL_ID = os.getenv("BEDROCK_MODEL_ID", "apac.anthropic.claude-3-7-sonnet-20250219-v1:0")
REGION   = os.getenv("AWS_REGION", "ap-southeast-2")

# chunking：以字元數估計，保守一些（模型與語言不同，token≈字元數/3~4，這裡用字元切）
CHUNK_SIZE = 4000         # 每段最大字元數
CHUNK_OVERLAP = 200       # 重疊字元，讓上下文更連貫
MAX_RETRY = 3
RETRY_BACKOFF = 2         # 指數退避基數（秒）

# ======== 基本讀檔工具 ========
def read_text_file(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def read_pdf(path: str) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError("未安裝 PyMuPDF。請執行：pip install pymupdf")
    doc = fitz.open(path)
    texts = []
    for page in doc:
        texts.append(page.get_text("text"))
    return "\n".join(texts)

def read_docx(path: str) -> str:
    try:
        import docx
    except ImportError:
        raise RuntimeError("未安裝 python-docx。請執行：pip install python-docx")
    d = docx.Document(path)
    # 基本策略：段落逐段讀，表格簡單以「 | 」連接儲存格
    parts = []
    for p in d.paragraphs:
        parts.append(p.text)
    # 表格（簡易轉文字）
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
        parts.append("")  # 表格與段落間空行
    return "\n".join(parts)

def load_text(path: str) -> str:
    ext = os.path.splitext(path.lower())[1]
    if ext == ".txt":
        return read_text_file(path)
    elif ext == ".pdf":
        return read_pdf(path)
    elif ext == ".docx":
        return read_docx(path)
    else:
        raise ValueError(f"不支援的副檔名：{ext}（支援 .txt/.pdf/.docx）")

# ======== chunk 工具 ========
def chunk_text(s: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    s = s.strip()
    if not s:
        return []
    chunks = []
    start = 0
    n = len(s)
    while start < n:
        end = min(start + size, n)
        chunk = s[start:end]
        chunks.append(chunk)
        if end == n:
            break
        # 下一段從 end - overlap 開始，確保不為負且不倒退
        start = max(0, end - overlap)
    return chunks

# ======== Bedrock Converse 呼叫 ========
def bedrock_client():
    return boto3.client("bedrock-runtime", region_name=REGION)

SYSTEM_PROMPT = (
    "You are a precise, domain-aware English translator.\n"
    "- Translate the user's content into **natural, clear English**.\n"
    "- Preserve **headings, lists, and basic structure**.\n"
    "- Keep numbers, units, chemical names, model names as-is.\n"
    "- If the source contains tables, keep them in Markdown rows when possible.\n"
    "- Do not add extra commentary; output **English only**.\n"
)

def translate_chunk(brt, model_id: str, text: str, temperature: float = 0.2, top_p: float = 0.9, max_tokens: int = 2048) -> str:
    # 帶重試與指數退避
    for attempt in range(1, MAX_RETRY + 1):
        try:
            resp = brt.converse(
                modelId=model_id,
                system=[{"text": SYSTEM_PROMPT}],
                messages=[{
                    "role": "user",
                    "content": [{"text": f"Translate the following content into English:\n\n{text}"}]
                }],
                inferenceConfig={
                    "maxTokens": max_tokens,
                    "temperature": temperature,
                    "topP": top_p
                },
            )
            out = resp["output"]["message"]["content"][0]["text"]
            return out.strip()
        except (ClientError, EndpointConnectionError) as e:
            if attempt == MAX_RETRY:
                raise
            sleep_sec = (RETRY_BACKOFF ** (attempt - 1))
            time.sleep(sleep_sec)

# ======== 主流程 ========
def translate_file(input_path: str, output_path: str, model_id: Optional[str] = None):
    model_id = model_id or MODEL_ID
    text = load_text(input_path)

    # 簡單的「文件導引」包裹，讓模型知道整體任務
    # header = (
    #     "# Translated Document (to English)\n\n"
    #     f"> Source file: `{os.path.basename(input_path)}`\n\n"
    #     "---\n\n"
    # )

    chunks = chunk_text(text)
    if not chunks:
        raise RuntimeError("文件內容為空，無法翻譯。")

    brt = bedrock_client()

    outputs: List[str] = []
    for i, ck in enumerate(chunks, 1):
        # 在段首加入章節提示，提升上下文銜接（可視需要移除）
        ck_prompt = f"[Part {i}/{len(chunks)}]\n{ck}"
        translated = translate_chunk(brt, model_id, ck_prompt)
        # outputs.append(f"<!-- Part {i}/{len(chunks)} -->\n{translated}\n")
        outputs.append(translated)

    # final_text = header + "\n\n".join(outputs)
    final_text = "\n\n".join(outputs)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(final_text)
    return output_path

def main():
    parser = argparse.ArgumentParser(description="Translate a document to English via Amazon Bedrock Converse API.")
    parser.add_argument("--input", "-i", required=True, help="輸入檔案路徑（.txt/.pdf/.docx）")
    parser.add_argument("--output", "-o", default=None, help="輸出檔案路徑（預設：同名 .md）")
    parser.add_argument("--model-id", "-m", default=None, help="模型或 Inference Profile ID（預設讀 BEDROCK_MODEL_ID 環境變數）")
    parser.add_argument("--region", "-r", default=None, help="AWS 區域（預設讀 AWS_REGION 環境變數）")
    args = parser.parse_args()

    global REGION
    if args.region:
        REGION = args.region

    in_path = args.input
    if not os.path.isfile(in_path):
        raise FileNotFoundError(f"找不到檔案：{in_path}")

    out_path = args.output or (os.path.splitext(in_path)[0] + "_translated.md")
    result_path = translate_file(in_path, out_path, model_id=args.model_id)
    print(f"翻譯完成：{result_path}")

if __name__ == "__main__":
    main()
