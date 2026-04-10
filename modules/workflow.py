
import os
import hashlib
import zipfile
import re
import shutil
from datetime import datetime
from typing import List, Dict, Any, Callable
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.oxml.ns import qn
from .Extract_AllFile_to_FinalWord import (
    extract_pdf_chapter_to_table,
    extract_word_all_content,
    extract_word_chapter,
    extract_specific_figure_from_word,
    extract_specific_table_from_word,
)
from .file_copier import copy_directories, copy_directory, copy_file, copy_files
from .docx_merger import merge_word_docs
from .template_manager import (
    display_order_template_mappings,
    order_template_mappings,
    parse_template_paragraphs,
    render_template_with_mappings,
)
from .docx_provenance import apply_final_provenance, build_provenance_descriptor
from modules.extract_pdf_img import extract_pdf_pages_to_docx
from app.services.execution_service import JobCanceledError


def _resolve_fragment_path(workdir: str, user_path: str | None, idx: int) -> str:
    """Build an absolute path for a fragment DOCX inside workdir."""
    if user_path:
        path = user_path if os.path.isabs(user_path) else os.path.join(workdir, user_path)
    else:
        path = os.path.join(workdir, f"fragment_{idx}.docx")
    os.makedirs(os.path.dirname(path) or workdir, exist_ok=True)
    return path


def _new_docx_fragment(path: str) -> DocxDocument:
    """Return a blank python-docx document ready to save to path."""
    os.makedirs(os.path.dirname(path), exist_ok=True)
    return DocxDocument()


def _docx_has_content(path: str) -> bool:
    if not path or not os.path.isfile(path):
        return False
    try:
        doc = DocxDocument(path)
        for para in doc.paragraphs:
            if (para.text or "").strip():
                return True
        if getattr(doc, "tables", None):
            if doc.tables:
                return True
    except Exception:
        pass
    try:
        with zipfile.ZipFile(path, "r") as zf:
            return any(name.startswith("word/media/") for name in zf.namelist())
    except Exception:
        return False

def _set_alignment(paragraph: DocxParagraph, align: str) -> None:
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    paragraph.alignment = align_map.get(align.lower(), WD_ALIGN_PARAGRAPH.LEFT)


def _clear_list_formatting(paragraph: DocxParagraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is not None:
        ppr.remove(num_pr)


def _clear_indent(paragraph: DocxParagraph) -> None:
    pf = paragraph.paragraph_format
    pf.left_indent = None
    pf.first_line_indent = None
    pf.right_indent = None
    pf.hanging_indent = None


def _to_roman(num: int) -> str:
    if num <= 0:
        return ""
    pairs = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    ]
    result = []
    for value, symbol in pairs:
        while num >= value:
            result.append(symbol)
            num -= value
    return "".join(result)

SUPPORTED_STEPS = {
    # "extract_pdf_chapter_to_table": {
    #     "label": "擷取 PDF 章節至表格（上傳 ZIP）",
    #     "inputs": ["pdf_zip", "target_section", "template_index", "template_mode"],
    #     "accepts": {
    #         "pdf_zip": "file:zip",
    #         "target_section": "text",
    #         "template_index": "text",
    #         "template_mode": "text",
    #     }
    # },
    "extract_pdf_pages_as_images": {
        "label": "擷取 PDF 標籤圖片",
        "inputs": ["input_file", "template_index", "template_mode"],
        "accepts": {
            "input_file": "file:pdf",
            "template_index": "text",
            "template_mode": "text",
        }
    },
    "extract_word_all_content": {
        "label": "擷取 Word 全部內容",
        "inputs": ["input_file", "ignore_toc", "ignore_header_footer", "template_index", "template_mode"],
        "accepts": {
            "input_file": "file:docx",
            "ignore_toc": "bool",
            "ignore_header_footer": "bool",
            "template_index": "text",
            "template_mode": "text",
        }
    },
    "extract_word_chapter": {
        "label": "擷取 Word 指定章節/標題",
        "inputs": [
            "input_file",
            "target_chapter_section",
            "target_chapter_ref_raw",
            "use_chapter_title",
            "target_chapter_title",
            "explicit_end_title",
            "explicit_end_number",
            "target_subtitle",
            "subheading_strict_match",
            "hide_chapter_title",
            "ignore_toc",
            "ignore_header_footer",
            "template_index",
            "template_mode",
        ],
        "accepts": {
            "input_file": "file:docx",
            "target_chapter_section": "text",
            "target_chapter_ref_raw": "text",
            "use_chapter_title": "bool",
            "target_chapter_title": "text",
            "explicit_end_title": "text",
            "explicit_end_number": "text",
            "target_subtitle": "text",
            "subheading_strict_match": "bool",
            "hide_chapter_title": "bool",
            "ignore_toc": "bool",
            "ignore_header_footer": "bool",
            "template_index": "text",
            "template_mode": "text",
        }
    },
    "extract_specific_figure_from_word": {
        "label": "插入 Word 指定章節/標題的特定圖片",
        "inputs": [
            "input_file",
            "target_chapter_section",
            "target_chapter_ref_raw",
            "target_chapter_title",
            "target_subtitle",
            "target_caption_label",
            "target_figure_title",
            "target_figure_index",
            "include_caption",
            "ignore_header_footer",
            "template_index",
            "template_mode",
        ],
        "accepts": {
            "input_file": "file:docx",
            "target_chapter_section": "text",
            "target_chapter_ref_raw": "text",
            "target_chapter_title": "text",
            "target_subtitle": "text",
            "target_caption_label": "text",
            "target_figure_title": "text",
            "target_figure_index": "int",
            "include_caption": "bool",
            "ignore_header_footer": "bool",
            "template_index": "text",
            "template_mode": "text",
        }
    },
    "extract_specific_table_from_word": {
        "label": "插入 Word 指定章節/標題的特定表格",
        "inputs": [
            "input_file",
            "target_chapter_section",
            "target_chapter_ref_raw",
            "target_chapter_title",
            "target_caption_label",
            "target_table_title",
            "target_table_index",
            "target_subtitle",
            "include_caption",
            "ignore_header_footer",
            "template_index",
            "template_mode",
        ],
        "accepts": {
            "input_file": "file:docx",
            "target_chapter_section": "text",
            "target_chapter_ref_raw": "text",
            "target_chapter_title": "text",
            "target_caption_label": "text",
            "target_table_title": "text",
            "target_table_index": "int",
            "target_subtitle": "text",
            "include_caption": "bool",
            "ignore_header_footer": "bool",
            "template_index": "text",
            "template_mode": "text",
        }
    },
    "insert_text": {
        "label": "插入純文字段落",
        "inputs": ["text", "align", "bold", "font_size", "page_break_before", "template_index", "template_mode"],
        "accepts": {
            "text": "text",
            "align": "align",
            "bold": "bool",
            "font_size": "float",
            "page_break_before": "bool",
            "template_index": "text",
            "template_mode": "text",
        }
    },
    "insert_roman_heading": {
        "label": "插入羅馬數字標題",
        "inputs": ["text", "level", "bold", "font_size", "template_index", "template_mode"],
        "accepts": {
            "text": "text",
            "level": "int",
            "bold": "bool",
            "font_size": "float",
            "template_index": "text",
            "template_mode": "text",
        }
    },
    "insert_bulleted_heading": {
        "label": "插入項目符號標題",
        "inputs": ["text", "bold", "font_size", "template_index", "template_mode"],
        "accepts": {
            "text": "text",
            "bold": "bool",
            "font_size": "float",
            "template_index": "text",
            "template_mode": "text",
        }
    },
     "insert_numbered_heading": {
        "label": "插入阿拉伯數字標題",
        "inputs": ["text", "level", "bold", "font_size", "template_index", "template_mode"],
        "accepts": {
            "text": "text",
            "level": "int",
            "bold": "bool",
            "font_size": "float",
            "template_index": "text",
            "template_mode": "text",
        }
    },
    "copy_files": {
        "label": "複製檔案",
        "inputs": ["source_dir", "dest_dir", "keywords", "target_name", "recursive_search"],
        "accepts": {
            "source_dir": "file:path",
            "dest_dir": "file:dir",
            "keywords": "text",
            "target_name": "text",
            "recursive_search": "bool",
        }
    },
    "copy_directory": {
        "label": "複製資料夾",
        "inputs": ["source_dir", "dest_dir", "keywords", "target_name", "recursive_search"],
        "accepts": {
            "source_dir": "file:dir",
            "dest_dir": "file:dir",
            "keywords": "text",
            "target_name": "text",
            "recursive_search": "bool",
        }
    },
    # "renumber_figures_tables": {
    #     "label": "重新編號圖表並更新參照",
    #     "inputs": ["numbering_scope", "figure_start", "table_start"],
    #     "accepts": {
    #         "numbering_scope": "text",
    #         "figure_start": "int",
    #         "table_start": "int",
    #     }
    # }
}

def boolish(v:str)->bool:
    return str(v).lower() in ["1","true","yes","y","on"]


def _dedupe_target_path(path: str) -> str:
    if not os.path.exists(path):
        return path
    root, ext = os.path.splitext(path)
    idx = 2
    while True:
        candidate = f"{root}_{idx}{ext}"
        if not os.path.exists(candidate):
            return candidate
        idx += 1


def _rename_single_copied_path(path: str, target_name: str) -> str:
    requested_name = (target_name or "").strip()
    if not requested_name:
        return path

    parent = os.path.dirname(path)
    current_name = os.path.basename(path)
    stem, ext = os.path.splitext(current_name)
    final_name = requested_name
    if ext and not os.path.splitext(final_name)[1]:
        final_name = f"{final_name}{ext}"
    target_path = os.path.join(parent, final_name)
    if os.path.abspath(target_path) == os.path.abspath(path):
        return path
    target_path = _dedupe_target_path(target_path)
    shutil.move(path, target_path)
    return target_path



def run_workflow(
    steps: List[Dict[str, Any]],
    workdir: str,
    template: Dict[str, Any] | None = None,
    cancel_check: Callable[[], None] | None = None,
) -> Dict[str, Any]:
    def _hash_file(path: str) -> str:
        sha = hashlib.sha256()
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(1024 * 1024), b""):
                sha.update(chunk)
        return sha.hexdigest()

    def _collect_file_metadata(
        steps_data: List[Dict[str, Any]], template_cfg: Dict[str, Any] | None
    ) -> List[Dict[str, Any]]:
        entries: list[Dict[str, Any]] = []
        seen: set[str] = set()

        def add_path(path_raw: str, kind: str) -> None:
            if not path_raw:
                return
            path = os.path.abspath(path_raw)
            if path in seen:
                return
            seen.add(path)
            exists = os.path.exists(path)
            meta = {
                "name": os.path.basename(path),
                "source": path,
                "type": kind,
                "version": "",
                "updated_at": "",
                "size_bytes": None,
                "exists": exists,
            }
            if exists:
                stat = os.stat(path)
                meta["updated_at"] = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
                if os.path.isfile(path):
                    meta["size_bytes"] = stat.st_size
                    try:
                        meta["version"] = _hash_file(path)
                    except Exception:
                        meta["version"] = ""
            entries.append(meta)

        for step in steps_data:
            stype = step.get("type")
            schema = SUPPORTED_STEPS.get(stype, {})
            accepts = schema.get("accepts", {})
            params = step.get("params", {}) or {}
            for key, acc in accepts.items():
                if not isinstance(acc, str) or not acc.startswith("file"):
                    continue
                path_val = params.get(key)
                if not path_val:
                    continue
                kind = "dir" if acc.endswith(":dir") else "file"
                add_path(str(path_val), kind)

        if template_cfg and template_cfg.get("path"):
            add_path(str(template_cfg.get("path")), "template")

        return entries

    log = []
    fragments: list[str] = []
    template_cfg = template or {}
    file_metadata = _collect_file_metadata(steps, template_cfg)
    if file_metadata:
        log.append({"type": "file_metadata", "files": file_metadata, "status": "ok"})
    template_mappings: list[Dict[str, Any]] = []
    template_merge_succeeded = False
    template_mode_default = (template_cfg.get("default_mode") or "insert_after").strip()
    template_paragraphs = template_cfg.get("paragraphs")
    copied_dir_registry: dict[str, dict[str, Any]] = {}
    arabic_counters: list[int] = [0, 0, 0]
    roman_counter = 0
    provenance_counter = 0

    def _next_arabic(level_raw: Any) -> str:
        try:
            level = int(level_raw)
        except Exception:
            level = 0
        if level < 0:
            level = 0
        if level >= len(arabic_counters):
            arabic_counters.extend([0] * (level + 1 - len(arabic_counters)))
        for i in range(level):
            if arabic_counters[i] == 0:
                arabic_counters[i] = 1
        arabic_counters[level] += 1
        for i in range(level + 1, len(arabic_counters)):
            arabic_counters[i] = 0
        parts = [str(arabic_counters[i]) for i in range(level + 1)]
        return ".".join(parts) + ". "

    def _next_roman() -> str:
        nonlocal roman_counter
        roman_counter += 1
        return _to_roman(roman_counter) + ". "

    def _content_type_for_step(stype: str) -> str:
        return {
            "extract_word_chapter": "paragraph",
            "extract_word_all_content": "paragraph",
            "extract_specific_table_from_word": "table",
            "extract_specific_figure_from_word": "figure",
            "extract_pdf_pages_as_images": "pdf_image",
            "extract_pdf_chapter_to_table": "table",
        }.get(stype, "fragment")

    def _build_fragment_provenance(path: str, stype: str) -> dict[str, Any] | None:
        nonlocal provenance_counter
        if not path or not os.path.isfile(path):
            return None

        provenance_counter += 1
        descriptor = build_provenance_descriptor(provenance_counter)
        descriptor["content_type"] = _content_type_for_step(stype)
        descriptor["fragment_path"] = path
        descriptor["fragment_order"] = provenance_counter
        return descriptor

    def _route_fragment(path: str, params: Dict[str, Any], stype: str) -> None:
        provenance = _build_fragment_provenance(path, stype)
        if provenance:
            log[-1]["provenance"] = provenance
        t_idx_raw = params.get("template_index")
        if template_cfg and t_idx_raw not in (None, "", "None"):
            try:
                t_idx = int(t_idx_raw)
            except Exception:
                fragments.append(path)
                return
            mode = (params.get("template_mode") or template_mode_default or "insert_after").strip() or "insert_after"
            template_mappings.append(
                {
                    "index": t_idx,
                    "mode": mode,
                    "content_docx_path": path,
                    "source_step": stype,
                    "source_order": len(template_mappings),
                    "source_id": provenance.get("source_id") if provenance else "",
                }
            )
            log[-1]["template_index"] = t_idx
            log[-1]["template_mode"] = mode
            if provenance:
                log[-1]["provenance"]["template_index"] = t_idx
                log[-1]["provenance"]["template_mode"] = mode
            return
        fragments.append(path)

    def _check_canceled() -> None:
        if cancel_check:
            cancel_check()

    for idx, step in enumerate(steps, start=1):
        _check_canceled()
        stype = step.get("type")
        params = step.get("params", {})
        log.append({"step": idx, "type": stype, "params": params})
        try:
            if stype == "extract_pdf_chapter_to_table":
                import zipfile
                zip_path = params.get("pdf_zip")
                target = params["target_section"]
                if not zip_path or not os.path.isfile(zip_path):
                    raise RuntimeError("Missing or invalid PDF ZIP path")
                extract_dir = os.path.join(workdir, "pdfs_extracted")
                os.makedirs(extract_dir, exist_ok=True)
                with zipfile.ZipFile(zip_path, 'r') as zf:
                    zf.extractall(extract_dir)
                frag_path = _resolve_fragment_path(workdir, params.get("output_docx_path"), idx)
                doc = DocxDocument()
                extract_pdf_chapter_to_table(extract_dir, target, output_doc=doc, section=None)
                doc.save(frag_path)
                _route_fragment(frag_path, params, stype)

            elif stype == "extract_word_all_content":
                infile = params["input_file"]
                ignore_toc = boolish(params.get("ignore_toc", params.get("ignore_toc_and_before", "true")))
                ignore_header_footer = boolish(params.get("ignore_header_footer", "true"))
                frag_path = _resolve_fragment_path(workdir, params.get("output_docx_path"), idx)
                result = extract_word_all_content(
                    infile,
                    output_doc=None,
                    section=None,
                    output_docx_path=frag_path,
                    ignore_toc=ignore_toc,
                    ignore_header_footer=ignore_header_footer,
                )
                out_path = None
                if isinstance(result, dict):
                    out_path = result.get("output_docx") or frag_path
                    log[-1]["output_docx"] = out_path
                if out_path:
                    _route_fragment(out_path, params, stype)

            elif stype == "extract_pdf_pages_as_images":
                infile = params["input_file"]
                frag_path = _resolve_fragment_path(workdir, params.get("output_docx_path"), idx)
                result = extract_pdf_pages_to_docx(
                    infile,
                    frag_path,
                )
                out_path = None
                if isinstance(result, dict):
                    out_path = result.get("output_docx") or frag_path
                    log[-1]["output_docx"] = out_path
                    log[-1]["pages"] = result.get("pages")
                if out_path:
                    _route_fragment(out_path, params, stype)

            elif stype == "extract_word_chapter":
                infile = params["input_file"]
                tsec = params.get("target_chapter_section", "")
                use_title = boolish(params.get("use_chapter_title", params.get("target_title", "false")))
                title_text = params.get("target_chapter_title", params.get("target_title_section", ""))
                end_number = params.get("explicit_end_number", "") or ""
                raw_tsec = str(tsec or "").strip()
                range_match = re.match(r"^(\d+(?:\.\d+)*)(?:\s*-\s*(\d+(?:\.\d+)*))?(?:\s+(.+))?$", raw_tsec)
                if range_match:
                    tsec = range_match.group(1)
                    if not end_number and range_match.group(2):
                        end_number = range_match.group(2)
                    if (not title_text or title_text == raw_tsec) and range_match.group(3):
                        title_text = range_match.group(3)
                    elif title_text == raw_tsec and not range_match.group(3):
                        title_text = ""
                target_subtitle = params.get("target_subtitle", params.get("subheading_text"))
                frag_path = _resolve_fragment_path(workdir, params.get("output_docx_path"), idx)
                result = extract_word_chapter(
                    infile,
                    tsec,
                    use_chapter_title=use_title,
                    target_chapter_title=title_text,
                    explicit_end_title=params.get("explicit_end_title") or None,
                    explicit_end_number=end_number or None,
                    target_subtitle=target_subtitle,
                    subheading_strict_match=boolish(params.get("subheading_strict_match", "true")),
                    hide_chapter_title=boolish(params.get("hide_chapter_title", "false")),
                    ignore_header_footer=boolish(params.get("ignore_header_footer", "true")),
                    ignore_toc=boolish(params.get("ignore_toc", "true")),
                    llm_boundary_fallback=(
                        boolish(params.get("llm_boundary_fallback", "false"))
                        if "llm_boundary_fallback" in params
                        else None
                    ),
                    llm_boundary_model_id=params.get("llm_boundary_model_id") or None,
                    output_docx_path=frag_path,
                    output_doc=None,
                    section=None
                )
                out_path = None
                if isinstance(result, dict):
                    log[-1]["captured_titles"] = result.get("captured_titles", [])
                    out_path = result.get("output_docx") or frag_path
                    log[-1]["output_docx"] = out_path
                if out_path:
                    _route_fragment(out_path, params, stype)

            elif stype == "extract_specific_figure_from_word":
                infile = params["input_file"]
                caption_label = (
                    params.get("target_caption_label")
                    or params.get("target_figure_label")
                    or params.get("target_table_label")
                    or ""
                )
                figure_title_raw = params.get("target_figure_title")
                figure_index_raw = params.get("target_figure_index")
                figure_title = str(figure_title_raw).strip() if figure_title_raw is not None else ""
                figure_index = str(figure_index_raw).strip() if figure_index_raw is not None else ""
                frag_path = _resolve_fragment_path(workdir, params.get("output_docx_path"), idx)
                include_caption = boolish(params.get("include_caption", "true"))
                result = extract_specific_figure_from_word(
                    infile,
                    params.get("target_chapter_section", ""),
                    caption_label,
                    target_subtitle=params.get("target_subtitle", params.get("subheading_text")),
                    target_chapter_title=params.get("target_chapter_title", params.get("target_title_section")),
                    target_figure_title=figure_title or None,
                    target_figure_index=figure_index or None,
                    output_docx_path=frag_path,
                    include_caption=include_caption,
                    ignore_header_footer=boolish(params.get("ignore_header_footer", "true")),
                    save_output=True,
                    return_reason=True,
                    output_doc=None,
                    section=None,
                )
                if isinstance(result, dict):
                    log[-1]["result"] = result
                if os.path.isfile(frag_path):
                    log[-1]["output_docx"] = frag_path
                    _route_fragment(frag_path, params, stype)

            elif stype == "extract_specific_table_from_word":
                infile = params["input_file"]
                caption_label = (
                    params.get("target_caption_label")
                    or params.get("target_table_label")
                    or params.get("target_figure_label")
                    or ""
                )
                table_title_raw = params.get("target_table_title")
                table_index_raw = params.get("target_table_index")
                table_title = str(table_title_raw).strip() if table_title_raw is not None else ""
                table_index = str(table_index_raw).strip() if table_index_raw is not None else ""
                frag_path = _resolve_fragment_path(workdir, params.get("output_docx_path"), idx)
                extract_specific_table_from_word(
                    infile,
                    frag_path,
                    params.get("target_chapter_section", ""),
                    caption_label,
                    params.get("target_subtitle", params.get("subheading_text")) or None,
                    target_chapter_title=params.get("target_chapter_title", params.get("target_title_section")),
                    target_table_title=table_title or None,
                    target_table_index=table_index or None,
                    include_caption=boolish(params.get("include_caption", "true")),
                    ignore_header_footer=boolish(params.get("ignore_header_footer", "true")),
                    save_output=True,
                )
                if os.path.isfile(frag_path):
                    _route_fragment(frag_path, params, stype)
                    log[-1]["output_docx"] = frag_path

            elif stype == "insert_text":
                frag_path = _resolve_fragment_path(workdir, params.get("output_docx_path"), idx)
                doc = _new_docx_fragment(frag_path)
                para = doc.add_paragraph(params.get("text",""))
                para.runs[0].bold = boolish(params.get("bold","false"))
                try:
                    para.runs[0].font.size = Pt(float(params.get("font_size", 12)))
                except Exception:
                    para.runs[0].font.size = None
                _set_alignment(para, params.get("align","left"))
                doc.save(frag_path)
                _route_fragment(frag_path, params, stype)

            elif stype == "insert_numbered_heading":
                frag_path = _resolve_fragment_path(workdir, params.get("output_docx_path"), idx)
                doc = _new_docx_fragment(frag_path)
                heading_text = params.get("text","")
                prefix = _next_arabic(params.get("level", 0))
                para = doc.add_paragraph(f"{prefix}{heading_text}")
                if "Normal" in doc.styles:
                    para.style = doc.styles["Normal"]
                _clear_list_formatting(para)
                _clear_indent(para)
                para.runs[0].bold = boolish(params.get("bold","true"))
                try:
                    para.runs[0].font.size = Pt(float(params.get("font_size", 12)))
                except Exception:
                    para.runs[0].font.size = None
                doc.save(frag_path)
                _route_fragment(frag_path, params, stype)

            elif stype == "insert_roman_heading":
                frag_path = _resolve_fragment_path(workdir, params.get("output_docx_path"), idx)
                doc = _new_docx_fragment(frag_path)
                heading_text = params.get("text","")
                prefix = _next_roman()
                para = doc.add_paragraph(f"{prefix}{heading_text}")
                if "Normal" in doc.styles:
                    para.style = doc.styles["Normal"]
                _clear_list_formatting(para)
                _clear_indent(para)
                para.runs[0].bold = boolish(params.get("bold","true"))
                try:
                    para.runs[0].font.size = Pt(float(params.get("font_size", 12)))
                except Exception:
                    para.runs[0].font.size = None
                doc.save(frag_path)
                _route_fragment(frag_path, params, stype)

            elif stype == "insert_bulleted_heading":
                frag_path = _resolve_fragment_path(workdir, params.get("output_docx_path"), idx)
                doc = _new_docx_fragment(frag_path)
                heading_text = params.get("text","")
                para = doc.add_paragraph(heading_text)
                if "List Bullet" in doc.styles:
                    para.style = doc.styles["List Bullet"]
                else:
                    para.text = f"• {heading_text}"
                _clear_indent(para)
                para.runs[0].bold = boolish(params.get("bold","true"))
                try:
                    para.runs[0].font.size = Pt(float(params.get("font_size", 12)))
                except Exception:
                    para.runs[0].font.size = None
                doc.save(frag_path)
                _route_fragment(frag_path, params, stype)

            elif stype == "copy_files":
                keywords = [k.strip() for k in params.get("keywords", "").split(",") if k.strip()]
                target_name = (params.get("target_name", "") or "").strip()
                source_path = params.get("source_dir", "")
                recursive_search = boolish(params.get("recursive_search", "true"))
                if os.path.isfile(source_path):
                    copied = [
                        copy_file(
                            source_path,
                            params.get("dest_dir", ""),
                            target_name=target_name or None,
                        )
                    ]
                    if keywords:
                        log[-1]["note"] = "已選擇單一來源檔案，已忽略關鍵字。"
                else:
                    copied = copy_files(
                        source_path,
                        params.get("dest_dir", ""),
                        keywords,
                        recursive=recursive_search,
                    )
                    if not copied:
                        log[-1]["copied_files"] = []
                        raise RuntimeError("未檢索到與關鍵字相符的檔案")
                    if target_name:
                        if len(copied) == 1:
                            copied = [_rename_single_copied_path(copied[0], target_name)]
                        elif copied:
                            log[-1]["note"] = "複製後名稱僅在實際複製 1 個檔案時生效；本次已忽略。"
                log[-1]["copied_files"] = copied
                if len(copied) == 1:
                    log[-1]["copied_file"] = copied[0]

            elif stype == "copy_directory":
                keywords = [k.strip() for k in params.get("keywords", "").split(",") if k.strip()]
                target_name = (params.get("target_name", "") or "").strip()
                recursive_search = boolish(params.get("recursive_search", "true"))
                if keywords:
                    copied_dirs = copy_directories(
                        params.get("source_dir", ""),
                        params.get("dest_dir", ""),
                        keywords,
                        recursive=recursive_search,
                        copied_registry=copied_dir_registry,
                        registry_entry_factory=lambda src_path: {"log_index": len(log) - 1, "source": os.path.abspath(src_path)},
                    )
                    if not copied_dirs:
                        log[-1]["copied_dirs"] = []
                        raise RuntimeError("未檢索到與關鍵字相符的資料夾")
                    if target_name:
                        if len(copied_dirs) == 1:
                            renamed = _rename_single_copied_path(copied_dirs[0], target_name)
                            existing_info = copied_dir_registry.pop(os.path.abspath(copied_dirs[0]), None)
                            copied_dirs = [renamed]
                            if existing_info is not None:
                                copied_dir_registry[os.path.abspath(renamed)] = existing_info
                        elif copied_dirs:
                            log[-1]["note"] = "複製後名稱僅在實際複製 1 個資料夾時生效；本次已忽略。"
                else:
                    copied_dir = copy_directory(
                        params.get("source_dir", ""),
                        params.get("dest_dir", ""),
                        target_name=target_name or None,
                        copied_registry=copied_dir_registry,
                        registry_entry={"log_index": len(log) - 1, "source": os.path.abspath(params.get("source_dir", ""))},
                    )
                    copied_dirs = [copied_dir]
                log[-1]["copied_dirs"] = copied_dirs
                if len(copied_dirs) == 1:
                    log[-1]["copied_dir"] = copied_dirs[0]

            elif stype == "renumber_figures_tables":
                # Skipped here to avoid Spire save (watermark); can be run externally if licensed.
                log[-1]["status"] = "skipped"
                log[-1]["note"] = "renumber_figures_tables skipped to avoid Spire watermark"
                continue

            else:
                raise RuntimeError(f"Unknown step type: {stype}")

            if "status" not in log[-1]:
                log[-1]["status"] = "ok"
            _check_canceled()

        except JobCanceledError:
            raise
        except Exception as e:
            log[-1]["status"] = "error"
            log[-1]["error"] = str(e)

    # Post-check: ensure extract steps actually produced content.
    for entry in log:
        _check_canceled()
        if not isinstance(entry, dict):
            continue
        if "step" not in entry:
            continue
        if entry.get("status") == "error":
            continue
        stype = entry.get("type")
        if stype in ("extract_word_chapter", "extract_word_all_content"):
            out_path = entry.get("output_docx")
            if not out_path or not _docx_has_content(out_path):
                entry["status"] = "error"
                entry["error"] = "No content extracted"
        elif stype == "extract_pdf_pages_as_images":
            out_path = entry.get("output_docx")
            if not out_path or not _docx_has_content(out_path):
                entry["status"] = "error"
                entry["error"] = "No PDF pages extracted"
        elif stype == "extract_specific_table_from_word":
            out_path = entry.get("output_docx")
            if not out_path or not _docx_has_content(out_path):
                entry["status"] = "error"
                entry["error"] = "Table not found"
        elif stype == "extract_specific_figure_from_word":
            out_path = entry.get("output_docx")
            if not out_path or not _docx_has_content(out_path):
                entry["status"] = "error"
                reason = ""
                result = entry.get("result")
                if isinstance(result, dict):
                    reason = str(result.get("reason") or "").strip()
                entry["error"] = reason or "Figure not found"

    if template_cfg.get("path") and template_mappings:
        try:
            _check_canceled()
            parsed = template_paragraphs or parse_template_paragraphs(template_cfg["path"])
            template_result_path = os.path.join(workdir, "template_result.docx")
            ordered_template_mappings = order_template_mappings(template_mappings)
            render_template_with_mappings(
                template_cfg["path"],
                template_result_path,
                ordered_template_mappings,
                parsed,
            )
            log.append(
                {
                    "step": len(log) + 1,
                    "type": "template_merge",
                    "template_file": template_cfg["path"],
                    "mappings": [
                        {
                            "index": m.get("index"),
                            "mode": m.get("mode"),
                            "source_step": m.get("source_step"),
                            "source_id": m.get("source_id"),
                        }
                        for m in ordered_template_mappings
                    ],
                    "output_docx": template_result_path,
                    "status": "ok",
                }
            )
            fragments.insert(0, template_result_path)
            template_merge_succeeded = True
        except JobCanceledError:
            raise
        except Exception as e:
            log.append(
                {
                    "step": len(log) + 1,
                    "type": "template_merge",
                    "template_file": template_cfg.get("path"),
                    "mappings": template_mappings,
                    "status": "error",
                    "error": str(e),
                }
            )
            for mp in template_mappings:
                cdoc = mp.get("content_docx_path")
                if cdoc:
                    fragments.append(cdoc)

    if not fragments:
        _check_canceled()
        os.makedirs(workdir, exist_ok=True)
        empty_path = os.path.join(workdir, "result.docx")
        DocxDocument().save(empty_path)
        fragments.append(empty_path)

    for final_path, info in copied_dir_registry.items():
        if not isinstance(info, dict):
            continue
        log_index = info.get("log_index")
        if isinstance(log_index, int) and 0 <= log_index < len(log):
            log[log_index]["copied_dir"] = final_path

    out_docx = os.path.join(workdir, "result.docx")
    _check_canceled()
    merge_word_docs(fragments, out_docx)
    final_template_source_ids = [
        str(item.get("source_id") or "")
        for item in display_order_template_mappings(template_mappings)
        if str(item.get("source_id") or "")
    ] if template_merge_succeeded else []
    final_fragment_source_ids = [
        str((entry.get("provenance") or {}).get("source_id") or "")
        for entry in log
        if isinstance(entry, dict)
        and isinstance(entry.get("provenance"), dict)
        and (entry.get("provenance") or {}).get("template_index") in (None, "", "None")
        and str((entry.get("provenance") or {}).get("source_id") or "")
    ]
    final_source_sequence = final_template_source_ids + final_fragment_source_ids
    provenance_by_source_id = {
        str((entry.get("provenance") or {}).get("source_id") or ""): dict(entry.get("provenance") or {})
        for entry in log
        if isinstance(entry, dict)
        and isinstance(entry.get("provenance"), dict)
        and str((entry.get("provenance") or {}).get("source_id") or "")
    }
    if provenance_by_source_id and final_source_sequence:
        _check_canceled()
        applied_ranges = apply_final_provenance(
            out_docx,
            [provenance_by_source_id[source_id] for source_id in final_source_sequence if source_id in provenance_by_source_id],
        )
        range_by_source_id = {
            str(item.get("source_id") or ""): item
            for item in applied_ranges
            if str(item.get("source_id") or "")
        }
        for entry in log:
            provenance = entry.get("provenance")
            if not isinstance(provenance, dict):
                continue
            source_id = str(provenance.get("source_id") or "")
            applied = range_by_source_id.get(source_id)
            if not applied:
                continue
            provenance["bookmark_start"] = applied.get("bookmark_start", provenance.get("bookmark_start"))
            provenance["bookmark_end"] = applied.get("bookmark_end", provenance.get("bookmark_end"))
            provenance["bookmark_id"] = applied.get("bookmark_id", provenance.get("bookmark_id"))
            provenance["result_block_start"] = applied.get("result_block_start")
            provenance["result_block_end"] = applied.get("result_block_end")

    out_log = os.path.join(workdir, "log.json")
    _check_canceled()
    with open(out_log, "w", encoding="utf-8") as f:
        import json
        json.dump(log, f, ensure_ascii=False, indent=2)

    return {"result_docx": out_docx, "log": out_log, "log_json": log}
