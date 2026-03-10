from __future__ import annotations

import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

from modules.docx_merger import merge_word_docs


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _create_table_only_docx(path: Path, cell_text: str) -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = cell_text

    # Make the fragment "table-only" to mimic extracted table fragments.
    for p in list(doc.paragraphs):
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
    doc.save(path)


def _body_block_types(docx_path: Path) -> list[str]:
    with zipfile.ZipFile(docx_path, "r") as zf:
        xml_bytes = zf.read("word/document.xml")
    root = ET.fromstring(xml_bytes)
    body = root.find(f".//{{{W_NS}}}body")
    if body is None:
        return []
    tags: list[str] = []
    for child in list(body):
        local = child.tag.split("}", 1)[-1]
        if local in {"p", "tbl"}:
            tags.append(local)
    return tags


def test_merge_word_docs_inserts_separator_between_adjacent_tables(tmp_path: Path) -> None:
    doc1 = tmp_path / "t1.docx"
    doc2 = tmp_path / "t2.docx"
    out = tmp_path / "merged.docx"
    _create_table_only_docx(doc1, "A")
    _create_table_only_docx(doc2, "B")

    merge_word_docs([str(doc1), str(doc2)], str(out))

    tags = _body_block_types(out)
    first_tbl = tags.index("tbl")
    second_tbl = tags.index("tbl", first_tbl + 1)
    assert "p" in tags[first_tbl + 1 : second_tbl]

