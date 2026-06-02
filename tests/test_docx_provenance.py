from pathlib import Path
import base64
from zipfile import ZipFile

from docx import Document as DocxDocument
from docx.shared import Inches
from lxml import etree

from modules.docx_merger import merge_word_docs
from modules.docx_provenance import (
    _installed_font_families,
    annotate_docx_with_provenance,
    _resolve_linux_east_asia_font,
    apply_final_provenance,
    build_provenance_descriptor,
    copy_docx_with_preview_fonts,
    create_provenance_preview_docx,
    extract_provenance_block_trace,
)
from modules.template_manager import render_template_with_mappings
from modules.workflow import run_workflow


_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


def _create_docx(path: Path, paragraphs: list[str] | None = None) -> None:
    doc = DocxDocument()
    for text in paragraphs or []:
        doc.add_paragraph(text)
    doc.save(path)


def _convert_first_inline_picture_to_anchor(path: Path) -> None:
    with ZipFile(path, "r") as zf:
        parts = {name: zf.read(name) for name in zf.namelist()}
    root = etree.fromstring(parts["word/document.xml"])
    inline = root.xpath("//wp:inline", namespaces=_NS)[0]
    inline.tag = "{%s}anchor" % _NS["wp"]
    simple_pos = etree.Element("{%s}simplePos" % _NS["wp"])
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    position_h = etree.Element("{%s}positionH" % _NS["wp"])
    position_h.set("relativeFrom", "column")
    pos_offset_h = etree.SubElement(position_h, "{%s}posOffset" % _NS["wp"])
    pos_offset_h.text = "0"
    position_v = etree.Element("{%s}positionV" % _NS["wp"])
    position_v.set("relativeFrom", "paragraph")
    pos_offset_v = etree.SubElement(position_v, "{%s}posOffset" % _NS["wp"])
    pos_offset_v.text = "0"
    wrap_none = etree.Element("{%s}wrapNone" % _NS["wp"])
    inline.insert(0, wrap_none)
    inline.insert(0, position_v)
    inline.insert(0, position_h)
    inline.insert(0, simple_pos)
    parts["word/document.xml"] = etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )
    with ZipFile(path, "w") as zf:
        for name, data in parts.items():
            zf.writestr(name, data)


def _create_table_docx(path: Path, rows: list[list[str]]) -> None:
    doc = DocxDocument()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for row_idx, row in enumerate(rows):
        for col_idx, text in enumerate(row):
            table.cell(row_idx, col_idx).text = text
    doc.save(path)


def _collect_provenance_pairs(path: Path) -> tuple[dict[str, str], set[str]]:
    with ZipFile(path, "r") as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
    starts = {
        str(node.get("{%s}id" % _NS["w"]) or ""): str(node.get("{%s}name" % _NS["w"]) or "")
        for node in root.xpath("//w:bookmarkStart", namespaces=_NS)
        if str(node.get("{%s}name" % _NS["w"]) or "").startswith("prov_src_")
    }
    ends = {
        str(node.get("{%s}id" % _NS["w"]) or "")
        for node in root.xpath("//w:bookmarkEnd", namespaces=_NS)
    }
    return starts, ends


def test_extract_provenance_block_trace_survives_docx_merge(tmp_path: Path) -> None:
    first_fragment = tmp_path / "fragment_a.docx"
    second_fragment = tmp_path / "fragment_b.docx"
    result_path = tmp_path / "result.docx"

    _create_docx(first_fragment, ["Alpha source paragraph"])
    _create_table_docx(
        second_fragment,
        [["Distinct table row content for provenance tracking", "Second cell content"]],
    )

    first_desc = build_provenance_descriptor(1)
    second_desc = build_provenance_descriptor(2)
    merge_word_docs([str(first_fragment), str(second_fragment)], str(result_path))

    applied = apply_final_provenance(
        str(result_path),
        [
            {
                **first_desc,
                "fragment_path": str(first_fragment),
                "content_type": "paragraph",
                "source_id": "src_000001",
            },
            {
                **second_desc,
                "fragment_path": str(second_fragment),
                "content_type": "table",
                "source_id": "src_000002",
            },
        ],
    )
    assert len(applied) == 2

    starts, ends = _collect_provenance_pairs(result_path)
    assert starts
    assert set(starts).issubset(ends)

    trace = extract_provenance_block_trace(
        str(result_path),
        {
            "src_000001": {
                **first_desc,
                "source_file": "A.docx",
                "source_step": "extract_word_chapter",
                "content_type": "paragraph",
            },
            "src_000002": {
                **second_desc,
                "source_file": "B.docx",
                "source_step": "extract_specific_table_from_word",
                "content_type": "table",
            },
        },
    )

    assert any(
        item["block_type"] == "paragraph"
        and item["source_file"] == "A.docx"
        and item["text"] == "Alpha source paragraph"
        for item in trace
    )
    assert any(
        item["block_type"] == "table"
        and item["source_file"] == "B.docx"
        and "Distinct table row content for provenance tracking" in item["probe_texts"]
        for item in trace
    )


def test_extract_provenance_block_trace_survives_template_merge_subdoc(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    content_path = tmp_path / "content.docx"
    output_path = tmp_path / "result.docx"

    _create_docx(template_path, ["Anchor"])
    _create_docx(content_path, ["Inserted provenance paragraph"])

    desc = build_provenance_descriptor(1)
    render_template_with_mappings(
        str(template_path),
        str(output_path),
        [
            {
                "index": 0,
                "mode": "insert_after",
                "content_docx_path": str(content_path),
                "source_order": 0,
            }
        ],
        [{"index": 0, "display": "", "text": "Anchor"}],
    )

    applied = apply_final_provenance(
        str(output_path),
        [
            {
                **desc,
                "fragment_path": str(content_path),
                "content_type": "paragraph",
                "source_id": "src_000001",
            }
        ],
    )
    assert len(applied) == 1

    trace = extract_provenance_block_trace(
        str(output_path),
        {
            "src_000001": {
                **desc,
                "source_file": "Template Source.docx",
                "source_step": "extract_word_all_content",
                "content_type": "paragraph",
            }
        },
    )

    assert any(
        item["block_type"] == "paragraph"
        and item["source_file"] == "Template Source.docx"
        and item["text"] == "Inserted provenance paragraph"
        for item in trace
    )


def test_extract_provenance_block_trace_uses_metadata_for_empty_figure_paragraph(tmp_path: Path) -> None:
    figure_docx = tmp_path / "figure.docx"
    image_path = tmp_path / "pixel.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9sX9n0cAAAAASUVORK5CYII="
        )
    )
    doc = DocxDocument()
    doc.add_paragraph().add_run().add_picture(str(image_path), width=Inches(0.2))
    doc.save(figure_docx)

    desc = build_provenance_descriptor(1)
    apply_final_provenance(
        str(figure_docx),
        [
            {
                **desc,
                "fragment_path": str(figure_docx),
                "content_type": "figure",
                "source_id": "src_000001",
                "primary_probe_texts": ["Figure 8 Packaging"],
            }
        ],
    )

    trace = extract_provenance_block_trace(
        str(figure_docx),
        {
            "src_000001": {
                **desc,
                "source_file": "Figure Source.docx",
                "source_step": "extract_specific_figure_from_word",
                "content_type": "figure",
                "primary_probe_texts": ["Figure 8 Packaging"],
            }
        },
    )

    assert any(
        item["block_type"] == "figure"
        and item["source_file"] == "Figure Source.docx"
        and "Figure 8 Packaging" in item["probe_texts"]
        for item in trace
    )


def test_create_provenance_preview_docx_inlines_floating_figure_images(tmp_path: Path) -> None:
    result_path = tmp_path / "result.docx"
    preview_path = tmp_path / "preview.docx"
    image_path = tmp_path / "pixel.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9sX9n0cAAAAASUVORK5CYII="
        )
    )
    doc = DocxDocument()
    doc.add_paragraph().add_run().add_picture(str(image_path), width=Inches(0.2))
    doc.save(result_path)
    _convert_first_inline_picture_to_anchor(result_path)

    desc = build_provenance_descriptor(1)
    apply_final_provenance(
        str(result_path),
        [
            {
                **desc,
                "fragment_path": str(result_path),
                "content_type": "figure",
                "source_id": "src_000001",
            }
        ],
    )

    created = create_provenance_preview_docx(
        str(result_path),
        str(preview_path),
        {
            "src_000001": {
                **desc,
                "source_file": "Figure Source.docx",
                "source_step": "extract_specific_figure_from_word",
                "content_type": "figure",
            }
        },
    )

    assert created is True
    with ZipFile(preview_path, "r") as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
    assert not root.xpath("//wp:anchor", namespaces=_NS)
    assert root.xpath("//wp:inline", namespaces=_NS)
    paragraph_texts = [
        paragraph.text.strip()
        for paragraph in DocxDocument(str(preview_path)).paragraphs
        if paragraph.text.strip()
    ]
    assert paragraph_texts[0] == "(此說明不會出現在實際輸出文件)來源: Figure Source.docx"


def test_create_provenance_preview_docx_skips_blank_bookmark_start_paragraph(tmp_path: Path) -> None:
    result_path = tmp_path / "result.docx"
    preview_path = tmp_path / "preview.docx"
    doc = DocxDocument()
    doc.add_paragraph("")
    doc.add_paragraph("Real source paragraph")
    doc.save(result_path)

    desc = build_provenance_descriptor(1)
    assert annotate_docx_with_provenance(
        str(result_path),
        bookmark_start=str(desc["bookmark_start"]),
        bookmark_end=str(desc["bookmark_end"]),
        bookmark_id=int(desc["bookmark_id"]),
    )

    created = create_provenance_preview_docx(
        str(result_path),
        str(preview_path),
        {
            "src_000001": {
                **desc,
                "source_file": "Alpha.docx",
                "source_step": "extract_word_chapter",
                "content_type": "paragraph",
            }
        },
    )

    assert created is True
    preview_doc = DocxDocument(str(preview_path))
    assert preview_doc.paragraphs[0].text == ""
    assert preview_doc.paragraphs[1].text == "(此說明不會出現在實際輸出文件)來源: Alpha.docx"
    assert preview_doc.paragraphs[2].text == "Real source paragraph"


def test_create_provenance_preview_docx_uses_result_block_ranges_over_open_bookmarks(tmp_path: Path) -> None:
    result_path = tmp_path / "result.docx"
    preview_path = tmp_path / "preview.docx"
    doc = DocxDocument()
    doc.add_paragraph("Figure A")
    doc.add_paragraph("Caption A")
    doc.add_paragraph("Figure B")
    doc.add_paragraph("Caption B")
    doc.add_paragraph("Template prompt must not inherit source A")
    doc.save(result_path)

    first_desc = build_provenance_descriptor(1)
    second_desc = build_provenance_descriptor(2)
    created = create_provenance_preview_docx(
        str(result_path),
        str(preview_path),
        {
            "src_000001": {
                **first_desc,
                "source_file": "A.docx",
                "content_type": "figure",
                "result_block_start": 0,
                "result_block_end": 1,
            },
            "src_000002": {
                **second_desc,
                "source_file": "B.docx",
                "content_type": "figure",
                "result_block_start": 2,
                "result_block_end": 3,
            },
        },
    )

    assert created is True
    paragraph_texts = [paragraph.text for paragraph in DocxDocument(str(preview_path)).paragraphs]
    assert paragraph_texts == [
        "(此說明不會出現在實際輸出文件)來源: A.docx",
        "Figure A",
        "Caption A",
        "(此說明不會出現在實際輸出文件)來源: B.docx",
        "Figure B",
        "Caption B",
        "Template prompt must not inherit source A",
    ]


def test_create_provenance_preview_docx_inserts_labels_without_highlighting_body_text(
    monkeypatch,
    tmp_path: Path,
) -> None:
    _installed_font_families.cache_clear()
    monkeypatch.setattr(
        "modules.docx_provenance._installed_font_families",
        lambda: {"calibri", "微軟正黑體"},
    )
    first_fragment = tmp_path / "fragment_a.docx"
    second_fragment = tmp_path / "fragment_b.docx"
    result_path = tmp_path / "result.docx"
    preview_path = tmp_path / "preview.docx"

    _create_docx(first_fragment, ["Alpha source paragraph"])
    _create_docx(second_fragment, ["Beta source paragraph"])

    first_desc = build_provenance_descriptor(1)
    second_desc = build_provenance_descriptor(2)
    merge_word_docs([str(first_fragment), str(second_fragment)], str(result_path))
    apply_final_provenance(
        str(result_path),
        [
            {
                **first_desc,
                "fragment_path": str(first_fragment),
                "content_type": "paragraph",
                "source_id": "src_000001",
            },
            {
                **second_desc,
                "fragment_path": str(second_fragment),
                "content_type": "paragraph",
                "source_id": "src_000002",
            },
        ],
    )

    created = create_provenance_preview_docx(
        str(result_path),
        str(preview_path),
        {
            "src_000001": {
                **first_desc,
                "source_file": "Alpha.docx",
                "source_step": "extract_word_chapter",
                "content_type": "paragraph",
            },
            "src_000002": {
                **second_desc,
                "source_file": "Beta.docx",
                "source_step": "extract_word_chapter",
                "content_type": "paragraph",
            },
        },
    )

    assert created is True

    preview_doc = DocxDocument(str(preview_path))
    paragraph_texts = [paragraph.text.strip() for paragraph in preview_doc.paragraphs if paragraph.text.strip()]

    assert "(此說明不會出現在實際輸出文件)來源: Alpha.docx" in paragraph_texts
    assert "(此說明不會出現在實際輸出文件)來源: Beta.docx" in paragraph_texts
    assert "Alpha source paragraph" in paragraph_texts
    assert "Beta source paragraph" in paragraph_texts
    assert paragraph_texts.index(
        "(此說明不會出現在實際輸出文件)來源: Alpha.docx"
    ) < paragraph_texts.index("Alpha source paragraph")
    assert paragraph_texts.index(
        "(此說明不會出現在實際輸出文件)來源: Beta.docx"
    ) < paragraph_texts.index("Beta source paragraph")

    with ZipFile(preview_path, "r") as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
    highlight_paragraphs = root.xpath("//w:p[w:r/w:rPr/w:highlight]", namespaces=_NS)
    highlight_texts = ["".join(paragraph.xpath(".//w:t/text()", namespaces=_NS)) for paragraph in highlight_paragraphs]
    assert not any("Alpha source paragraph" in text for text in highlight_texts)
    assert not any("Beta source paragraph" in text for text in highlight_texts)
    label_colors = root.xpath(
        "//w:p[contains(string(.), '來源: ')]//w:color/@w:val",
        namespaces=_NS,
    )
    assert "C00000" in label_colors
    label_fonts = root.xpath(
        "//w:p[contains(string(.), '來源: ')]//w:rFonts/@w:eastAsia",
        namespaces=_NS,
    )
    assert "微軟正黑體" in label_fonts
    complex_script_fonts = root.xpath(
        "//w:p[contains(string(.), '來源: ')]//w:rFonts/@w:cs",
        namespaces=_NS,
    )
    assert "微軟正黑體" in complex_script_fonts
    font_hints = root.xpath(
        "//w:p[contains(string(.), '來源: ')]//w:rFonts/@w:hint",
        namespaces=_NS,
    )
    assert "eastAsia" in font_hints
    east_asia_langs = root.xpath(
        "//w:p[contains(string(.), '來源: ')]//w:lang/@w:eastAsia",
        namespaces=_NS,
    )
    assert "zh-TW" in east_asia_langs


def test_create_provenance_preview_docx_uses_configured_linux_friendly_font(
    monkeypatch,
    tmp_path: Path,
) -> None:
    _installed_font_families.cache_clear()
    monkeypatch.setattr(
        "modules.docx_provenance._installed_font_families",
        lambda: {"calibri", "noto sans cjk tc"},
    )
    first_fragment = tmp_path / "fragment_a.docx"
    result_path = tmp_path / "result.docx"
    preview_path = tmp_path / "preview.docx"

    _create_docx(first_fragment, ["Alpha source paragraph"])

    first_desc = build_provenance_descriptor(1)
    merge_word_docs([str(first_fragment)], str(result_path))
    apply_final_provenance(
        str(result_path),
        [
            {
                **first_desc,
                "fragment_path": str(first_fragment),
                "content_type": "paragraph",
                "source_id": "src_000001",
            },
        ],
    )

    monkeypatch.setenv("PROVENANCE_PREVIEW_LABEL_EAST_ASIA_FONT", "Noto Sans CJK TC")

    created = create_provenance_preview_docx(
        str(result_path),
        str(preview_path),
        {
            "src_000001": {
                **first_desc,
                "source_file": "中文檔名.docx",
                "source_step": "extract_word_chapter",
                "content_type": "paragraph",
            },
        },
    )

    assert created is True

    with ZipFile(preview_path, "r") as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
    label_fonts = root.xpath(
        "//w:p[contains(string(.), '來源: 中文檔名.docx')]//w:rFonts/@w:eastAsia",
        namespaces=_NS,
    )
    assert "Noto Sans CJK TC" in label_fonts


def test_resolve_linux_east_asia_font_falls_back_to_installed_cjk_font(monkeypatch) -> None:
    _installed_font_families.cache_clear()
    monkeypatch.setattr(
        "modules.docx_provenance._installed_font_families",
        lambda: {"wenquanyi zen hei"},
    )

    resolved = _resolve_linux_east_asia_font("Noto Sans CJK TC")

    assert resolved == "WenQuanYi Zen Hei"


def test_resolve_linux_east_asia_font_ignores_non_cjk_font_fallbacks(monkeypatch) -> None:
    _installed_font_families.cache_clear()
    monkeypatch.setattr(
        "modules.docx_provenance._installed_font_families",
        lambda: {"noto sans", "noto serif"},
    )

    resolved = _resolve_linux_east_asia_font("Noto Sans CJK TC")

    assert resolved == ""


def test_copy_docx_with_preview_fonts_applies_east_asia_font_to_runs_and_defaults(tmp_path: Path) -> None:
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "preview.docx"

    doc = DocxDocument()
    doc.add_paragraph("中文檔名來源")
    doc.save(source_path)

    copied = copy_docx_with_preview_fonts(str(source_path), str(output_path), east_asia_font="Noto Sans CJK TC")

    assert copied is True

    with ZipFile(output_path, "r") as zf:
        document_root = etree.fromstring(zf.read("word/document.xml"))
        styles_root = etree.fromstring(zf.read("word/styles.xml"))

    run_fonts = document_root.xpath("//w:rPr/w:rFonts/@w:eastAsia", namespaces=_NS)
    assert "Noto Sans CJK TC" in run_fonts
    hints = document_root.xpath("//w:rPr/w:rFonts/@w:hint", namespaces=_NS)
    assert "eastAsia" in hints
    langs = document_root.xpath("//w:rPr/w:lang/@w:eastAsia", namespaces=_NS)
    assert "zh-TW" in langs
    default_fonts = styles_root.xpath(
        "//w:docDefaults/w:rPrDefault/w:rPr/w:rFonts/@w:eastAsia",
        namespaces=_NS,
    )
    assert "Noto Sans CJK TC" in default_fonts


def test_apply_final_provenance_covers_all_merged_fragments(tmp_path: Path) -> None:
    fragments: list[Path] = []
    source_records = []
    for idx in range(1, 9):
        fragment = tmp_path / f"fragment_{idx}.docx"
        _create_docx(fragment, [f"Fragment {idx}"])
        desc = build_provenance_descriptor(idx)
        fragments.append(fragment)
        source_records.append(
            {
                **desc,
                "fragment_path": str(fragment),
                "content_type": "paragraph",
                "source_id": f"src_{idx:06d}",
            }
        )

    result_path = tmp_path / "result.docx"
    merge_word_docs([str(path) for path in fragments], str(result_path))
    applied = apply_final_provenance(str(result_path), source_records)

    starts, ends = _collect_provenance_pairs(result_path)
    assert len(applied) == 8
    assert len(starts) == 8
    assert set(starts).issubset(ends)


def test_apply_final_provenance_inserts_exact_range_on_final_docx(tmp_path: Path) -> None:
    fragment = tmp_path / "fragment.docx"
    result_path = tmp_path / "result.docx"
    _create_docx(fragment, ["First paragraph", "Second paragraph"])
    _create_docx(result_path, ["Prefix", "First paragraph", "Second paragraph", "Suffix"])

    desc = build_provenance_descriptor(1)
    applied = apply_final_provenance(
        str(result_path),
        [
            {
                **desc,
                "fragment_path": str(fragment),
                "content_type": "paragraph",
                "source_id": "src_000001",
            }
        ],
    )

    assert len(applied) == 1
    assert applied[0]["result_block_start"] == 1
    assert applied[0]["result_block_end"] == 2

    starts, ends = _collect_provenance_pairs(result_path)
    assert starts
    assert set(starts).issubset(ends)


def test_run_workflow_applies_template_provenance_in_display_order(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    job_dir = tmp_path / "job_template_order"

    _create_docx(
        template_path,
        [
            "Anchor 0",
            "Anchor 1",
            "Anchor 2",
            "Anchor 3",
        ],
    )

    steps = []
    for idx in range(4):
        steps.append(
            {
                "type": "insert_text",
                "params": {
                    "text": f"Section {idx} - First",
                    "template_index": str(idx),
                    "template_mode": "insert_after",
                },
            }
        )
        steps.append(
            {
                "type": "insert_text",
                "params": {
                    "text": f"Section {idx} - Second",
                    "template_index": str(idx),
                    "template_mode": "insert_after",
                },
            }
        )

    result = run_workflow(steps, str(job_dir), template={"path": str(template_path)})
    provenance_entries = [
        entry["provenance"]
        for entry in result["log_json"]
        if isinstance(entry, dict) and isinstance(entry.get("provenance"), dict)
    ]

    assert len(provenance_entries) == 8
    assert all(entry.get("result_block_start") is not None for entry in provenance_entries)
    assert all(entry.get("result_block_end") is not None for entry in provenance_entries)

    starts, ends = _collect_provenance_pairs(Path(result["result_docx"]))
    assert len(starts) == 8
    assert set(starts).issubset(ends)
