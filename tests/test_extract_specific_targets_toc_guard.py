from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Inches

import modules.extract_specific_figure_xml as figure_xml
import modules.extract_specific_table_xml as table_xml
from modules.docx_toc import TocEntry


def _create_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("Device Under Evaluation")
    doc.add_paragraph("Body text")
    doc.save(path)


def _create_figure_docx_with_long_caption(path: Path, image_path: Path) -> None:
    image_path.write_bytes(
        bytes.fromhex(
            "89504E470D0A1A0A0000000D4948445200000001000000010804000000B51C0C02"
            "0000000B4944415478DA63FCFF1F0002EB01F6C5FD9F470000000049454E44AE426082"
        )
    )
    doc = DocxDocument()
    doc.add_paragraph("Device Under Evaluation")
    doc.add_paragraph().add_run().add_picture(str(image_path), width=Inches(0.2))
    doc.add_paragraph(
        "Figure 1. The overlapping graph of predicate device (K082424) and the current "
        "submission device. The above figure shows the locking mechanism (Taper at outer "
        "surface, Thread in inner surface) of current submission (in blue) and predicate "
        "device (in black) are identical."
    )
    doc.save(path)


def _create_table_figure_docx(path: Path, image_path: Path) -> None:
    image_path.write_bytes(
        bytes.fromhex(
            "89504E470D0A1A0A0000000D4948445200000001000000010804000000B51C0C02"
            "0000000B4944415478DA63FCFF1F0002EB01F6C5FD9F470000000049454E44AE426082"
        )
    )
    doc = DocxDocument()
    doc.add_paragraph("Packaging for the Implant")
    tbl = doc.add_table(rows=1, cols=2)
    for idx, cell in enumerate(tbl.rows[0].cells, start=1):
        cell.paragraphs[0].add_run().add_picture(str(image_path), width=Inches(0.2))
        cell.add_paragraph(f"View {idx}")
    doc.add_paragraph("Figure 2 Packaging of U2 Total Knee System PSA type extension line")
    doc.save(path)


def test_extract_specific_figure_uses_toc_backed_mismatch_guard(monkeypatch, tmp_path: Path) -> None:
    src = tmp_path / "figure_source.docx"
    _create_docx(src)
    captured: dict[str, object] = {}

    def fake_find_section_range_children(*args, **kwargs):
        captured["strict_heading_number_match"] = kwargs.get("strict_heading_number_match")
        captured["allow_start_number_mismatch_fallback"] = kwargs.get(
            "allow_start_number_mismatch_fallback"
        )
        captured["rule_based_boundary_fallback"] = kwargs.get("rule_based_boundary_fallback")
        captured["style_numpr"] = kwargs.get("style_numpr")
        return 0, len(args[0])

    def fake_extract_toc_entries_from_parts(_parts):
        return [
            TocEntry(
                order=1,
                level=3,
                number="3.2.1",
                title="Device Under Evaluation",
                page="11",
                anchor="_Toc1",
                raw_text="3.2.1 Device Under Evaluation 11",
                style_id="33",
                style_name="toc 3",
            )
        ]

    monkeypatch.setattr(figure_xml, "find_section_range_children", fake_find_section_range_children)
    monkeypatch.setattr(
        "modules.docx_toc.extract_toc_entries_from_parts",
        fake_extract_toc_entries_from_parts,
    )

    result = figure_xml.extract_specific_figure_from_word_xml(
        input_file=str(src),
        output_docx_path=None,
        target_chapter_section="3.2.1",
        target_caption_label="Figure 1.",
        target_chapter_title="Device Under Evaluation",
        save_output=False,
        return_reason=True,
    )

    assert result["ok"] is False
    assert captured["strict_heading_number_match"] is True
    assert captured["allow_start_number_mismatch_fallback"] is True
    assert isinstance(captured["style_numpr"], dict)


def test_extract_specific_table_disables_mismatch_guard_without_toc_match(
    monkeypatch, tmp_path: Path
) -> None:
    src = tmp_path / "table_source.docx"
    _create_docx(src)
    captured: dict[str, object] = {}

    def fake_find_section_range_children(*args, **kwargs):
        captured["strict_heading_number_match"] = kwargs.get("strict_heading_number_match")
        captured["allow_start_number_mismatch_fallback"] = kwargs.get(
            "allow_start_number_mismatch_fallback"
        )
        captured["rule_based_boundary_fallback"] = kwargs.get("rule_based_boundary_fallback")
        captured["style_numpr"] = kwargs.get("style_numpr")
        return 0, len(args[0])

    def fake_extract_toc_entries_from_parts(_parts):
        return [
            TocEntry(
                order=1,
                level=3,
                number="3.2.1",
                title="Device Under Evaluation",
                page="11",
                anchor="_Toc1",
                raw_text="3.2.1 Device Under Evaluation 11",
                style_id="33",
                style_name="toc 3",
            )
        ]

    monkeypatch.setattr(table_xml, "find_section_range_children", fake_find_section_range_children)
    monkeypatch.setattr(
        "modules.docx_toc.extract_toc_entries_from_parts",
        fake_extract_toc_entries_from_parts,
    )

    result = table_xml.extract_specific_table_from_word_xml(
        input_file=str(src),
        output_docx_path=None,
        target_chapter_section="9.9.9",
        target_caption_label="Table 1.",
        target_chapter_title="Device Under Evaluation",
        save_output=False,
        return_reason=True,
    )

    assert result["ok"] is False
    assert captured["strict_heading_number_match"] is True
    assert captured["allow_start_number_mismatch_fallback"] is False
    assert captured["rule_based_boundary_fallback"] is False
    assert isinstance(captured["style_numpr"], dict)


def test_extract_specific_figure_index_keeps_long_figure_prefix_caption(tmp_path: Path) -> None:
    src = tmp_path / "figure_long_caption.docx"
    image_path = tmp_path / "pixel.png"
    out = tmp_path / "figure_out.docx"
    _create_figure_docx_with_long_caption(src, image_path)

    result = figure_xml.extract_specific_figure_from_word_xml(
        input_file=str(src),
        output_docx_path=str(out),
        target_chapter_section="",
        target_caption_label="",
        target_figure_index=1,
        save_output=True,
        return_reason=True,
    )

    assert result["ok"] is True
    assert result["match_mode"] == "figure_index"
    assert result["selected_caption_analysis"]["accepted"] is True
    assert result["selected_caption_analysis"]["reason"] == "figure_number_prefix"

    out_doc = DocxDocument(out)
    lines = [p.text.strip() for p in out_doc.paragraphs if p.text.strip()]
    assert any(line.startswith("Figure 1. The overlapping graph") for line in lines)


def test_extract_specific_figure_detects_table_wrapped_image_with_caption_below(tmp_path: Path) -> None:
    src = tmp_path / "table_figure.docx"
    image_path = tmp_path / "pixel_table.png"
    out = tmp_path / "table_figure_out.docx"
    _create_table_figure_docx(src, image_path)

    result = figure_xml.extract_specific_figure_from_word_xml(
        input_file=str(src),
        output_docx_path=str(out),
        target_chapter_section="",
        target_caption_label="Figure 2",
        allow_table_figure_container=True,
        save_output=True,
        return_reason=True,
    )

    assert result["ok"] is True
    assert result["match_mode"] == "caption"
    assert result["selected_caption_text"].startswith("Figure 2 Packaging")

    out_doc = DocxDocument(out)
    assert len(out_doc.tables) == 1
    lines = [p.text.strip() for p in out_doc.paragraphs if p.text.strip()]
    assert any(line.startswith("Figure 2 Packaging") for line in lines)


def test_extract_specific_figure_index_accepts_long_caption_without_space_after_number(tmp_path: Path) -> None:
    src = tmp_path / "table_figure_nospace.docx"
    image_path = tmp_path / "pixel_table_nospace.png"
    out = tmp_path / "table_figure_nospace_out.docx"
    _create_table_figure_docx(src, image_path)

    doc = DocxDocument(src)
    doc.paragraphs[-1].text = (
        "Figure 3.Packaging of U2 Total Knee System PSA type extension line "
        "(Stem length 75 mm) Double blister system. (D) Side View (E) The position "
        "of the label in front view (F) Bottom-view; Inner blister"
    )
    doc.save(src)

    result = figure_xml.extract_specific_figure_from_word_xml(
        input_file=str(src),
        output_docx_path=str(out),
        target_chapter_section="",
        target_caption_label="",
        target_figure_index=1,
        allow_table_figure_container=True,
        save_output=True,
        return_reason=True,
    )

    assert result["ok"] is True
    assert result["match_mode"] == "figure_index"
    assert result["selected_caption_analysis"]["accepted"] is True
    assert result["selected_caption_analysis"]["reason"] == "figure_number_prefix"


def test_extract_specific_figure_ignores_table_wrapped_image_by_default(tmp_path: Path) -> None:
    src = tmp_path / "table_figure_default_off.docx"
    image_path = tmp_path / "pixel_table_default_off.png"
    _create_table_figure_docx(src, image_path)

    result = figure_xml.extract_specific_figure_from_word_xml(
        input_file=str(src),
        output_docx_path=None,
        target_chapter_section="",
        target_caption_label="Figure 2",
        save_output=False,
        return_reason=True,
    )

    assert result["ok"] is False
    assert result["reason"] == "figure_not_found"
