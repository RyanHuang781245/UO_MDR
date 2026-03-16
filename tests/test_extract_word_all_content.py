from pathlib import Path

from docx import Document as DocxDocument

from modules.extract_word_all_content import extract_body_with_options


def test_extract_body_with_ignore_toc_keeps_content_when_no_toc_or_heading(
    tmp_path: Path,
) -> None:
    src = tmp_path / "source.docx"
    out = tmp_path / "output.docx"

    doc = DocxDocument()
    doc.add_paragraph("This document starts directly with body text.")
    doc.add_paragraph("No TOC and no heading styles are present.")
    doc.save(src)

    extract_body_with_options(
        str(src),
        str(out),
        ignore_toc_and_before=True,
        ignore_header_footer=False,
    )

    result = DocxDocument(out)
    texts = [p.text for p in result.paragraphs if p.text]
    assert texts == [
        "This document starts directly with body text.",
        "No TOC and no heading styles are present.",
    ]
