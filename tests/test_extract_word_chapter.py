from pathlib import Path

from docx import Document as DocxDocument
from spire.doc import Document, FileFormat

from modules.Extract_AllFile_to_FinalWord import extract_word_chapter


def test_extract_word_chapter_keeps_title(tmp_path: Path) -> None:
    src = Document()
    sec = src.AddSection()
    sec.AddParagraph().AppendText("1.1 Sample Title")
    sec.AddParagraph().AppendText("Body text")
    src_path = tmp_path / "source.docx"
    src.SaveToFile(str(src_path), FileFormat.Docx)
    src.Close()

    out_doc = Document()
    out_section = out_doc.AddSection()

    result = extract_word_chapter(
        str(src_path),
        "1.1",
        output_doc=out_doc,
        section=out_section,
    )

    out_path = tmp_path / "out.docx"
    out_doc.SaveToFile(str(out_path), FileFormat.Docx)
    out_doc.Close()

    assert result == {"captured_titles": ["1.1 Sample Title"]}

    docx_doc = DocxDocument(out_path)
    paragraphs = [p for p in docx_doc.paragraphs if p.text.strip()]
    title_para = next((p for p in paragraphs if p.text == "1.1 Sample Title"), None)
    assert title_para is not None
    assert all(run.font.hidden for run in title_para.runs)
    assert any("Body text" in p.text for p in paragraphs)


def test_extract_word_chapter_stops_at_next_heading_text(tmp_path: Path) -> None:
    src = Document()
    sec = src.AddSection()
    sec.AddParagraph().AppendText("1.1 Sample Title")
    sec.AddParagraph().AppendText("Body text")
    sec.AddParagraph().AppendText("1.2 Another Section")
    sec.AddParagraph().AppendText("Next section body")
    src_path = tmp_path / "source.docx"
    src.SaveToFile(str(src_path), FileFormat.Docx)
    src.Close()

    out_doc = Document()
    out_section = out_doc.AddSection()

    result = extract_word_chapter(
        str(src_path),
        "1.1",
        output_doc=out_doc,
        section=out_section,
    )

    out_path = tmp_path / "out.docx"
    out_doc.SaveToFile(str(out_path), FileFormat.Docx)
    out_doc.Close()

    assert result == {"captured_titles": ["1.1 Sample Title"]}

    docx_doc = DocxDocument(out_path)
    paragraphs = [p.text.strip() for p in docx_doc.paragraphs if p.text.strip()]
    assert "1.2 Another Section" not in paragraphs
    assert all("Next section body" not in p for p in paragraphs)
