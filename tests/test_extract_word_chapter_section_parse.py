from pathlib import Path

from docx import Document

import modules.Extract_AllFile_to_FinalWord as chapter_module
from modules.Extract_AllFile_to_FinalWord import _parse_chapter_section_expression, extract_word_chapter
from modules.extract_word_chapter import get_all_text


def test_parse_section_with_dot_and_title() -> None:
    assert _parse_chapter_section_expression("1. 測試1.1") == ("1", "", "測試1.1")


def test_parse_section_range_with_title() -> None:
    assert _parse_chapter_section_expression("1.1.1 - 1.1.3 測試標題") == ("1.1.1", "1.1.3", "測試標題")


def test_parse_section_without_title() -> None:
    assert _parse_chapter_section_expression("1.1.1-1.1.3") == ("1.1.1", "1.1.3", "")


def test_parse_end_marker_with_trailing_dot_and_title() -> None:
    assert _parse_chapter_section_expression("2. Materials") == ("2", "", "Materials")


def test_extract_word_chapter_splits_end_marker_with_trailing_dot(monkeypatch, tmp_path: Path) -> None:
    src = tmp_path / "source.docx"
    src.write_bytes(b"stub")
    captured: dict = {}

    def fake_extract_section_docx_xml(*_args, **kwargs):
        captured.update(kwargs)

    monkeypatch.setattr(chapter_module, "extract_section_docx_xml", fake_extract_section_docx_xml)
    monkeypatch.setattr(chapter_module, "_read_first_paragraph_text", lambda *_args, **_kwargs: "")

    extract_word_chapter(
        str(src),
        "1.1",
        target_chapter_title="Start title",
        explicit_end_title="2. Materials",
        output_docx_path=str(tmp_path / "out.docx"),
    )

    assert captured["explicit_end_number"] == "2"
    assert captured["explicit_end_title"] == "Materials"


def test_get_all_text_ignores_hidden_runs(tmp_path: Path) -> None:
    src = tmp_path / "source.docx"
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Visible")
    hidden = para.add_run("Hidden prompt")
    hidden.font.hidden = True
    doc.save(src)

    loaded = Document(src)

    assert get_all_text(loaded.paragraphs[0]._p) == "Visible"


def test_extract_word_chapter_ignores_hidden_heading_prompt(tmp_path: Path) -> None:
    src = tmp_path / "source.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("目的/Objectives:")
    hidden = heading.add_run("請輸入文字、網址，你也可翻譯文件。")
    hidden.font.hidden = True
    doc.add_paragraph("Body text")
    doc.save(src)

    extract_word_chapter(
        str(src),
        "",
        target_chapter_title="目的/Objectives:",
        output_docx_path=str(out),
        llm_boundary_fallback=False,
    )

    result = Document(out)
    assert get_all_text(result.paragraphs[0]._p) == "目的/Objectives:"
