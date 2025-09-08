import os
import pytest
from spire.doc import Document, FileFormat, HorizontalAlignment
from docx import Document as DocxDocument

from modules.mapping_processor import process_mapping_excel

openpyxl = pytest.importorskip("openpyxl")
from openpyxl import Workbook


def test_process_mapping_centers_and_renumbers(tmp_path):
    # Create source document with misnumbered captions
    doc = Document()
    sec = doc.AddSection()
    p1 = sec.AddParagraph()
    p1.AppendText("Figure 5 Sample figure")
    p2 = sec.AddParagraph()
    p2.AppendText("Table 9 Sample table")
    src_path = tmp_path / "src.docx"
    doc.SaveToFile(str(src_path), FileFormat.Docx)
    doc.Close()

    # Build mapping file
    wb = Workbook()
    ws = wb.active
    ws.append(["A", "B", "C", "D"])
    ws.append(["OutDoc", "", "src.docx", "all"])
    mapping_path = tmp_path / "map.xlsx"
    wb.save(mapping_path)

    out_dir = tmp_path / "out"
    result = process_mapping_excel(str(mapping_path), str(tmp_path), str(out_dir))
    out_path = os.path.join(out_dir, "OutDoc.docx")
    assert out_path in result["outputs"]

    out = Document()
    out.LoadFromFile(out_path)
    sec = out.Sections.get_Item(0)
    fig = sec.Paragraphs.get_Item(0)
    tab = sec.Paragraphs.get_Item(1)
    assert "Figure 1" in fig.Text
    assert fig.Format.HorizontalAlignment == HorizontalAlignment.Center
    assert "Table 1" in tab.Text
    assert tab.Format.HorizontalAlignment == HorizontalAlignment.Center
    out.Close()

    # verify basic style applied (font and line spacing)
    docx_doc = DocxDocument(out_path)
    p = docx_doc.paragraphs[0]
    assert p.paragraph_format.line_spacing == pytest.approx(1.5)
    run = p.runs[0]
    assert run.font.name == "Times New Roman"
    assert run.font.size.pt == pytest.approx(12)


def test_process_mapping_strips_chapter_numbers(tmp_path):
    doc = Document()
    sec = doc.AddSection()
    p1 = sec.AddParagraph()
    p1.AppendText("6.4.2 Heading")
    p2 = sec.AddParagraph()
    p2.AppendText("Body")
    src_path = tmp_path / "src.docx"
    doc.SaveToFile(str(src_path), FileFormat.Docx)
    doc.Close()

    wb = Workbook()
    ws = wb.active
    ws.append(["A", "B", "C", "D"])
    ws.append(["Out", "6.4.2 Heading", "src.docx", "6.4.2"])
    mapping_path = tmp_path / "map.xlsx"
    wb.save(mapping_path)

    out_dir = tmp_path / "out"
    process_mapping_excel(str(mapping_path), str(tmp_path), str(out_dir))
    out_path = os.path.join(out_dir, "Out.docx")
    docx_doc = DocxDocument(out_path)
    text = "\n".join(p.text for p in docx_doc.paragraphs)
    assert "6.4.2" not in text
    assert "Heading" in text


def test_process_mapping_folder_input(tmp_path):
    # Create a folder containing a single document
    subdir = tmp_path / "Folder"
    subdir.mkdir()
    doc = Document()
    sec = doc.AddSection()
    sec.AddParagraph().AppendText("content")
    src = subdir / "inner.docx"
    doc.SaveToFile(str(src), FileFormat.Docx)
    doc.Close()

    # Mapping file specifies the folder instead of the file name
    wb = Workbook()
    ws = wb.active
    ws.append(["A", "B", "C", "D"])
    ws.append(["Out", "Title", "Folder", "all"])
    mapping = tmp_path / "map.xlsx"
    wb.save(mapping)

    out_dir = tmp_path / "out"
    process_mapping_excel(str(mapping), str(tmp_path), str(out_dir))
    assert (out_dir / "Out.docx").exists()


def test_process_mapping_copy_from_folder(tmp_path):
    # Source directories
    root = tmp_path
    sub = root / "sub"
    other = root / "other"
    sub.mkdir()
    other.mkdir()
    (sub / "match EO.txt").write_text("data")
    (other / "match EO.txt").write_text("should not copy")

    wb = Workbook()
    ws = wb.active
    ws.append(["A", "B", "C", "D"])
    ws.append(["Dest", "", "sub", "match, EO"])
    mapping = root / "map.xlsx"
    wb.save(mapping)

    out_dir = root / "out"
    process_mapping_excel(str(mapping), str(root), str(out_dir))
    dest_dir = root / "Dest"
    assert dest_dir.joinpath("match EO.txt").exists()
    # ensure only file from specified folder is copied
    assert len(list(dest_dir.glob("*.txt"))) == 1
