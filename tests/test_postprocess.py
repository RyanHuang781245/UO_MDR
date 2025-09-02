import os
from zipfile import ZipFile

from docx import Document

from modules.postprocess import renumber_figures_tables


def test_renumber_figures_tables(tmp_path):
    doc_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Figure 5: A cat")
    doc.add_paragraph("See Figure 5 for details. Table 3 shows numbers.")
    doc.add_paragraph("Table 3: Data table")
    doc.add_paragraph("Another reference to Table 3 and Figure 5.")
    doc.save(doc_path)

    # Simulate an existing image in the document package to ensure it is preserved
    with ZipFile(doc_path, "a") as zf:
        zf.writestr("word/media/fake.png", b"data")
    original_entries = set(ZipFile(doc_path).namelist())

    renumber_figures_tables(str(doc_path))

    processed = Document(doc_path)
    texts = [p.text for p in processed.paragraphs]
    assert texts[0].startswith("Figure 1")
    assert "Figure 1" in texts[1]
    assert texts[2].startswith("Table 1")
    assert "Table 1" in texts[3]

    # Ensure no package parts (such as images) were lost during processing
    assert original_entries == set(ZipFile(doc_path).namelist())
