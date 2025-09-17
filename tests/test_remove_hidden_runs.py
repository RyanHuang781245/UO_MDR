from pathlib import Path

from docx import Document

from modules.Extract_AllFile_to_FinalWord import (
    remove_hidden_runs,
    remove_paragraphs_with_text,
)


def test_remove_hidden_runs_keeps_paragraph_in_table_cell(tmp_path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    para = cell.paragraphs[0]
    run = para.add_run("to hide")
    run.font.hidden = True

    doc_path = tmp_path / "table.docx"
    doc.save(doc_path)

    assert remove_hidden_runs(str(doc_path))

    updated = Document(doc_path)
    cell_after = updated.tables[0].cell(0, 0)

    # Even though all visible text was removed, the table cell must keep a paragraph
    assert len(cell_after.paragraphs) == 1
    assert cell_after.paragraphs[0].text == ""


def test_remove_hidden_runs_preserves_titles(tmp_path: Path) -> None:
    doc = Document()
    para = doc.add_paragraph("1.1 Sample Title")
    for run in para.runs:
        run.font.hidden = True

    doc_path = tmp_path / "title.docx"
    doc.save(doc_path)

    assert remove_hidden_runs(str(doc_path), preserve_texts=["1.1 Sample Title"])

    updated = Document(doc_path)
    assert updated.paragraphs[0].text == "1.1 Sample Title"
    assert all(run.font.hidden for run in updated.paragraphs[0].runs)


def test_remove_paragraphs_with_text_removes_titles(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("1.1 Sample Title")
    doc.add_paragraph("Body")

    doc_path = tmp_path / "strip.docx"
    doc.save(doc_path)

    assert remove_paragraphs_with_text(str(doc_path), ["1.1 Sample Title"])

    updated = Document(doc_path)
    texts = [p.text for p in updated.paragraphs]
    assert "1.1 Sample Title" not in texts
    assert texts == ["Body"]


def test_remove_paragraphs_with_text_preserves_table_structure(tmp_path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "1.1 Sample Title"

    doc_path = tmp_path / "table_strip.docx"
    doc.save(doc_path)

    assert remove_paragraphs_with_text(str(doc_path), ["1.1 Sample Title"])

    updated = Document(doc_path)
    cell_after = updated.tables[0].cell(0, 0)
    assert len(cell_after.paragraphs) == 1
    assert cell_after.paragraphs[0].text == ""
