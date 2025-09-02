import os
from docx import Document
from modules.Extract_AllFile_to_FinalWord import renumber_figures_tables

def test_renumber_figures_tables(tmp_path):
    doc = Document()
    doc.add_paragraph("Figure 5. Sample figure")
    doc.add_paragraph("As shown in Figure 5 above.")
    doc.add_paragraph("Table 10: Sample table")
    doc.add_paragraph("Refer to Table 10 for details.")
    file_path = tmp_path / "test.docx"
    doc.save(file_path)

    assert renumber_figures_tables(str(file_path))

    updated = Document(file_path)
    texts = [p.text for p in updated.paragraphs]
    assert texts[0].startswith("Figure 1")
    assert "Figure 1" in texts[1]
    assert texts[2].startswith("Table 1")
    assert "Table 1" in texts[3]
