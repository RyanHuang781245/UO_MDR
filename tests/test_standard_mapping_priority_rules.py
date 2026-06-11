import os
import tempfile
import zipfile

from docx import Document
from lxml import etree
from openpyxl import Workbook

from app.services.standard_mapping_service import (
    build_diff_segments,
    build_preserve_original_segments,
    build_title_with_amendment,
    classify_standard_level,
    extract_harmonised_reference_entries,
    extract_harmonised_reference_keys,
    extract_standard_match_key,
    extract_latest_year_from_astm_style,
    find_latest_year_from_excel,
    is_harmonised_standard,
    is_accepted_eu_harmonised_form,
    load_harmonised_reference_index,
    load_excel_index,
    normalize_harmonised_standard_text,
    normalize_iso_priority,
    process_document,
    resolve_target_table_label_map,
    resolve_target_table_scan_label_map,
    resolve_target_table_indexes,
)
from app.blueprints.tasks.standard_mapping_routes import _build_stats


def _add_single_cell_table(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = text


def _write_standard_workbook(path, standard_no: str, title: str = "New Title") -> None:
    wb = Workbook()
    default = wb.active
    default.title = "BS-EN-DIN(歐洲國家標準)"
    for sheet_name in ("BS-EN-DIN(歐洲國家標準)", "ISO", "ASTM"):
        ws = wb[sheet_name] if sheet_name in wb.sheetnames else wb.create_sheet(sheet_name)
        ws.cell(row=1, column=1, value="標準名稱")
        ws.cell(row=1, column=6, value="Standards")
    ws = wb["BS-EN-DIN(歐洲國家標準)"]
    ws.cell(row=2, column=1, value=title)
    ws.cell(row=2, column=6, value=standard_no)
    wb.save(path)


def _write_harmonised_workbook(path, reference_text: str) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "2017-745-Medical Devices"
    ws.cell(row=3, column=3, value="Reference and title Provision")
    ws.cell(row=4, column=3, value=reference_text)
    wb.save(path)


def _write_standard_workbook_rows(path, rows: list[tuple[str, str]]) -> None:
    wb = Workbook()
    default = wb.active
    default.title = "BS-EN-DIN(歐洲國家標準)"
    for sheet_name in ("BS-EN-DIN(歐洲國家標準)", "ISO", "ASTM"):
        ws = wb[sheet_name] if sheet_name in wb.sheetnames else wb.create_sheet(sheet_name)
        ws.cell(row=1, column=1, value="標準名稱")
        ws.cell(row=1, column=6, value="Standards")
    ws = wb["BS-EN-DIN(歐洲國家標準)"]
    for row_idx, (standard_no, title) in enumerate(rows, start=2):
        ws.cell(row=row_idx, column=1, value=title)
        ws.cell(row=row_idx, column=6, value=standard_no)
    wb.save(path)


def test_normalize_iso_priority_expands_full_priority_order():
    assert normalize_iso_priority(("ISO", "BS EN ISO")) == (
        "ISO",
        "BS EN ISO",
        "BS EN",
        "EN",
        "EN ISO",
        "BS ISO",
        "BS",
    )


def test_extract_standard_match_key_removes_prefix_year_and_amendment():
    assert (
        extract_standard_match_key("EN ISO 14971:2019/A11:2021", "BS-EN-DIN(歐洲國家標準)")
        == "14971"
    )
    assert extract_standard_match_key("IEC 60601-1-2:2014", "ISO") == "60601-1-2"


def test_extract_latest_year_from_astm_style_prefers_bracket_year():
    assert extract_latest_year_from_astm_style("ASTM F1140/F1140M-13(2025)") == 2025


def test_classify_standard_level_checks_en_iso_before_en():
    assert classify_standard_level("EN ISO 11137-2:2015") == ("EN ISO", 4)
    assert classify_standard_level("DIN EN ISO 11137-2:2015") == ("EN ISO", 4)
    assert classify_standard_level("EN 455-1:2020") == ("EN", 5)


def test_build_title_with_amendment_supports_plus_a_and_amd_formats():
    assert (
        build_title_with_amendment(
            "Application of risk management to medical devices",
            "BS EN ISO 14971:2019+A11:2021",
        )
        == "Application of risk management to medical devices - Amendment 11"
    )
    assert (
        build_title_with_amendment(
            "Hexalobular internal driving feature for bolts and screws",
            "ISO 16047：2005/Amd 1：2012",
        )
        == "Hexalobular internal driving feature for bolts and screws - Amendment 1"
    )


def test_build_preserve_original_segments_keeps_old_and_marks_new_red():
    assert build_preserve_original_segments("No", "Yes") == [
        ("No ", False),
        ("Yes", True),
    ]
    assert build_preserve_original_segments("Yes", "No") == [
        ("Yes ", False),
        ("No", True),
    ]


def test_build_diff_segments_matches_preview_for_standard_number_change():
    assert build_diff_segments("EN ISO 15223-1", "EN ISO 29176") == [
        ("EN ISO ", False),
        ("29176", True),
    ]


def test_harmonised_matching_accepts_same_standard_match_key():
    lookup = {
        normalize_harmonised_standard_text("EN ISO 14971:2019\nApplication of risk management to medical devices"): {2019}
    }
    assert is_harmonised_standard("EN ISO 14971:2019", "Different title", lookup) is True
    assert is_harmonised_standard("BS EN ISO 14971:2019", "Application of risk management to medical devices", lookup) is False
    assert is_harmonised_standard("ISO 14971:2019", "Application of risk management to medical devices", lookup) is False


def test_harmonised_matching_accepts_same_standard_number_from_reference_key():
    lookup = {
        key: {2019}
        for key in extract_harmonised_reference_keys("EN ISO 14971:2019\nApplication of risk management to medical devices")
    }
    assert is_harmonised_standard("BS EN ISO 14971:2019", "Application of risk management to medical devices", lookup) is True
    assert is_harmonised_standard("ISO 14971:2019", "Application of risk management to medical devices", lookup) is False
    assert is_harmonised_standard("ISO 13485:2016", "Different standard", lookup) is False


def test_harmonised_matching_requires_accepted_en_adoption_form():
    lookup = {
        "14971": {2019},
        "IEC 60601-1": {2020},
    }

    assert is_accepted_eu_harmonised_form("EN ISO 14971:2019") is True
    assert is_accepted_eu_harmonised_form("BS EN ISO 14971:2019") is True
    assert is_accepted_eu_harmonised_form("EN IEC 60601-1:2020") is True
    assert is_accepted_eu_harmonised_form("ISO 14971:2019") is False
    assert is_accepted_eu_harmonised_form("BS ISO 14971:2019") is False
    assert is_accepted_eu_harmonised_form("IEC 60601-1:2020") is False

    assert is_harmonised_standard("EN ISO 14971:2019", "", lookup) is True
    assert is_harmonised_standard("BS EN ISO 14971:2019", "", lookup) is True
    assert is_harmonised_standard("EN IEC 60601-1:2020", "", lookup) is True
    assert is_harmonised_standard("ISO 14971:2019", "", lookup) is False
    assert is_harmonised_standard("BS ISO 14971:2019", "", lookup) is False
    assert is_harmonised_standard("IEC 60601-1:2020", "", lookup) is False


def test_extract_harmonised_reference_entries_collects_all_standard_lines():
    value = (
        "EN ISO 13485:2016\n"
        "Medical devices - Quality management systems - Requirements for regulatory purposes (ISO 13485:2016)\n"
        "EN ISO 13485:2016/AC:2018\n"
        "EN ISO 13485:2016/A11:2021"
    )
    assert extract_harmonised_reference_entries(value) == [
        "EN ISO 13485:2016",
        "EN ISO 13485:2016/AC:2018",
        "EN ISO 13485:2016/A11:2021",
    ]


def test_build_stats_counts_harmonised_yes_fallback_as_updated():
    stats = _build_stats([
        {"status": "UPDATED"},
        {"status": "SAME_NO_UPDATE"},
        {"status": "NOT_FOUND"},
        {"status": "HARMONISED_YES_FALLBACK"},
    ])

    assert stats == {
        "updated": 2,
        "same": 1,
        "missing": 1,
        "harmonised_fallback": 1,
        "total": 4,
    }


def test_find_latest_year_prefers_newer_year_before_priority():
    excel_index = {
        "BS-EN-DIN(歐洲國家標準)": [
            {
                "sheet_name": "BS-EN-DIN(歐洲國家標準)",
                "excel_row_index": 2,
                "excel_col_letter": "F",
                "standard_no": "BS EN ISO 14971:2019",
                "standard_match_key": "14971",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "BS EN ISO 14971",
                "standard_level": "BS EN ISO",
                "standard_level_rank": 7,
            },
            {
                "sheet_name": "BS-EN-DIN(歐洲國家標準)",
                "excel_row_index": 3,
                "excel_col_letter": "F",
                "standard_no": "EN ISO 14971:2021",
                "standard_match_key": "14971",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "EN ISO 14971",
                "standard_level": "EN ISO",
                "standard_level_rank": 4,
            },
        ],
        "ISO": [],
        "ASTM": [],
    }

    result = find_latest_year_from_excel(
        "ISO 14971:2019",
        excel_index,
        ("BS EN ISO", "BS EN", "EN", "EN ISO", "BS ISO", "ISO", "BS"),
    )

    assert result is not None
    assert result["matched_standard_no"] == "EN ISO 14971:2021"
    assert result["latest_year"] == 2021


def test_find_latest_year_prefers_en_group_over_newer_non_en_candidate():
    excel_index = {
        "BS-EN-DIN(歐洲國家標準)": [
            {
                "sheet_name": "BS-EN-DIN(歐洲國家標準)",
                "excel_row_index": 2,
                "excel_col_letter": "F",
                "standard_no": "EN ISO 14971:2021",
                "standard_match_key": "14971",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "EN ISO 14971",
                "standard_level": "EN ISO",
                "standard_level_rank": 4,
            }
        ],
        "ISO": [
            {
                "sheet_name": "ISO",
                "excel_row_index": 4,
                "excel_col_letter": "F",
                "standard_no": "ISO 14971:2023",
                "standard_match_key": "14971",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "ISO 14971",
                "standard_level": "ISO",
                "standard_level_rank": 2,
            }
        ],
        "ASTM": [],
    }

    result = find_latest_year_from_excel(
        "ISO 14971:2019",
        excel_index,
        ("BS EN ISO", "BS EN", "EN", "EN ISO", "BS ISO", "ISO", "BS"),
    )

    assert result is not None
    assert result["matched_standard_no"] == "EN ISO 14971:2021"
    displayed_standards = [candidate["matched_standard_no"] for candidate in result["all_candidates"]]
    assert displayed_standards.index("EN ISO 14971:2021") < displayed_standards.index("ISO 14971:2023")


def test_find_latest_year_displays_candidates_by_priority_when_no_levels_are_checked():
    excel_index = {
        "BS-EN-DIN(歐洲國家標準)": [
            {
                "sheet_name": "BS-EN-DIN(歐洲國家標準)",
                "excel_row_index": 2,
                "excel_col_letter": "F",
                "standard_no": "BS EN ISO 14971:2021",
                "standard_match_key": "14971",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "BS EN ISO 14971",
                "standard_level": "BS EN ISO",
                "standard_level_rank": 7,
            },
        ],
        "ISO": [
            {
                "sheet_name": "ISO",
                "excel_row_index": 3,
                "excel_col_letter": "F",
                "standard_no": "ISO 14971:2024",
                "standard_match_key": "14971",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "ISO 14971",
                "standard_level": "ISO",
                "standard_level_rank": 2,
            },
        ],
        "ASTM": [],
    }

    result = find_latest_year_from_excel(
        "ISO 14971:2019",
        excel_index,
        ("BS EN ISO", "BS EN", "EN", "EN ISO", "BS ISO", "ISO", "BS"),
        enabled_standard_levels=(),
    )

    assert result is not None
    assert result["matched_standard_no"] == "BS EN ISO 14971:2021"
    displayed_standards = [candidate["matched_standard_no"] for candidate in result["all_candidates"]]
    assert displayed_standards[:2] == ["BS EN ISO 14971:2021", "ISO 14971:2024"]


def test_find_latest_year_does_not_treat_en_iso_as_en_for_priority():
    excel_index = {
        "BS-EN-DIN(歐洲國家標準)": [
            {
                "sheet_name": "BS-EN-DIN(歐洲國家標準)",
                "excel_row_index": 13,
                "excel_col_letter": "F",
                "standard_no": "EN ISO 11137-2:2015",
                "standard_match_key": "11137-2",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "EN ISO 11137-2",
                "standard_level": "EN ISO",
                "standard_level_rank": 4,
            },
            {
                "sheet_name": "BS-EN-DIN(歐洲國家標準)",
                "excel_row_index": 14,
                "excel_col_letter": "F",
                "standard_no": "BS EN ISO 11137-2:2015+A1:2023",
                "standard_match_key": "11137-2",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "BS EN ISO 11137-2",
                "standard_level": "BS EN ISO",
                "standard_level_rank": 7,
            },
        ],
        "ISO": [],
        "ASTM": [],
    }

    result = find_latest_year_from_excel(
        "ISO 11137-2:2015",
        excel_index,
        ("BS EN", "EN", "BS EN ISO", "EN ISO", "BS ISO", "ISO", "BS"),
        enabled_standard_levels=(),
    )

    assert result is not None
    assert result["matched_standard_no"] == "BS EN ISO 11137-2:2015+A1:2023"
    displayed_standards = [candidate["matched_standard_no"] for candidate in result["all_candidates"]]
    assert displayed_standards[:2] == [
        "BS EN ISO 11137-2:2015+A1:2023",
        "EN ISO 11137-2:2015",
    ]


def test_find_latest_year_classifies_en_iso_correctly_from_workbook(tmp_path):
    standard_excel_path = tmp_path / "standards.xlsx"
    _write_standard_workbook_rows(
        standard_excel_path,
        [
            ("EN ISO 11137-2:2015", "Sterilization older EN ISO"),
            ("BS EN ISO 11137-2:2015+A1:2023", "Sterilization newer BS EN ISO"),
        ],
    )

    result = find_latest_year_from_excel(
        "ISO 11137-2:2015",
        load_excel_index(str(standard_excel_path)),
        ("BS EN", "EN", "BS EN ISO", "EN ISO", "BS ISO", "ISO", "BS"),
        enabled_standard_levels=(),
    )

    assert result is not None
    assert result["matched_standard_no"] == "BS EN ISO 11137-2:2015+A1:2023"
    displayed = [(candidate["matched_standard_no"], candidate["standard_level"]) for candidate in result["all_candidates"]]
    assert displayed[:2] == [
        ("BS EN ISO 11137-2:2015+A1:2023", "BS EN ISO"),
        ("EN ISO 11137-2:2015", "EN ISO"),
    ]


def test_find_latest_year_keeps_harmonised_amendment_candidate_when_no_levels_are_checked(tmp_path):
    standard_excel_path = tmp_path / "standards.xlsx"
    harmonised_path = tmp_path / "harmonised.xlsx"
    _write_standard_workbook_rows(
        standard_excel_path,
        [
            ("BS EN ISO 14971:2019", "Risk management"),
            ("BS EN ISO 14971:2019+A11:2021", "Risk management"),
        ],
    )
    _write_harmonised_workbook(
        harmonised_path,
        "EN ISO 14971:2019\n"
        "Medical devices - Application of risk management to medical devices (ISO 14971:2019)\n"
        "EN ISO 14971:2019/A11:2021",
    )

    result = find_latest_year_from_excel(
        "ISO 14971:2019",
        load_excel_index(str(standard_excel_path)),
        ("BS EN ISO", "BS EN", "EN", "EN ISO", "BS ISO", "ISO", "BS"),
        enabled_standard_levels=(),
        harmonised_reference_index=load_harmonised_reference_index(str(harmonised_path)),
    )

    assert result is not None
    assert result["matched_standard_no"] == "BS EN ISO 14971:2019+A11:2021"
    assert result["matched_harmonised"] == "YES"
    assert result["harmonised_yes_fallback"] is False
    displayed_standards = [candidate["matched_standard_no"] for candidate in result["all_candidates"]]
    assert displayed_standards[:2] == [
        "BS EN ISO 14971:2019+A11:2021",
        "BS EN ISO 14971:2019",
    ]


def test_find_latest_year_uses_priority_as_tiebreaker():
    excel_index = {
        "BS-EN-DIN(歐洲國家標準)": [
            {
                "sheet_name": "BS-EN-DIN(歐洲國家標準)",
                "excel_row_index": 2,
                "excel_col_letter": "F",
                "standard_no": "EN ISO 14971:2021",
                "standard_match_key": "14971",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "EN ISO 14971",
                "standard_level": "EN ISO",
                "standard_level_rank": 4,
            },
            {
                "sheet_name": "BS-EN-DIN(歐洲國家標準)",
                "excel_row_index": 3,
                "excel_col_letter": "F",
                "standard_no": "BS EN ISO 14971:2021",
                "standard_match_key": "14971",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "BS EN ISO 14971",
                "standard_level": "BS EN ISO",
                "standard_level_rank": 7,
            },
        ],
        "ISO": [],
        "ASTM": [],
    }

    result = find_latest_year_from_excel(
        "ISO 14971:2019",
        excel_index,
        ("BS EN ISO", "BS EN", "EN", "EN ISO", "BS ISO", "ISO", "BS"),
    )

    assert result is not None
    assert result["matched_standard_no"] == "BS EN ISO 14971:2021"


def test_find_latest_year_candidate_title_includes_amendment_suffix():
    excel_index = {
        "BS-EN-DIN(歐洲國家標準)": [
            {
                "sheet_name": "BS-EN-DIN(歐洲國家標準)",
                "excel_row_index": 3,
                "excel_col_letter": "F",
                "standard_no": "BS EN ISO 14971:2019+A11:2021",
                "standard_title": "Application of risk management to medical devices",
                "standard_match_key": "14971",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "BS EN ISO 14971",
                "standard_level": "BS EN ISO",
                "standard_level_rank": 7,
            },
        ],
        "ISO": [],
        "ASTM": [],
    }

    result = find_latest_year_from_excel(
        "ISO 14971:2019",
        excel_index,
        ("BS EN ISO", "BS EN", "EN", "EN ISO", "BS ISO", "ISO", "BS"),
    )

    assert result is not None
    assert result["matched_title"] == "Application of risk management to medical devices - Amendment 11"
    assert result["all_candidates"][0]["matched_title"] == "Application of risk management to medical devices - Amendment 11"


def test_find_latest_year_marks_harmonised_yes_from_reference_index():
    excel_index = {
        "BS-EN-DIN(歐洲國家標準)": [
            {
                "sheet_name": "BS-EN-DIN(歐洲國家標準)",
                "excel_row_index": 3,
                "excel_col_letter": "F",
                "standard_no": "BS EN ISO 14971:2019",
                "standard_title": "Application of risk management to medical devices",
                "standard_match_key": "14971",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "BS EN ISO 14971",
                "standard_level": "BS EN ISO",
                "standard_level_rank": 7,
            },
        ],
        "ISO": [],
        "ASTM": [],
    }

    harmonised_lookup = {
        normalize_harmonised_standard_text("BS EN ISO 14971:2019\nApplication of risk management to medical devices"): {2019}
    }
    result = find_latest_year_from_excel(
        "ISO 14971:2019",
        excel_index,
        ("BS EN ISO", "BS EN", "EN", "EN ISO", "BS ISO", "ISO", "BS"),
        harmonised_reference_index=harmonised_lookup,
    )

    assert result is not None
    assert result["matched_harmonised"] == "YES"
    assert result["all_candidates"][0]["candidate_harmonised"] == "YES"


def test_find_latest_year_marks_harmonised_no_when_not_found():
    excel_index = {
        "BS-EN-DIN(歐洲國家標準)": [
            {
                "sheet_name": "BS-EN-DIN(歐洲國家標準)",
                "excel_row_index": 3,
                "excel_col_letter": "F",
                "standard_no": "BS EN ISO 14971:2019",
                "standard_title": "Application of risk management to medical devices",
                "standard_match_key": "14971",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "BS EN ISO 14971",
                "standard_level": "BS EN ISO",
                "standard_level_rank": 7,
            },
        ],
        "ISO": [],
        "ASTM": [],
    }

    result = find_latest_year_from_excel(
        "ISO 14971:2019",
        excel_index,
        ("BS EN ISO", "BS EN", "EN", "EN ISO", "BS ISO", "ISO", "BS"),
        harmonised_reference_index=set(),
    )

    assert result is not None
    assert result["matched_harmonised"] == "NO"
    assert result["all_candidates"][0]["candidate_harmonised"] == "NO"


def test_find_latest_year_selects_harmonised_yes_candidate_over_newer_no_candidate():
    excel_index = {
        "BS-EN-DIN(歐洲國家標準)": [
            {
                "sheet_name": "BS-EN-DIN(歐洲國家標準)",
                "excel_row_index": 2,
                "excel_col_letter": "F",
                "standard_no": "BS EN ISO 14971:2021",
                "standard_title": "Newer non-harmonised",
                "standard_match_key": "14971",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "BS EN ISO 14971",
                "standard_level": "BS EN ISO",
                "standard_level_rank": 7,
            },
            {
                "sheet_name": "BS-EN-DIN(歐洲國家標準)",
                "excel_row_index": 3,
                "excel_col_letter": "F",
                "standard_no": "EN ISO 14971:2019",
                "standard_title": "Older harmonised",
                "standard_match_key": "14971",
                "search_family": "ISO_FAMILY",
                "standard_display_no": "EN ISO 14971",
                "standard_level": "EN ISO",
                "standard_level_rank": 4,
            },
        ],
        "ISO": [],
        "ASTM": [],
    }
    harmonised_lookup = {"14971": {2019}}

    result = find_latest_year_from_excel(
        "ISO 14971:2019",
        excel_index,
        ("BS EN ISO", "BS EN", "EN", "EN ISO", "BS ISO", "ISO", "BS"),
        harmonised_reference_index=harmonised_lookup,
    )

    assert result is not None
    assert result["matched_standard_no"] == "EN ISO 14971:2019"
    assert result["latest_year"] == 2019
    assert result["matched_harmonised"] == "YES"
    assert result["harmonised_yes_fallback"] is True
    assert result["harmonised_rejected_candidate_id"]
    assert result["all_candidates"][0]["matched_standard_no"] == "BS EN ISO 14971:2021"
    assert any(
        candidate["matched_standard_no"] == "EN ISO 14971:2019" and candidate["decision"] == "selected"
        for candidate in result["all_candidates"]
    )


def test_process_document_skips_entire_row_when_harmonised_result_is_no(tmp_path):
    word_path = tmp_path / "input.docx"
    standard_excel_path = tmp_path / "standards.xlsx"
    harmonised_path = tmp_path / "harmonised.xlsx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    table = doc.add_table(rows=2, cols=4)
    headers = [
        "Standards",
        "Issued Year",
        "EU Harmonised Standards under MDR 2017/745 (YES/NO)",
        "Title",
    ]
    original_values = ["ISO 14971", "2019", "YES", "Old Title"]
    for col_idx, header in enumerate(headers):
        table.cell(0, col_idx).text = header
        table.cell(1, col_idx).text = original_values[col_idx]
    doc.save(word_path)

    _write_standard_workbook(standard_excel_path, "BS EN ISO 14971:2021", "New Title")
    _write_harmonised_workbook(
        harmonised_path,
        "EN ISO 13485:2016\nMedical devices - Quality management systems",
    )

    result = process_document(
        str(word_path),
        str(standard_excel_path),
        harmonised_reference_path=str(harmonised_path),
        output_path=str(output_path),
    )

    output_doc = Document(output_path)
    updated_values = [cell.text for cell in output_doc.tables[0].rows[1].cells]
    row_meta = next(iter(result["reference_payload"].values()))

    assert result["updated_count"] == 0
    assert updated_values == original_values
    assert row_meta["status"] == "SAME_NO_UPDATE"
    assert row_meta["matched_standard_no"] == "ISO 14971"


def test_process_document_marks_harmonised_yes_fallback_update(tmp_path):
    word_path = tmp_path / "input_fallback.docx"
    standard_excel_path = tmp_path / "standards_fallback.xlsx"
    harmonised_path = tmp_path / "harmonised_fallback.xlsx"
    output_path = tmp_path / "output_fallback.docx"

    doc = Document()
    table = doc.add_table(rows=2, cols=4)
    headers = [
        "Standards",
        "Issued Year",
        "EU Harmonised Standards under MDR 2017/745 (YES/NO)",
        "Title",
    ]
    for col_idx, header in enumerate(headers):
        table.cell(0, col_idx).text = header
    for col_idx, value in enumerate(["ISO 14971", "2018", "YES", "Old Title"]):
        table.cell(1, col_idx).text = value
    doc.save(word_path)

    _write_standard_workbook_rows(
        standard_excel_path,
        [
            ("BS EN ISO 14971:2021", "Newer non-harmonised"),
            ("EN ISO 14971:2019", "Older harmonised"),
        ],
    )
    _write_harmonised_workbook(
        harmonised_path,
        "EN ISO 14971:2019\nApplication of risk management to medical devices",
    )

    result = process_document(
        str(word_path),
        str(standard_excel_path),
        harmonised_reference_path=str(harmonised_path),
        output_path=str(output_path),
    )
    stats = _build_stats(result["report"])
    row_meta = next(iter(result["reference_payload"].values()))
    output_doc = Document(output_path)
    updated_values = [cell.text for cell in output_doc.tables[0].rows[1].cells]

    assert row_meta["status"] == "HARMONISED_YES_FALLBACK"
    assert row_meta["matched_standard_no"] == "EN ISO 14971"
    assert row_meta["excel_year"] == "2019"
    assert updated_values[:3] == ["EN ISO 14971", "2019", "YES"]
    assert stats["updated"] == 1
    assert stats["harmonised_fallback"] == 1


def test_process_document_applies_manual_edits_even_when_harmonised_changes_yes_to_no(tmp_path):
    word_path = tmp_path / "input_manual_edit.docx"
    standard_excel_path = tmp_path / "standards_manual_edit.xlsx"
    harmonised_path = tmp_path / "harmonised_manual_edit.xlsx"
    output_path = tmp_path / "output_manual_edit.docx"

    doc = Document()
    table = doc.add_table(rows=2, cols=4)
    headers = [
        "Standards",
        "Issued Year",
        "EU Harmonised Standards under MDR 2017/745 (YES/NO)",
        "Title",
    ]
    original_values = ["ISO 14971", "2019", "YES", "Old Title"]
    for col_idx, header in enumerate(headers):
        table.cell(0, col_idx).text = header
        table.cell(1, col_idx).text = original_values[col_idx]
    doc.save(word_path)

    _write_standard_workbook(standard_excel_path, "BS EN ISO 14971:2021", "New Title")
    _write_harmonised_workbook(
        harmonised_path,
        "EN ISO 13485:2016\nMedical devices - Quality management systems",
    )

    result = process_document(
        str(word_path),
        str(standard_excel_path),
        harmonised_reference_path=str(harmonised_path),
        edit_map={
            "table-0-row-1": {
                "standards": "Manual Standard",
                "issued_year": "2030",
                "harmonised": "NO",
            }
        },
        output_path=str(output_path),
    )

    output_doc = Document(output_path)
    updated_values = [cell.text for cell in output_doc.tables[0].rows[1].cells]
    row_meta = next(iter(result["reference_payload"].values()))

    assert updated_values[:3] == ["Manual Standard", "2030", "NO"]
    assert row_meta["status"] == "UPDATED"
    assert row_meta["manual_edits"] == {
        "standards": "Manual Standard",
        "issued_year": "2030",
        "harmonised": "NO",
    }


def test_process_document_manual_edits_preserve_unedited_candidate_title_amendment(tmp_path):
    word_path = tmp_path / "input_manual_edit_amendment.docx"
    standard_excel_path = tmp_path / "standards_manual_edit_amendment.xlsx"
    harmonised_path = tmp_path / "harmonised_manual_edit_amendment.xlsx"
    output_path = tmp_path / "output_manual_edit_amendment.docx"

    doc = Document()
    table = doc.add_table(rows=2, cols=4)
    headers = [
        "Standards",
        "Issued Year",
        "EU Harmonised Standards under MDR 2017/745 (YES/NO)",
        "Title",
    ]
    original_values = ["ISO 14971", "2019", "YES", "Old Title"]
    for col_idx, header in enumerate(headers):
        table.cell(0, col_idx).text = header
        table.cell(1, col_idx).text = original_values[col_idx]
    doc.save(word_path)

    _write_standard_workbook(
        standard_excel_path,
        "BS EN ISO 14971:2019+A11:2021",
        "Application of risk management to medical devices",
    )
    _write_harmonised_workbook(
        harmonised_path,
        "BS EN ISO 14971:2019+A11:2021\nApplication of risk management to medical devices",
    )

    result = process_document(
        str(word_path),
        str(standard_excel_path),
        harmonised_reference_path=str(harmonised_path),
        edit_map={
            "table-0-row-1": {
                "standards": "Manual Standard",
                "issued_year": "2030",
                "harmonised": "NO",
            }
        },
        output_path=str(output_path),
    )

    output_doc = Document(output_path)
    updated_values = [cell.text for cell in output_doc.tables[0].rows[1].cells]
    row_meta = next(iter(result["reference_payload"].values()))

    assert updated_values == [
        "Manual Standard",
        "2030",
        "NO",
        "Application of risk management to medical devices - Amendment 11",
    ]
    assert row_meta["matched_title"] == "Application of risk management to medical devices - Amendment 11"
    assert "title" not in row_meta["manual_edits"]


def test_process_document_keeps_no_to_no_out_of_harmonised_fallback_stats(tmp_path):
    word_path = tmp_path / "input_same.docx"
    standard_excel_path = tmp_path / "standards_same.xlsx"
    harmonised_path = tmp_path / "harmonised_same.xlsx"
    output_path = tmp_path / "output_same.docx"

    doc = Document()
    table = doc.add_table(rows=2, cols=4)
    headers = [
        "Standards",
        "Issued Year",
        "EU Harmonised Standards under MDR 2017/745 (YES/NO)",
        "Title",
    ]
    original_values = ["ISO 14971", "2019", "NO", "Old Title"]
    for col_idx, header in enumerate(headers):
        table.cell(0, col_idx).text = header
        table.cell(1, col_idx).text = original_values[col_idx]
    doc.save(word_path)

    _write_standard_workbook(standard_excel_path, "BS EN ISO 14971:2021", "New Title")
    _write_harmonised_workbook(
        harmonised_path,
        "EN ISO 13485:2016\nMedical devices - Quality management systems",
    )

    result = process_document(
        str(word_path),
        str(standard_excel_path),
        harmonised_reference_path=str(harmonised_path),
        output_path=str(output_path),
    )
    stats = _build_stats(result["report"])
    output_doc = Document(output_path)
    updated_values = [cell.text for cell in output_doc.tables[0].rows[1].cells]
    row_meta = next(iter(result["reference_payload"].values()))

    assert updated_values == ["BS EN ISO 14971", "2021", "NO", "New Title"]
    assert row_meta["status"] == "UPDATED"
    assert stats["updated"] == 1
    assert stats["harmonised_fallback"] == 0


def test_find_latest_year_keeps_iec_candidate_as_fallback():
    excel_index = {
        "BS-EN-DIN(歐洲國家標準)": [],
        "ISO": [
            {
                "sheet_name": "ISO",
                "excel_row_index": 2,
                "excel_col_letter": "F",
                "standard_no": "IEC 60601-1:2005",
                "standard_match_key": "60601-1",
                "search_family": "IEC_FAMILY",
                "standard_display_no": "IEC 60601-1",
                "standard_level": "IEC",
                "standard_level_rank": 1,
            }
        ],
        "ASTM": [],
    }

    result = find_latest_year_from_excel(
        "IEC 60601-1:2005",
        excel_index,
        ("BS EN ISO", "BS EN", "EN", "EN ISO", "BS ISO", "ISO", "BS"),
    )

    assert result is not None
    assert result["matched_standard_no"] == "IEC 60601-1:2005"


def test_resolve_target_table_indexes_keeps_standards_applied_table_after_up_to_date_block():
    src = os.path.join(os.getcwd(), "table1.docx")
    with tempfile.TemporaryDirectory() as tmpdir:
        with zipfile.ZipFile(src, "r") as archive:
            archive.extractall(tmpdir)

        document_xml_path = os.path.join(tmpdir, "word", "document.xml")
        tree = etree.parse(document_xml_path)

        result = resolve_target_table_indexes(
            tree,
            document_xml_path=document_xml_path,
            target_chapter_ref="4.1.2 Standards applied",
            target_table_index=1,
        )

    assert result == {0}


def test_resolve_target_table_indexes_accepts_multiple_chapter_scopes(tmp_path):
    src = tmp_path / "multi_scope.docx"
    doc = Document()
    doc.add_heading("1 First Section", level=1)
    _add_single_cell_table(doc, "first")
    doc.add_heading("2 Second Section", level=1)
    _add_single_cell_table(doc, "second-a")
    _add_single_cell_table(doc, "second-b")
    doc.add_heading("3 Third Section", level=1)
    _add_single_cell_table(doc, "third")
    doc.save(src)

    extracted = tmp_path / "extracted"
    extracted.mkdir()
    with zipfile.ZipFile(src, "r") as archive:
        archive.extractall(extracted)

    document_xml_path = extracted / "word" / "document.xml"
    tree = etree.parse(str(document_xml_path))

    result = resolve_target_table_indexes(
        tree,
        document_xml_path=str(document_xml_path),
        target_scopes=[
            {"chapter_ref": "1 First Section"},
            {"chapter_ref": "2 Second Section", "table_indexes": (2,)},
            {"chapter_ref": "3 Third Section"},
        ],
    )

    assert result == {0, 2, 3}


def test_resolve_target_table_label_map_uses_chapter_and_optional_scope_index(tmp_path):
    src = tmp_path / "multi_scope_labels.docx"
    doc = Document()
    doc.add_heading("1 First Section", level=1)
    _add_single_cell_table(doc, "first")
    doc.add_heading("2 Second Section", level=1)
    _add_single_cell_table(doc, "second-a")
    _add_single_cell_table(doc, "second-b")
    doc.save(src)

    extracted = tmp_path / "extracted_labels"
    extracted.mkdir()
    with zipfile.ZipFile(src, "r") as archive:
        archive.extractall(extracted)

    document_xml_path = extracted / "word" / "document.xml"
    tree = etree.parse(str(document_xml_path))

    result = resolve_target_table_label_map(
        tree,
        document_xml_path=str(document_xml_path),
        target_scopes=[
            {"chapter_ref": "1 First Section"},
            {"chapter_ref": "2 Second Section", "table_indexes": (2,)},
        ],
    )

    assert result[0] == "1 First Section - 表格 1"
    assert result[2] == "2 Second Section - 表格索引 2"
    assert 1 not in result


def test_resolve_target_table_scan_label_map_omits_long_chapter_title(tmp_path):
    src = tmp_path / "multi_scope_scan_labels.docx"
    doc = Document()
    doc.add_heading("1 Very Long Section Name That Should Not Compress The Board", level=1)
    _add_single_cell_table(doc, "first")
    doc.add_heading("2 Second Section", level=1)
    _add_single_cell_table(doc, "second-a")
    _add_single_cell_table(doc, "second-b")
    doc.save(src)

    extracted = tmp_path / "extracted_scan_labels"
    extracted.mkdir()
    with zipfile.ZipFile(src, "r") as archive:
        archive.extractall(extracted)

    document_xml_path = extracted / "word" / "document.xml"
    tree = etree.parse(str(document_xml_path))

    result = resolve_target_table_scan_label_map(
        tree,
        document_xml_path=str(document_xml_path),
        target_scopes=[
            {"chapter_ref": "1 Very Long Section Name That Should Not Compress The Board"},
            {"chapter_ref": "2 Second Section", "table_indexes": (2,)},
        ],
    )

    assert result[0] == "表格 1"
    assert result[2] == "表格索引 2"
