from pathlib import Path
import zipfile

from docx import Document as DocxDocument
from lxml import etree

from modules.template_manager import parse_template_paragraphs, render_template_with_mappings
from modules.template_mapping import NS


def _create_blank_template(path: Path) -> None:
    doc = DocxDocument()
    doc.sections[0].header.paragraphs[0].text = "Header only"
    doc.sections[0].footer.paragraphs[0].text = "Footer only"
    doc.add_paragraph("")
    doc.save(path)


def _remove_numbering_defaults(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as zin:
        parts = {name: zin.read(name) for name in zin.namelist()}

    numbering_xml = parts.get("word/numbering.xml")
    assert numbering_xml is not None

    root = etree.fromstring(numbering_xml)
    first_lvl = root.find(".//w:abstractNum/w:lvl", namespaces=NS)
    assert first_lvl is not None
    for tag in ("w:start", "w:numFmt", "w:lvlText"):
        node = first_lvl.find(tag, namespaces=NS)
        if node is not None:
            first_lvl.remove(node)
    parts["word/numbering.xml"] = etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )

    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)


def test_parse_template_paragraphs_returns_anchor_for_blank_template(tmp_path: Path) -> None:
    template_path = tmp_path / "blank_template.docx"
    _create_blank_template(template_path)
    _remove_numbering_defaults(template_path)

    paragraphs = parse_template_paragraphs(str(template_path), use_cache=False)

    assert paragraphs == [
        {
            "index": 0,
            "display": "",
            "text": "空白模板起始段落",
            "style": None,
        }
    ]


def test_render_template_with_mappings_supports_blank_template_anchor(tmp_path: Path) -> None:
    template_path = tmp_path / "blank_template.docx"
    output_path = tmp_path / "result.docx"
    _create_blank_template(template_path)

    parsed_results = parse_template_paragraphs(str(template_path), use_cache=False)
    assert parsed_results[0]["index"] == 0

    render_template_with_mappings(
        str(template_path),
        str(output_path),
        [
            {
                "index": 0,
                "mode": "replace",
                "content_text": "Inserted body",
                "source_order": 0,
            }
        ],
        parsed_results,
    )

    result_doc = DocxDocument(output_path)
    texts = [(p.text or "").strip() for p in result_doc.paragraphs if (p.text or "").strip()]
    assert texts == ["Inserted body"]
    assert result_doc.sections[0].header.paragraphs[0].text == "Header only"
    assert result_doc.sections[0].footer.paragraphs[0].text == "Footer only"
