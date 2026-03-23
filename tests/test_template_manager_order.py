from pathlib import Path
import zipfile

from docx import Document as DocxDocument

from modules.template_manager import render_template_with_mappings


def _create_docx(path: Path, paragraphs: list[str]) -> None:
    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


def test_render_template_with_mappings_preserves_insert_after_row_order(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "result.docx"

    _create_docx(template_path, ["Anchor"])

    parsed_results = [{"index": 0, "display": "", "text": "Anchor"}]
    mappings = [
        {
            "index": 0,
            "mode": "insert_after",
            "content_text": "First",
            "source_order": 0,
        },
        {
            "index": 0,
            "mode": "insert_after",
            "content_text": "Second",
            "source_order": 1,
        },
    ]

    render_template_with_mappings(
        str(template_path),
        str(output_path),
        mappings,
        parsed_results,
    )

    result_doc = DocxDocument(output_path)
    texts = [(p.text or "").strip() for p in result_doc.paragraphs]
    texts = [text for text in texts if text]
    assert texts[:3] == ["Anchor", "First", "Second"]


def test_render_template_with_mappings_subdoc_placeholder_renders_as_paragraph_xml(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    content_path = tmp_path / "content.docx"
    output_path = tmp_path / "result.docx"

    _create_docx(template_path, ["Anchor"])
    _create_docx(content_path, ["Inserted body"])

    render_template_with_mappings(
        str(template_path),
        str(output_path),
        [
            {
                "index": 0,
                "mode": "insert_after",
                "content_docx_path": str(content_path),
                "source_order": 0,
            }
        ],
        [{"index": 0, "display": "", "text": "Anchor"}],
    )

    with zipfile.ZipFile(output_path, "r") as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")

    assert "<w:t xml:space=\"preserve\"><w:p>" not in document_xml
    assert "<w:t>Inserted body</w:t>" in document_xml
