from pathlib import Path

from docx import Document as DocxDocument

from modules.workflow import run_workflow


def _create_docx(path: Path, text: str = "source") -> None:
    doc = DocxDocument()
    doc.add_paragraph(text)
    doc.save(path)


def _extract_chapter_step(input_file: str, **extra_params):
    params = {
        "input_file": input_file,
        "target_chapter_section": "6.4.2",
        "target_subtitle": "",
        "ignore_toc": "true",
        "ignore_header_footer": "true",
    }
    params.update(extra_params)
    return {"type": "extract_word_chapter", "params": params}


def test_workflow_result_keeps_heading_chapter_numbers(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source.docx"
    _create_docx(src)

    def fake_extract_word_chapter(*_args, **kwargs):
        out_path = Path(str(kwargs.get("output_docx_path")))
        doc = DocxDocument()
        doc.add_paragraph("6.4.2 General description", style="Heading 2")
        doc.add_paragraph("Body A")
        doc.add_paragraph("8.1 Intended users", style="Heading 2")
        doc.add_paragraph("Body B")
        doc.save(out_path)
        return {"output_docx": str(out_path), "captured_titles": []}

    monkeypatch.setattr("modules.workflow.extract_word_chapter", fake_extract_word_chapter)

    steps = [_extract_chapter_step(str(src))]
    result = run_workflow(steps, str(tmp_path / "job_chapter_keep_numbers"))

    output_doc = DocxDocument(result["result_docx"])
    lines = [p.text.strip() for p in output_doc.paragraphs if p.text.strip()]

    assert "6.4.2 General description" in lines
    assert "8.1 Intended users" in lines


def test_workflow_result_keeps_plain_text_chapter_numbers(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source_plain.docx"
    _create_docx(src)

    def fake_extract_word_chapter(*_args, **kwargs):
        out_path = Path(str(kwargs.get("output_docx_path")))
        doc = DocxDocument()
        doc.add_paragraph("6.4.2 General description")
        doc.add_paragraph("Body A")
        doc.add_paragraph("8.1 Intended users")
        doc.add_paragraph("Body B")
        doc.save(out_path)
        return {"output_docx": str(out_path), "captured_titles": []}

    monkeypatch.setattr("modules.workflow.extract_word_chapter", fake_extract_word_chapter)

    steps = [_extract_chapter_step(str(src))]
    result = run_workflow(steps, str(tmp_path / "job_chapter_keep_numbers_plain"))

    output_doc = DocxDocument(result["result_docx"])
    lines = [p.text.strip() for p in output_doc.paragraphs if p.text.strip()]

    assert "6.4.2 General description" in lines
    assert "8.1 Intended users" in lines


def test_workflow_chapter_extract_triggers_figure_renumber(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source_with_figure.docx"
    _create_docx(src)

    def fake_extract_word_chapter(*_args, **kwargs):
        out_path = Path(str(kwargs.get("output_docx_path")))
        doc = DocxDocument()
        doc.add_paragraph("6.4.2 General description", style="Heading 2")
        doc.add_paragraph("Figure 2. Packaging overview")
        doc.add_paragraph("See Figure 2 for details.")
        doc.add_paragraph("8.1 Intended users", style="Heading 2")
        doc.add_paragraph("Figure 3. Device architecture")
        doc.add_paragraph("See Figure 3 for details.")
        doc.save(out_path)
        return {"output_docx": str(out_path), "captured_titles": []}

    monkeypatch.setattr("modules.workflow.extract_word_chapter", fake_extract_word_chapter)

    steps = [_extract_chapter_step(str(src))]
    result = run_workflow(steps, str(tmp_path / "job_chapter_figure_renumber"))

    output_doc = DocxDocument(result["result_docx"])
    lines = [p.text.strip() for p in output_doc.paragraphs if p.text.strip()]

    assert "Figure 1. Packaging overview" in lines
    assert "See Figure 1 for details." in lines
    assert "Figure 2. Device architecture" in lines
    assert "See Figure 2 for details." in lines


def test_workflow_figure_references_reset_mapping_after_heading_boundary(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source_with_repeated_local_figure.docx"
    _create_docx(src)

    def fake_extract_word_chapter(*_args, **kwargs):
        out_path = Path(str(kwargs.get("output_docx_path")))
        doc = DocxDocument()
        doc.add_paragraph("6.4.2 General description", style="Heading 2")
        doc.add_paragraph("The local figure is shown as Figure 1.")
        doc.add_paragraph("Figure 1. Packaging overview")
        doc.add_paragraph("8.1 Intended users", style="Heading 2")
        doc.add_paragraph("The local figure is shown as Figure 1.")
        doc.add_paragraph("Figure 1. Device architecture")
        doc.save(out_path)
        return {"output_docx": str(out_path), "captured_titles": []}

    monkeypatch.setattr("modules.workflow.extract_word_chapter", fake_extract_word_chapter)

    steps = [_extract_chapter_step(str(src))]
    result = run_workflow(steps, str(tmp_path / "job_chapter_figure_reference_reset"))

    output_doc = DocxDocument(result["result_docx"])
    lines = [p.text.strip() for p in output_doc.paragraphs if p.text.strip()]

    assert "The local figure is shown as Figure 1." in lines
    assert "Figure 1. Packaging overview" in lines
    assert "The local figure is shown as Figure 2." in lines
    assert "Figure 2. Device architecture" in lines


def test_workflow_removes_orphan_caption_only_paragraph(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source_orphan_caption.docx"
    _create_docx(src)

    def fake_extract_word_chapter(*_args, **kwargs):
        out_path = Path(str(kwargs.get("output_docx_path")))
        doc = DocxDocument()
        doc.add_paragraph("6.4.2 General description", style="Heading 2")
        doc.add_paragraph("Figure 1. Packaging overview")
        doc.add_paragraph("Figure 2.")
        doc.add_paragraph("Figure 3. Device architecture")
        doc.add_paragraph("Direct marking on product")
        doc.save(out_path)
        return {"output_docx": str(out_path), "captured_titles": []}

    monkeypatch.setattr("modules.workflow.extract_word_chapter", fake_extract_word_chapter)

    steps = [_extract_chapter_step(str(src))]
    result = run_workflow(steps, str(tmp_path / "job_orphan_caption_cleanup"))

    output_doc = DocxDocument(result["result_docx"])
    lines = [p.text.strip() for p in output_doc.paragraphs if p.text.strip()]

    assert "Figure 1. Packaging overview" in lines
    assert "Figure 2." not in lines
    assert "Figure 2. Device architecture" in lines
    assert "Direct marking on product" in lines
