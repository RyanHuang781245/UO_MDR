from pathlib import Path
import zipfile

from docx import Document as DocxDocument
from lxml import etree

from modules.workflow import run_workflow


def _create_docx(path: Path, text: str = "source") -> None:
    doc = DocxDocument()
    doc.add_paragraph(text)
    doc.save(path)


def _wqn(local_name: str) -> str:
    return "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}" + local_name


def _create_numbered_caption_fragment(path: Path, caption_text: str, reference_text: str) -> None:
    doc = DocxDocument()
    doc.add_paragraph(caption_text)
    doc.add_paragraph(reference_text)
    doc.save(path)

    with zipfile.ZipFile(path, "r") as zin:
        parts = {name: zin.read(name) for name in zin.namelist()}

    root = etree.fromstring(parts["word/document.xml"])
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = root.xpath("./w:body/w:p", namespaces=ns)
    caption = paragraphs[0]
    p_pr = caption.find("w:pPr", namespaces=ns)
    if p_pr is None:
        p_pr = etree.Element(_wqn("pPr"))
        caption.insert(0, p_pr)
    num_pr = etree.SubElement(p_pr, _wqn("numPr"))
    ilvl = etree.SubElement(num_pr, _wqn("ilvl"))
    ilvl.set(_wqn("val"), "0")
    num_id = etree.SubElement(num_pr, _wqn("numId"))
    num_id.set(_wqn("val"), "7")

    numbering_xml = parts.get("word/numbering.xml")
    if numbering_xml:
        numbering_root = etree.fromstring(numbering_xml)
    else:
        numbering_root = etree.Element(_wqn("numbering"), nsmap={"w": ns["w"]})

    abstract = etree.SubElement(numbering_root, _wqn("abstractNum"))
    abstract.set(_wqn("abstractNumId"), "51")
    lvl = etree.SubElement(abstract, _wqn("lvl"))
    lvl.set(_wqn("ilvl"), "0")
    start = etree.SubElement(lvl, _wqn("start"))
    start.set(_wqn("val"), "1")
    num_fmt = etree.SubElement(lvl, _wqn("numFmt"))
    num_fmt.set(_wqn("val"), "decimal")
    lvl_text = etree.SubElement(lvl, _wqn("lvlText"))
    lvl_text.set(_wqn("val"), "Figure %1.")

    num = etree.SubElement(numbering_root, _wqn("num"))
    num.set(_wqn("numId"), "7")
    abstract_ref = etree.SubElement(num, _wqn("abstractNumId"))
    abstract_ref.set(_wqn("val"), "51")

    parts["word/document.xml"] = etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )
    parts["word/numbering.xml"] = etree.tostring(
        numbering_root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)


def _create_two_numbered_caption_fragment(path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("Packaging overview")
    doc.add_paragraph("See Figure 3 for details.")
    doc.add_paragraph("Device architecture")
    doc.add_paragraph("Refer Figure 4 when reviewing setup.")
    doc.save(path)

    with zipfile.ZipFile(path, "r") as zin:
        parts = {name: zin.read(name) for name in zin.namelist()}

    root = etree.fromstring(parts["word/document.xml"])
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = root.xpath("./w:body/w:p", namespaces=ns)
    for idx in (0, 2):
        caption = paragraphs[idx]
        p_pr = caption.find("w:pPr", namespaces=ns)
        if p_pr is None:
            p_pr = etree.Element(_wqn("pPr"))
            caption.insert(0, p_pr)
        num_pr = etree.SubElement(p_pr, _wqn("numPr"))
        ilvl = etree.SubElement(num_pr, _wqn("ilvl"))
        ilvl.set(_wqn("val"), "0")
        num_id = etree.SubElement(num_pr, _wqn("numId"))
        num_id.set(_wqn("val"), "17")

    numbering_xml = parts.get("word/numbering.xml")
    if numbering_xml:
        numbering_root = etree.fromstring(numbering_xml)
    else:
        numbering_root = etree.Element(_wqn("numbering"), nsmap={"w": ns["w"]})

    abstract = etree.SubElement(numbering_root, _wqn("abstractNum"))
    abstract.set(_wqn("abstractNumId"), "61")
    lvl = etree.SubElement(abstract, _wqn("lvl"))
    lvl.set(_wqn("ilvl"), "0")
    start = etree.SubElement(lvl, _wqn("start"))
    start.set(_wqn("val"), "3")
    num_fmt = etree.SubElement(lvl, _wqn("numFmt"))
    num_fmt.set(_wqn("val"), "decimal")
    lvl_text = etree.SubElement(lvl, _wqn("lvlText"))
    lvl_text.set(_wqn("val"), "Figure %1.")

    num = etree.SubElement(numbering_root, _wqn("num"))
    num.set(_wqn("numId"), "17")
    abstract_ref = etree.SubElement(num, _wqn("abstractNumId"))
    abstract_ref.set(_wqn("val"), "61")

    parts["word/document.xml"] = etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )
    parts["word/numbering.xml"] = etree.tostring(
        numbering_root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)


def _extract_step(input_file: str, **extra_params):
    params = {
        "input_file": input_file,
        "target_chapter_section": "1.1",
        "target_caption_label": "Figure 1.",
    }
    params.update(extra_params)
    return {"type": "extract_specific_figure_from_word", "params": params}


def _extract_table_step(input_file: str, **extra_params):
    params = {
        "input_file": input_file,
        "target_chapter_section": "1.1",
        "target_caption_label": "Table 1.",
    }
    params.update(extra_params)
    return {"type": "extract_specific_table_from_word", "params": params}


def test_workflow_extract_figure_success_saves_fragment_and_result(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source.docx"
    _create_docx(src)
    captured: dict[str, object] = {}

    def fake_extract_specific_figure_from_word(*_args, **kwargs):
        captured["include_caption"] = kwargs.get("include_caption")
        captured["ignore_header_footer"] = kwargs.get("ignore_header_footer")
        output_docx_path = kwargs.get("output_docx_path")
        assert output_docx_path
        out_path = Path(str(output_docx_path))
        _create_docx(out_path, text="Figure 1. Demo")
        return {"ok": True, "reason": "ok"}

    monkeypatch.setattr(
        "modules.workflow.extract_specific_figure_from_word",
        fake_extract_specific_figure_from_word,
    )

    steps = [_extract_step(str(src), include_caption="false")]
    result = run_workflow(steps, str(tmp_path / "job_success"))
    figure_entry = next(e for e in result["log_json"] if e.get("type") == "extract_specific_figure_from_word")

    assert figure_entry["status"] == "ok"
    assert Path(figure_entry["output_docx"]).is_file()
    assert figure_entry["result"]["ok"] is True
    assert captured["include_caption"] is False
    assert captured["ignore_header_footer"] is True


def test_workflow_extract_figure_not_found_uses_result_reason(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source.docx"
    _create_docx(src)

    def fake_extract_specific_figure_from_word(*_args, **_kwargs):
        return {"ok": False, "reason": "figure_not_found"}

    monkeypatch.setattr(
        "modules.workflow.extract_specific_figure_from_word",
        fake_extract_specific_figure_from_word,
    )

    steps = [_extract_step(str(src))]
    result = run_workflow(steps, str(tmp_path / "job_not_found"))
    figure_entry = next(e for e in result["log_json"] if e.get("type") == "extract_specific_figure_from_word")

    assert figure_entry["status"] == "error"
    assert figure_entry["error"] == "figure_not_found"


def test_workflow_extract_figure_include_caption_defaults_true(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source.docx"
    _create_docx(src)
    captured: dict[str, object] = {}

    def fake_extract_specific_figure_from_word(*_args, **kwargs):
        captured["include_caption"] = kwargs.get("include_caption")
        output_docx_path = kwargs.get("output_docx_path")
        out_path = Path(str(output_docx_path))
        _create_docx(out_path, text="Figure 1. Default Caption")
        return {"ok": True, "reason": "ok"}

    monkeypatch.setattr(
        "modules.workflow.extract_specific_figure_from_word",
        fake_extract_specific_figure_from_word,
    )

    steps = [_extract_step(str(src))]
    result = run_workflow(steps, str(tmp_path / "job_default_caption"))
    figure_entry = next(e for e in result["log_json"] if e.get("type") == "extract_specific_figure_from_word")

    assert figure_entry["status"] == "ok"
    assert captured["include_caption"] is True


def test_workflow_extract_figure_forwards_title_and_index(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source.docx"
    _create_docx(src)
    captured: dict[str, object] = {}

    def fake_extract_specific_figure_from_word(*_args, **kwargs):
        captured["target_figure_title"] = kwargs.get("target_figure_title")
        captured["target_figure_index"] = kwargs.get("target_figure_index")
        output_docx_path = kwargs.get("output_docx_path")
        out_path = Path(str(output_docx_path))
        _create_docx(out_path, text="Figure result")
        return {"ok": True, "reason": "ok"}

    monkeypatch.setattr(
        "modules.workflow.extract_specific_figure_from_word",
        fake_extract_specific_figure_from_word,
    )

    steps = [
        _extract_step(
            str(src),
            target_caption_label="",
            target_figure_title="System architecture",
            target_figure_index="2",
        )
    ]
    result = run_workflow(steps, str(tmp_path / "job_title_index"))
    figure_entry = next(e for e in result["log_json"] if e.get("type") == "extract_specific_figure_from_word")

    assert figure_entry["status"] == "ok"
    assert captured["target_figure_title"] == "System architecture"
    assert captured["target_figure_index"] == "2"


def test_workflow_extract_figure_forwards_table_container_flag(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source.docx"
    _create_docx(src)
    captured: dict[str, object] = {}

    def fake_extract_specific_figure_from_word(*_args, **kwargs):
        captured["allow_table_figure_container"] = kwargs.get("allow_table_figure_container")
        output_docx_path = kwargs.get("output_docx_path")
        out_path = Path(str(output_docx_path))
        _create_docx(out_path, text="Figure result")
        return {"ok": True, "reason": "ok"}

    monkeypatch.setattr(
        "modules.workflow.extract_specific_figure_from_word",
        fake_extract_specific_figure_from_word,
    )

    steps = [_extract_step(str(src), allow_table_figure_container="true")]
    result = run_workflow(steps, str(tmp_path / "job_table_container"))
    figure_entry = next(e for e in result["log_json"] if e.get("type") == "extract_specific_figure_from_word")

    assert figure_entry["status"] == "ok"
    assert captured["allow_table_figure_container"] is True


def test_workflow_extract_table_include_caption_defaults_true(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source.docx"
    _create_docx(src)
    captured: dict[str, object] = {}

    def fake_extract_specific_table_from_word(*args, **kwargs):
        captured["include_caption"] = kwargs.get("include_caption")
        output_docx_path = kwargs.get("output_docx_path") or (args[1] if len(args) > 1 else None)
        out_path = Path(str(output_docx_path))
        _create_docx(out_path, text="Table 1. Default Caption")
        return {"ok": True, "reason": "ok"}

    monkeypatch.setattr(
        "modules.workflow.extract_specific_table_from_word",
        fake_extract_specific_table_from_word,
    )

    steps = [_extract_table_step(str(src))]
    result = run_workflow(steps, str(tmp_path / "job_table_default_caption"))
    table_entry = next(e for e in result["log_json"] if e.get("type") == "extract_specific_table_from_word")

    assert table_entry["status"] == "ok"
    assert captured["include_caption"] is True


def test_workflow_extract_figure_forwards_ignore_header_footer(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source.docx"
    _create_docx(src)
    captured: dict[str, object] = {}

    def fake_extract_specific_figure_from_word(*_args, **kwargs):
        captured["ignore_header_footer"] = kwargs.get("ignore_header_footer")
        output_docx_path = kwargs.get("output_docx_path")
        out_path = Path(str(output_docx_path))
        _create_docx(out_path, text="Figure result")
        return {"ok": True, "reason": "ok"}

    monkeypatch.setattr(
        "modules.workflow.extract_specific_figure_from_word",
        fake_extract_specific_figure_from_word,
    )

    steps = [_extract_step(str(src), ignore_header_footer="false")]
    result = run_workflow(steps, str(tmp_path / "job_figure_ignore_header_footer"))
    figure_entry = next(e for e in result["log_json"] if e.get("type") == "extract_specific_figure_from_word")

    assert figure_entry["status"] == "ok"
    assert captured["ignore_header_footer"] is False


def test_workflow_extract_table_forwards_ignore_header_footer(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source.docx"
    _create_docx(src)
    captured: dict[str, object] = {}

    def fake_extract_specific_table_from_word(*args, **kwargs):
        captured["ignore_header_footer"] = kwargs.get("ignore_header_footer")
        output_docx_path = kwargs.get("output_docx_path") or (args[1] if len(args) > 1 else None)
        out_path = Path(str(output_docx_path))
        _create_docx(out_path, text="Table result")
        return {"ok": True, "reason": "ok"}

    monkeypatch.setattr(
        "modules.workflow.extract_specific_table_from_word",
        fake_extract_specific_table_from_word,
    )

    steps = [_extract_table_step(str(src), ignore_header_footer="false")]
    result = run_workflow(steps, str(tmp_path / "job_table_ignore_header_footer"))
    table_entry = next(e for e in result["log_json"] if e.get("type") == "extract_specific_table_from_word")

    assert table_entry["status"] == "ok"
    assert captured["ignore_header_footer"] is False


def test_workflow_result_renumbers_figure_captions_and_references(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source.docx"
    _create_docx(src)
    call_idx = {"value": 0}

    def fake_extract_specific_figure_from_word(*_args, **kwargs):
        call_idx["value"] += 1
        output_docx_path = kwargs.get("output_docx_path")
        out_path = Path(str(output_docx_path))
        doc = DocxDocument()
        if call_idx["value"] == 1:
            doc.add_paragraph("Figure 8. Packaging overview")
            doc.add_paragraph("See Figure 8 for details.")
        else:
            doc.add_paragraph("Figure 3. Device architecture")
            doc.add_paragraph("Refer Figure 3 when reviewing setup.")
        doc.save(out_path)
        return {"ok": True, "reason": "ok"}

    monkeypatch.setattr(
        "modules.workflow.extract_specific_figure_from_word",
        fake_extract_specific_figure_from_word,
    )

    steps = [
        _extract_step(str(src), target_caption_label="Figure 8."),
        _extract_step(str(src), target_caption_label="Figure 3."),
    ]
    result = run_workflow(steps, str(tmp_path / "job_renumber_figure"))

    output_doc = DocxDocument(result["result_docx"])
    lines = [p.text.strip() for p in output_doc.paragraphs if p.text.strip()]

    assert "Figure 1. Packaging overview" in lines
    assert "See Figure 1 for details." in lines
    assert "Figure 2. Device architecture" in lines
    assert any(
        entry.get("type") == "postprocess_renumber_figure_table" and entry.get("status") == "ok"
        for entry in result["log_json"]
    )


def test_workflow_result_renumbers_numbering_style_figure_captions(
    tmp_path: Path,
    monkeypatch,
) -> None:
    src = tmp_path / "source.docx"
    _create_docx(src)
    def fake_extract_specific_figure_from_word(*_args, **kwargs):
        output_docx_path = kwargs.get("output_docx_path")
        out_path = Path(str(output_docx_path))
        _create_two_numbered_caption_fragment(out_path)
        return {"ok": True, "reason": "ok"}

    monkeypatch.setattr(
        "modules.workflow.extract_specific_figure_from_word",
        fake_extract_specific_figure_from_word,
    )

    steps = [_extract_step(str(src), target_caption_label="Figure 3.")]
    result = run_workflow(steps, str(tmp_path / "job_renumber_figure_numbered_style"))

    output_doc = DocxDocument(result["result_docx"])
    lines = [p.text.strip() for p in output_doc.paragraphs if p.text.strip()]

    assert "Figure 1. Packaging overview" in lines
    assert "See Figure 1 for details." in lines
    assert "Figure 2. Device architecture" in lines
